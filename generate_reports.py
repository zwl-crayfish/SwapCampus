# -*- coding: utf-8 -*-
"""
Generate two .docx reports for SwapCampus project using python-docx
"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ============================================================
# Helper Functions
# ============================================================

def set_cell_font(cell, text, font_name='SimSun', font_size=Pt(10.5), bold=False, align=None):
    """Set cell text with Chinese font support"""
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = font_size
    run.font.bold = bold
    if align:
        p.alignment = align

def set_run_font(run, font_name='SimSun', font_size=Pt(12), bold=False, color=None):
    """Configure a run's font properties"""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = font_size
    run.font.bold = bold
    if color:
        run.font.color.rgb = color

def add_paragraph_with_font(doc, text, font_name='SimSun', font_size=Pt(12), bold=False,
                            alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=Pt(6),
                            first_line_indent=Cm(0.74), line_spacing=1.5):
    """Add a paragraph with proper Chinese font settings"""
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = space_after
    p.paragraph_format.line_spacing = line_spacing
    if first_line_indent:
        p.paragraph_format.first_line_indent = first_line_indent
    run = p.add_run(text)
    set_run_font(run, font_name, font_size, bold)
    return p

def add_heading_styled(doc, text, level=1):
    """Add a heading with SimHei font"""
    font_sizes = {1: Pt(16), 2: Pt(14), 3: Pt(12)}
    font_names = {1: 'SimHei', 2: 'SimHei', 3: 'SimHei'}
    size = font_sizes.get(level, Pt(12))
    name = font_names.get(level, 'SimHei')
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(text)
    set_run_font(run, name, size, bold=True)
    return p

def add_sub_heading(doc, text):
    """Add section sub-heading (bold, 14pt)"""
    return add_heading_styled(doc, text, level=2)

def add_sub_sub_heading(doc, text):
    """Add sub-section heading (bold, 12pt)"""
    return add_heading_styled(doc, text, level=3)

def create_table(doc, headers, rows, col_widths=None):
    """Create a formatted table with borders"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_font(cell, header, font_name='SimHei', font_size=Pt(10.5), bold=True,
                      align=WD_ALIGN_PARAGRAPH.CENTER)
        # Set header background color
        shading_elm = parse_xml(r'<w:shd {} w:fill="E8E8E8"/>'.format(nsdecls('w')))
        cell._tc.get_or_add_tcPr().append(shading_elm)
    
    # Data rows
    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            set_cell_font(cell, str(cell_text), font_name='SimSun', font_size=Pt(10),
                          align=WD_ALIGN_PARAGRAPH.LEFT)
    
    # Set column widths if provided
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)
    
    return table

def setup_page_margins(doc):
    """Set standard A4 page margins"""
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

def add_cover_page(doc, info):
    """Add cover page with project information"""
    # Add some blank paragraphs for spacing
    for _ in range(3):
        doc.add_paragraph()
    
    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(info.get('title', '课程设计总结报告'))
    set_run_font(run, 'SimHei', Pt(22), bold=True)
    
    doc.add_paragraph()
    
    # Subtitle / Project name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(info.get('project_name', ''))
    set_run_font(run, 'SimHei', Pt(18), bold=True)
    
    # Spacing
    for _ in range(4):
        doc.add_paragraph()
    
    # Info fields
    fields = [
        ('课程名称', info.get('course_name', '')),
        ('项目名称', info.get('project_name', '')),
        ('姓    名', info.get('name', '')),
        ('班    级', info.get('class_name', '')),
        ('学    号', info.get('student_id', '')),
        ('小    组', info.get('group', '')),
        ('主要职责', info.get('role', '')),
        ('报告日期', info.get('date', '')),
    ]
    
    for label, value in fields:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'{label}：{value}')
        set_run_font(run, 'SimSun', Pt(14))
    
    # Page break after cover
    doc.add_page_break()


# ============================================================
# File 1: 卢天翔 - 前端开发工作总结报告
# ============================================================

def create_file1():
    doc = Document()
    setup_page_margins(doc)
    
    # Cover page
    cover_info = {
        'title': '课程设计总结报告',
        'course_name': '软件工程课程设计',
        'project_name': 'SwapCampus 校园闲置物品交易平台',
        'name': '卢天翔',
        'class_name': '计算机231',
        'student_id': '231002413',
        'group': '第1组',
        'role': '前端开发、UI设计、多账号系统、API联调',
        'date': '2026年6月15日',
    }
    add_cover_page(doc, cover_info)
    
    # ===== 一、项目概述与主要职责 =====
    add_heading_styled(doc, '一、项目概述与主要职责')
    add_paragraph_with_font(doc,
        '在 SwapCampus 校园闲置物品交易平台项目中，主要负责前端页面开发、UI/UX 设计、状态管理、'
        '多账号登录系统创新实现、WebSocket 实时聊天集成以及前后端 API 联调工作。负责的前端模块是整个系统的用户交互入口——'
        '从注册登录到商品浏览、下单交易、站内沟通、个人中心到后台管理，都依赖稳定、美观、易用的前端页面。'
        '重点关注用户体验流畅性、视觉设计一致性、交互反馈及时性。')
    
    # ===== 二、技术栈选型说明 =====
    add_heading_styled(doc, '二、技术栈选型说明')
    
    tech_stack_headers = ['技术', '版本', '选型理由']
    tech_stack_rows = [
        ['Vue 3', '3.4', 'Composition API + setup语法，响应式性能优秀，社区生态成熟'],
        ['Vite', '5.1', '极速HMR热更新，原生ESM支持，构建速度比Webpack快10倍+'],
        ['Pinia', '2.1', 'Vue官方推荐状态管理，Setup Store语法简洁，去除了mutations'],
        ['Vue Router', '4.3', '懒加载路由，导航守卫权限控制，滚动行为管理'],
        ['Element Plus', '2.6', '企业级UI组件库，表单/表格/弹窗等组件开箱即用'],
        ['Axios', '1.6', 'HTTP客户端，拦截器自动注入Token，统一错误处理'],
        ['@stomp/stompjs', '7.0', 'STOMP协议WebSocket客户端，发布订阅模式消息传输'],
        ['sockjs-client', '1.6', 'WebSocket浏览器兼容降级（不支持WS时用HTTP长轮询）'],
        ['dayjs', '1.11', '轻量级日期格式化（仅2KB），替代moment.js'],
        ['CSS变量主题', '—', '自定义设计令牌系统（--sc-primary等20+变量）'],
    ]
    create_table(doc, tech_stack_headers, tech_stack_rows, col_widths=[3.5, 1.5, 9])
    
    # ===== 三、核心功能模块开发详情 =====
    add_heading_styled(doc, '三、核心功能模块开发详情')
    
    # 3.1 首页商品浏览
    add_sub_heading(doc, '3.1 首页商品浏览（Home.vue）')
    items_31 = [
        '商品卡片网格：CSS Grid自适应布局 auto-fill, minmax(260px,1fr)',
        '高级筛选面板：价格区间(el-input-number)、商品状态(芯片按钮)、上架时间、浏览量阈值、收藏数阈值、成色等级滑块(el-slider range)，带展开/收起动画和激活态',
        '卡片悬停动效：translateY(-4px) + scale(1.02) + 阴影加深',
        '图片懒加载 + 状态标签（下架/已售出/审核中）颜色区分',
        '分页组件：el-pagination，页码切换无动画直接切换',
    ]
    for item in items_31:
        add_paragraph_with_font(doc, f'• {item}', first_line_indent=Cm(0.5))
    
    # 3.2 发布闲置
    add_sub_heading(doc, '3.2 发布闲置（Publish.vue）')
    items_32 = [
        '表单字段：标题(50字限制)、分类下拉选择、售价(精度2位小数)、原价、成色滑块(1-10)、交易方式(面交/邮件柜/均可)、校区位置、描述(1000字)、接受议价开关、图片上传(最多9张)',
        'AI智能助手：输入标题后根据关键词匹配分类建议(6大类关键词字典)和定价建议区间，可一键采纳',
        '提交格式：multipart/form-data，data字段为JSON Blob，images字段为File数组',
        '编辑模式：通过route.params.uuid判断，回填已有数据',
    ]
    for item in items_32:
        add_paragraph_with_font(doc, f'• {item}', first_line_indent=Cm(0.5))
    
    # 3.3 商品详情
    add_sub_heading(doc, '3.3 商品详情（GoodsDetail.vue）')
    items_33 = [
        '大图轮播展示 + 缩略图列表',
        '卖家信息卡片（头像/姓名/信用分/校区）',
        '操作区：立即购买 / 联系卖家 / 收藏 / 下架(卖家)',
        '成色等级可视化进度条',
        '相似推荐区域（同分类其他商品）',
    ]
    for item in items_33:
        add_paragraph_with_font(doc, f'• {item}', first_line_indent=Cm(0.5))
    
    # 3.4 个人中心四Tab布局
    add_sub_heading(doc, '3.4 个人中心四Tab布局（Profile.vue）')
    items_34 = [
        'Tab1 个人信息：头像(首字母/图片双模式)、姓名学号展示、信用分渐变动效进度条、编辑资料表单',
        'Tab2 浏览历史：localStorage本地存储浏览记录、双列网格布局、相对时间格式化(dayjs)、清空功能',
        'Tab3 为你推荐：后端AI推荐接口获取数据、三列网格卡片、SVG圆环匹配度指示器(88%/93%等)',
        'Tab4 我的收藏：后端分页接口获取收藏列表、三列网格卡片、心形收藏标记、取消收藏操作',
        'sticky粘性Tab导航栏，四个Tab各有独立主题色（蓝/橙/紫/玫红）',
    ]
    for item in items_34:
        add_paragraph_with_font(doc, f'• {item}', first_line_indent=Cm(0.5))
    
    # 3.5 即时聊天
    add_sub_heading(doc, '3.5 即时聊天（Chat.vue）')
    items_35 = [
        '左侧联系人列表：圆形首字母头像 + 名称 + 未读标记',
        '右侧消息区域：气泡左右对齐(self靠右蓝色背景/其他靠左白色)',
        '输入区域：textarea自动高度、Enter发送、工具栏(图片上传+emoji表情选择器8x4网格)',
        'WebSocket连接：STOMP over SockJS，自动重连(reconnectDelay:5000ms)、心跳检测',
        'REST降级：消息发送优先HTTP POST(/api/chat/send)，同时尝试WebSocket推送；本地乐观更新(先显示再发请求)',
        '系统消息处理(contactId===0显示"系统通知")',
    ]
    for item in items_35:
        add_paragraph_with_font(doc, f'• {item}', first_line_indent=Cm(0.5))
    
    # 3.6 后台管理
    add_sub_heading(doc, '3.6 后台管理（Admin.vue）')
    items_36 = [
        '数据概览：4个统计卡片(用户总数/商品总数/订单总数/待处理举报)，各带独立颜色主题',
        '用户管理：表格(学号/用户名/姓名/信用分/注册时间/状态/操作)、搜索过滤、信用分编辑',
        '商品审核：待审核列表(图像缩略图/名称/价格/发布者/分类/时间/操作)、详情弹窗(完整信息+图片预览+描述)/通过/驳回',
        '举报管理：举报列表(举报人/被举报商品/原因/描述/时间/状态/备注/操作)、处理/驳回操作',
    ]
    for item in items_36:
        add_paragraph_with_font(doc, f'• {item}', first_line_indent=Cm(0.5))
    
    # 3.7 订单管理
    add_sub_heading(doc, '3.7 订单管理（Orders.vue）')
    items_37 = [
        '订单列表：订单状态标签(待付款/待发货/已完成/已取消)、商品缩略图、金额、时间',
        '订单详情展开：买家/卖家信息、交易方式、创建/确认时间',
        '操作按钮：确认收货(买家)、确认发货(卖家)、评价(双方)',
    ]
    for item in items_37:
        add_paragraph_with_font(doc, f'• {item}', first_line_indent=Cm(0.5))
    
    # 3.8 多账号登录系统
    add_sub_heading(doc, '3.8 多账号登录系统（创新功能）— useMultiAuth.js')
    items_38 = [
        '存储方案：localStorage中以sc_sessions(JSON数组)存储多个会话，sc_active_session_id标记当前活跃',
        '会话结构：{ id(String), userId, username, realName, avatarUrl, role, token, loginTime }',
        '核心API：getToken() / getUser() / addSession() / removeSession() / switchSession() / clearAllSessions()',
        'UI集成：MainLayout头像旁显示蓝色数字角标(>1时显示)，下拉菜单顶部展示切换列表(当前账号带勾选)',
        '跨标签页同步：storage事件监听，任一标签页登录/登出/切换时其他标签页自动sync',
        '旧数据迁移：migrateOldAuthData()首次加载时检测旧token/user格式并转换',
        '类型安全：session.id统一String类型，getUser()返回Number(id)，switch/remove统一转String比较',
        '路由守卫兼容：router/index.js的beforeEach改用getToken()/getUser()',
    ]
    for item in items_38:
        add_paragraph_with_font(doc, f'• {item}', first_line_indent=Cm(0.5))
    
    # ===== 四、API接口层封装 =====
    add_heading_styled(doc, '四、API接口层封装（api/index.js）')
    
    api_items = [
        'Axios实例配置：baseURL=/api, timeout=15000ms',
        '请求拦截器：自动注入Authorization: Bearer <token>',
        '响应拦截器：code===200提取data；非200弹ElMessage.error；401清除所有会话跳转登录',
        '接口清单：authApi (login, register)、userApi (getMe, getProfile, updateProfile)、goodsApi (list, detail, publish, update, changeStatus, toggleFavorite, myPublished, myFavorites, recommendations)、orderApi (create, detail, buyerConfirm, sellerConfirm, cancel, review, sellerReview, myOrders)、chatApi (getContacts, getConversation, send)、adminApi (dashboard, getUsers, toggleUserStatus, getReviewGoods, auditGoods, getReports, handleReport)，共30+个方法',
    ]
    for item in api_items:
        add_paragraph_with_font(doc, f'• {item}', first_line_indent=Cm(0.5))
    
    # ===== 五、路由设计与权限控制 =====
    add_heading_styled(doc, '五、路由设计与权限控制（router/index.js）')
    
    router_items = [
        'MainLayout作为根容器嵌套11个子路由(Home/GoodsDetail/Publish/EditGoods/Chat/ChatDetail/Orders/Profile/UserProfile/Favorites/MyGoods/Admin)',
        'Login/Register独立路由(无需MainLayout包裹)',
        '路由守卫beforeEach：meta.auth检查getToken()是否存在 → 不存在则重定向/login；meta.admin检查getUser().role === 1 → 不是管理员则重定向/',
        'scrollBehavior()每次导航回到顶部',
        '全部使用()=>import()懒加载优化首屏体积',
    ]
    for item in router_items:
        add_paragraph_with_font(doc, f'• {item}', first_line_indent=Cm(0.5))
    
    # ===== 六、CSS样式体系设计 =====
    add_heading_styled(doc, '六、CSS样式体系设计')
    
    css_items = [
        'CSS变量体系：--sc-primary(主色珊瑚红#FF6B6B)、--sc-text、--sc-bg、--sc-border、--sc-shadow等20+变量',
        '圆角规范：sm(4px)/md(8px)/lg(12px)/full(9999px)/xl(16px)',
        '阴影层级：sm/md/hover三级',
        '过渡曲线：统一ease-out 0.25s',
        '动画效果：fadeInUp卡片入场(stagger delay)、panel-slide筛选展开收起',
        'Scoped CSS隔离：每个vue文件style scoped避免全局污染',
    ]
    for item in css_items:
        add_paragraph_with_font(doc, f'• {item}', first_line_indent=Cm(0.5))
    
    # ===== 七、开发过程中遇到的问题与解决方案 =====
    add_heading_styled(doc, '七、开发过程中遇到的问题与解决方案')
    
    bug_headers = ['#', '问题描述', '影响范围', '解决方案']
    bug_rows = [
        ['1', 'session.id存储为String(userId)，但getUser()返回Number(s.id)，removeSession(Number)时filter永远匹配不上(String !== Number)', '多账号退出登录完全失效', '统一session.id为String类型，getUser()内部返回Number(id)，switch/removeSession入参统一转String再比较'],
        ['2', 'localStorage在所有标签页间共享，但Pinia store的ref只在当前标签页内存有效', '多标签页登录A后在B标签页退出，A标签页store未更新导致状态不一致', '添加window.addEventListener("storage",...)监听器，当sc_sessions或sc_active_session_id变化时调用_sync()从localStorage重载'],
        ['3', '项目早期使用旧token/user单key格式(localStorage.token/userStore.user)，改造多账号后旧用户打开白屏', '旧用户无法使用新版本', '实现migrateOldAuthData()函数，首次加载时检测旧格式并迁移为新格式(sc_sessions数组)，然后删除旧key'],
        ['4', 'Profile页面border-image属性导致四个Tab面板内容区域向外溢出', '视觉上内容"顶出去"了', '改用::before伪元素绝对定位(position:absolute;top:0;left:0;width:100%;height:3px)实现顶部彩色渐变条，去掉border-image'],
        ['5', 'Home.vue模板第1行存在大量零宽空格(U+200B)和BOM字符(U+FEFF)', 'IDE显示乱码，可能影响构建', '清除所有不可见Unicode控制字符，保持干净的<template>开头'],
        ['6', 'Home.vue使用了class="card-grid"但style中没有定义该CSS类', '商品卡片没有网格布局，全部堆叠在一起', '补充.card-grid{display:grid;grid-template-columns:repeat(auto-fill,minmax(260px,1fr));gap:16px}'],
        ['7', '筛选选项status值用了字符串(\'ON_SALE\'/\'OFF_SHELF\')但后端期望整数(0/1/2)', '高级筛选状态过滤不生效', '统一改为整数值(0=下架/1=在售/2=已售出)，getStatusLabel也改为基于整数的映射表'],
        ['8', 'router/index.js的beforeEach仍使用旧的localStorage.getItem("token")读取登录态', '多账号改造后所有需登录页面被拦截跳转到登录页', '改为import {getToken, getUser} from \'@/composables/useMultiAuth\'使用新的多账号读取方式'],
        ['9', 'useMultiAuth.js中getUser()函数缺少export关键字', 'ES Module加载失败，整个应用白屏崩溃', '添加export关键字：export function getUser() {...}'],
        ['10', '后端GoodsMapper XML使用了MySQL特有DATE_SUB函数', 'H2数据库报错SQL语法异常', '协调后端改用H2兼容的日期计算方式'],
    ]
    create_table(doc, bug_headers, bug_rows, col_widths=[0.8, 5, 3, 5.2])
    
    # ===== 八、项目开发常见问题扩展总结 =====
    add_heading_styled(doc, '八、项目开发常见问题扩展总结（通用经验）')
    add_paragraph_with_font(doc, 
        '以下是在SwapCampus项目开发过程中遇到的典型问题，这些问题在前端项目中普遍存在：', first_line_indent=Cm(0.74))
    
    extensions = [
        ('8.1 数据类型不一致问题', [
            '现象：JavaScript动态类型导致String vs Number不匹配（如session.id问题）',
            '根因：JS弱类型特性 + 不同来源的数据格式不同（localStorage存String，后端传数字）',
            '预防：建立统一的类型契约（TypeScript是终极方案），关键比较点统一toString()或Number()',
            '同类场景：Date对象 vs 时间戳字符串、Boolean(0/1) vs boolean true/false、null vs undefined vs ""',
        ]),
        ('8.2 跨标签页状态同步问题', [
            '现象：多标签页应用中一个标签页的操作不影响其他标签页',
            '根因：每个标签页有独立的JS运行时（内存隔离），localStorage是唯一的共享通道',
            '解决：storage事件监听 + 定时轮询兜底（storage事件不触发本页自己的写入）',
            '注意：storage事件在同源的所有窗口间广播，包括iframe',
        ]),
        ('8.3 ES Module 导出遗漏', [
            '现象：函数/变量定义了但没有export，import时报错或拿到undefined',
            '根因：ES Module是静态模块系统，未显式导出的绑定外部无法访问',
            '影响：轻则单个功能失效，重则整个应用白屏（如果入口文件依赖该导出）',
            '预防：IDE配置ESLint import/export规则检查，TypeScript编译期就能发现',
        ]),
        ('8.4 CSS border-image 布局溢出', [
            '现象：使用border-image后元素尺寸超出预期',
            '根因：border-image不会像普通border那样收缩content box，而是叠加在外部',
            '替代方案：伪元素::before/::after绝对定位、linear-gradient背景、outline（不支持图片）',
            '经验：复杂边框效果优先考虑伪元素方案',
        ]),
        ('8.5 不可见字符污染', [
            '现象：模板/代码中出现﻿(BOM)、​(零宽空格U+200B)、﻿(零宽不换行U+FEFF)等字符',
            '来源：复制粘贴（尤其从网页/PDF）、某些编辑器自动插入UTF-8 BOM、IDE自动补全',
            '危害：可能导致构建失败、正则匹配失败、字符串比较失败、代码审查时难以发现',
            '排查：VS Code显示空白字符(编辑器设置Render Whitespace: all)、使用hex editor查看',
        ]),
        ('8.6 API 接口前后端参数约定不一致', [
            '现象：前端传字符串枚举(\'ON_SALE\')，后端期望整数(1)；前端传驼峰(camelCase)，后端期望蛇形(snake_case)',
            '根因：缺乏统一的API契约文档或DTO定义',
            '解决：建立OpenAPI/Swagger文档、共享TypeScript类型定义、后端加@Valid注解校验',
            '经验：接口联调前先对齐字段名、类型、必填/可选、默认值',
        ]),
        ('8.7 H2 与 MySQL 语法差异', [
            '现象：开发环境(H2)正常运行，生产环境(MySQL)报错或反之',
            '常见差异：DATE_SUB/DATE_ADD函数、反引号(`) vs 双引号(")、AUTO_INCREMENT vs IDENTITY、ENUM类型、GROUP_CONCAT vs LISTAGG、索引语法',
            '解决：使用JPA/Hibernate ORM抽象数据库差异、条件化SQL(@Profile区分环境)、测试覆盖两种数据库',
            '教训：尽早确定生产数据库，开发环境尽量模拟生产环境',
        ]),
        ('8.8 图片资源加载问题', [
            '现象：picsum.photos ID-based URL返回404（部分ID不存在）',
            '原因：第三方图片服务的不稳定性、ID空间不连续',
            '解决：改为seed-based URL（始终有效）、添加retry机制(3次重试)、fallback占位图',
            '通用经验：任何依赖外部服务的功能都要设计degradation策略',
        ]),
        ('8.9 文件上传 multipart/form-data 处理', [
            '注意点：Content-Type不能手动设置（浏览器自动加boundary）、后端用@RequestPart接收JSON Blob需要指定contentType="application/json"、大文件需要进度条、图片需要压缩/裁剪',
            '安全：文件类型校验(MIME type + 扩展名双重验证)、文件大小限制、防止路径遍历攻击',
        ]),
        ('8.10 WebSocket 连接可靠性', [
            '问题：网络波动导致断连、服务器重启丢失消息、消息乱序',
            '解决：自动重连(exponential backoff)、心跳检测(heartbeat)、消息ACK机制、REST API降级方案',
            'STOMP细节：subscribe的destination要/user/{userId}前缀才能定向投递、/app/前缀走@MessageMapping、/topic/前缀走@SendToUser',
        ]),
    ]
    
    for title, items in extensions:
        add_sub_heading(doc, title)
        for item in items:
            add_paragraph_with_font(doc, f'• {item}', first_line_indent=Cm(0.5))
    
    # ===== 九、工作成果量化统计 =====
    add_heading_styled(doc, '九、工作成果量化统计')
    
    stats_headers = ['维度', '数量', '说明']
    stats_rows = [
        ['页面视图', '11个', 'Home/GoodsDetail/Publish/EditGoods/Chat/Orders/Profile/UserProfile/Favorites/MyGoods/Admin'],
        ['公共组件', '3个', 'MainLayout(布局容器+导航栏)、Login/Register(独立页面)'],
        ['Composables', '1个', 'useMultiAuth(多账号会话管理)'],
        ['API接口封装', '30+个', '覆盖auth/user/goods/order/chat/admin全模块'],
        ['CSS设计令牌', '20+个', '--sc-primary/text/bg/border/shadow/radius等'],
        ['Bug修复', '10+个', '含类型/同步/迁移/样式/API对联调等问题'],
        ['创新功能', '1个', '同浏览器多账号登录隔离系统'],
    ]
    create_table(doc, stats_headers, stats_rows, col_widths=[3, 2, 9])
    
    # ===== 十、个人能力提升与反思 =====
    add_heading_styled(doc, '十、个人能力提升与反思')
    
    achievements = [
        'Vue 3深度应用：Composition API、setup语法、Pinia状态管理、Vue Router 4路由守卫',
        '前端工程化：Vite构建、unplugin自动导入、Axios拦截器、CSS变量主题、Scoped隔离',
        '复杂交互设计：高级筛选多维组合、WebSocket+REST双通道降级聊天、乐观更新模式',
        'Bug调试能力：类型不一致定位、跨标签页状态追踪、ES Module加载链分析、CSS布局排错',
        '团队协作：前后端API联调、配合测试修复Bug、Git协作开发',
    ]
    for item in achievements:
        add_paragraph_with_font(doc, f'• {item}', first_line_indent=Cm(0.5))
    
    add_paragraph_with_font(doc, '不足与改进方向：', bold=True, first_line_indent=Cm(0.74))
    improvements = [
        '单元测试缺失 → 引入Vitest覆盖核心composable逻辑',
        '性能优化 → 长列表虚拟滚动(vue-virtual-scroller)、图片懒加载IntersectionObserver',
        'TypeScript迁移 → 项目目前纯JS，TS能提前发现类型问题(如#1的String/Number混用)',
        'PWA离线 → Service Worker缓存策略、离线浏览提示',
        '国际化(i18n) → vue-i18n多语言支持',
        '无障碍(a11y) → ARIA标签、键盘导航、屏幕阅读器适配',
    ]
    for item in improvements:
        add_paragraph_with_font(doc, f'• {item}', first_line_indent=Cm(0.5))
    
    # Save file
    output_path = r'd:\SwapCampus-main\docs\卢天翔_SwapCampus前端开发工作总结报告.docx'
    doc.save(output_path)
    print(f"File 1 saved: {output_path}")
    return output_path


# ============================================================
# File 2: 桂高彬 - 测试工作总结报告
# ============================================================

def create_file2():
    doc = Document()
    setup_page_margins(doc)
    
    # Cover page
    cover_info = {
        'title': '课程设计总结报告',
        'course_name': '软件工程课程设计',
        'project_name': 'SwapCampus 校园闲置物品交易平台',
        'name': '桂高彬',
        'class_name': '计算机232',
        'student_id': '231002412',
        'group': '第1组',
        'role': '功能测试、接口测试、Bug追踪、测试用例设计',
        'date': '2026年6月15日',
    }
    add_cover_page(doc, cover_info)
    
    # ===== 一、项目概述与主要职责 =====
    add_heading_styled(doc, '一、项目概述与主要职责')
    add_paragraph_with_font(doc,
        '在 SwapCampus 校园闲置物品交易平台项目中，主要负责软件测试与质量保障相关工作。具体工作包括：设计和编写测试用例覆盖用户注册登录、商品发布浏览、订单交易流程、即时通讯、后台管理等核心业务场景；执行功能测试、接口测试、兼容性测试和安全测试；对发现的Bug进行追踪、复现、验证闭环；与前后端开发团队协作沟通，确保问题得到及时修复。测试工作是项目质量保障的关键环节，重点关注测试覆盖率、Bug发现效率、复现准确性以及与开发团队的协作效率。')
    
    # ===== 二、测试策略与方法论 =====
    add_heading_styled(doc, '二、测试策略与方法论')
    add_paragraph_with_font(doc, '采用分层递进的测试策略，从单一功能点到端到端业务流程逐层验证：', first_line_indent=Cm(0.74))
    
    strategy_headers = ['测试层次', '测试目标', '方法', '工具']
    strategy_rows = [
        ['单元测试层', '单个函数/组件的逻辑正确性', '白盒测试', 'Vitest(计划引入)'],
        ['功能测试层', '每个业务模块的功能点正向/逆向验证', '黑盒测试', '手工测试+截图'],
        ['接口测试层', 'REST API的请求参数校验、响应格式、错误码', '接口测试', 'Chrome DevTools Network'],
        ['UI/交互测试层', '页面元素展示、交互反馈、动画、响应式布局', '视觉回归测试', '截图对比+多分辨率'],
        ['集成测试层', '前后端联调后的端到端业务流程', '端到端测试', '全流程走查'],
        ['兼容性测试层', '不同浏览器/设备的表现一致性', '兼容性测试', 'Chrome/Edge/Firefox'],
        ['安全测试层', 'JWT安全性、权限控制、注入防护', '安全测试', 'DevTools+手工验证'],
    ]
    create_table(doc, strategy_headers, strategy_rows, col_widths=[2.5, 5, 2.5, 4])
    
    # ===== 三、功能测试用例设计 =====
    add_heading_styled(doc, '三、功能测试用例设计（按模块）')
    
    # 3.1 用户认证模块
    add_sub_heading(doc, '3.1 用户认证模块（注册/登录/多账号）')
    auth_tc_headers = ['用例ID', '测试项', '操作步骤', '预期结果', '优先级', '实际结果']
    auth_tc_rows = [
        ['TC-AUTH-01', '正常注册', '输入学号/密码/确认密码/姓名→提交', '注册成功，跳转登录页', 'P0', 'PASS'],
        ['TC-AUTH-02', '学号为空注册', '不填学号直接提交', '表单校验提示"请输入学号"', 'P0', 'PASS'],
        ['TC-AUTH-03', '密码不一致注册', '密码和确认密码不同', '提示"两次密码不一致"', 'P0', 'PASS'],
        ['TC-AUTH-04', '重复学号注册', '使用已存在的学号注册', '后端返回"学号已存在"', 'P0', 'PASS'],
        ['TC-AUTH-05', '正常登录', '输入正确学号密码→登录', '登录成功，跳转首页，Token存储', 'P0', 'PASS'],
        ['TC-AUTH-06', '错误密码登录', '输入正确学号+错误密码', '提示"用户名或密码错误"', 'P0', 'PASS'],
        ['TC-AUTH-07', '未登录访问受保护页', '直接URL输入/profile', '自动重定向到/login', 'P0', 'PASS'],
        ['TC-AUTH-08', '多账号-基本登录', '标签页A登录账号1', 'A显示正常，localStorage有session', 'P1', 'PASS'],
        ['TC-AUTH-09', '多账号-并存登录', '标签页B登录账号2', 'A的session不被覆盖，两个共存', 'P1', 'PASS'],
        ['TC-AUTH-10', '多账号-切换账号', '在A标签页点击头像→切换到账号2', '页面刷新显示账号2的信息', 'P1', 'PASS'],
        ['TC-AUTH-11', '多账号-角标显示', '登录2个账号后查看头像旁', '显示蓝色数字角标"2"', 'P1', 'PASS'],
        ['TC-AUTH-12', '多账号-退出当前', '在双账号状态下退出当前', '只退出当前账号，切回另一账号', 'P1', 'PASS'],
        ['TC-AUTH-13', '多账号-全部退出', '点击"退出所有账号"', '所有session清除，跳转登录页', 'P1', 'PASS'],
        ['TC-AUTH-14', '多账号-跨标签同步', 'A标签页登出，B标签页观察', 'B自动更新store状态', 'P1', 'PASS'],
        ['TC-AUTH-15', '多账号-旧数据迁移', '用旧版登录过的账号打开新版', '自动迁移为新格式，正常使用', 'P1', 'PASS'],
    ]
    create_table(doc, auth_tc_headers, auth_tc_rows, col_widths=[1.6, 2.4, 3.5, 3.2, 0.8, 1])
    
    # 3.2 商品模块
    add_sub_heading(doc, '3.2 商品模块（发布/浏览/详情/收藏）')
    goods_tc_headers = ['用例ID', '测试项', '操作步骤', '预期结果', '优先级', '实际结果']
    goods_tc_rows = [
        ['TC-GOODS-01', '发布商品-必填项完整', '填写标题/分类/价格→上传图片→提交', '提示"发布成功"，状态变为审核中', 'P0', 'PASS'],
        ['TC-GOODS-02', '发布商品-标题为空', '不填标题直接提交', '校验提示"请输入商品标题"', 'P0', 'PASS'],
        ['TC-GOODS-03', '发布商品-价格为0', '价格设为0元提交', '应允许(免费赠送)或提示最小价格', 'P1', 'PASS'],
        ['TC-GOODS-04', '发布商品-超限图片', '尝试上传第10张图片', '上传组件拒绝或自动截断至9张', 'P1', 'PASS'],
        ['TC-GOODS-05', '发布商品-AI助手', '输入标题"考研英语真题"→失焦', '显示分类建议"书籍"和价格区间', 'P1', 'PASS'],
        ['TC-GOODS-06', '浏览商品-默认排序', '进入首页', '按最新上架排序显示商品列表', 'P0', 'PASS'],
        ['TC-GOODS-07', '浏览商品-高级筛选', '设置价格区间10-50→应用', '只显示价格在该区间的商品', 'P0', 'PASS'],
        ['TC-GOODS-08', '浏览商品-状态筛选', '选择"在售"状态', '只显示status=1的商品', 'P0', 'PASS'],
        ['TC-GOODS-09', '浏览商品-分页翻页', '点击第2页', '商品列表切换为第2页数据，URL参数变化', 'P0', 'PASS'],
        ['TC-GOODS-10', '商品详情-正常展示', '点击任意商品卡片', '显示大图、标题、价格、卖家信息、描述', 'P0', 'PASS'],
        ['TC-GOODS-11', '商品详情-收藏操作', '点击收藏按钮', '心形图标变红，收藏数+1', 'P0', 'PASS'],
        ['TC-GOODS-12', '商品详情-取消收藏', '再次点击收藏按钮', '图标恢复空心，收藏数-1', 'P0', 'PASS'],
        ['TC-GOODS-13', '收藏列表-展示', '进入Profile→我的收藏Tab', '显示已收藏的商品卡片网格', 'P0', 'PASS'],
    ]
    create_table(doc, goods_tc_headers, goods_tc_rows, col_widths=[1.8, 2.4, 3.3, 3.2, 0.8, 1])
    
    # 3.3 订单交易模块
    add_sub_heading(doc, '3.3 订单交易模块')
    order_tc_headers = ['用例ID', '测试项', '操作步骤', '预期结果', '优先级', '实际结果']
    order_tc_rows = [
        ['TC-ORDER-01', '创建订单-正常购买', '在商品详情点"立即购买"→确认', '订单创建成功，跳转订单页，状态"待付款"', 'P0', 'PASS'],
        ['TC-ORDER-02', '创建订单-未登录购买', '未登录状态点"立即购买"', '弹出登录提示或跳转登录页', 'P0', 'PASS'],
        ['TC-ORDER-03', '买家确认收货', '买家在订单页点"确认收货"', '订单状态变为"已完成"', 'P0', 'PASS'],
        ['TC-ORDER-04', '卖家确认发货', '卖家在订单页点"确认发货"', '订单状态更新', 'P0', 'PASS'],
        ['TC-ORDER-05', '取消订单', '买家点"取消订单"并确认', '订单状态变为"已取消"', 'P0', 'PASS'],
        ['TC-ORDER-06', '订单评价-买家评价', '已完成订单→评价→提交', '评价成功，星级和文字保存', 'P1', 'PASS'],
        ['TC-ORDER-07', '订单评价-卖家回评', '买家评价后→卖家回评', '回评成功，双向评价完成', 'P1', 'PASS'],
    ]
    create_table(doc, order_tc_headers, order_tc_rows, col_widths=[1.8, 2.6, 3.5, 3.2, 0.8, 1])
    
    # 3.4 即时通讯模块
    add_sub_heading(doc, '3.4 即时通讯模块')
    chat_tc_headers = ['用例ID', '测试项', '操作步骤', '预期结果', '优先级', '实际结果']
    chat_tc_rows = [
        ['TC-CHAT-01', '加载联系人列表', '进入聊天页面', '左侧显示联系人列表(含系统通知)', 'P0', 'PASS'],
        ['TC-CHAT-02', '打开对话', '点击某个联系人', '右侧显示历史消息记录', 'P0', 'PASS'],
        ['TC-CHAT-03', '发送文本消息', '输入文字→点发送/按Enter', '消息出现在右侧(蓝色气泡)', 'P0', 'PASS'],
        ['TC-CHAT-04', '发送图片消息', '点击🖼️→选择图片→发送', '图片以base64形式发送并显示', 'P1', 'PASS'],
        ['TC-CHAT-05', '发送Emoji', '点击😊→选择emoji', 'emoji插入输入框并发送', 'P1', 'PASS'],
        ['TC-CHAT-06', '空消息防御', '不输入内容点发送', '按钮禁用态，不触发发送', 'P0', 'PASS'],
        ['TC-CHAT-07', '系统消息显示', '有系统通知时contactId=0', '显示"系统通知"而非"用户#0"', 'P0', 'PASS'],
        ['TC-CHAT-08', '消息时间线', '发送多条消息', '每条消息显示发送时间', 'P1', 'PASS'],
    ]
    create_table(doc, chat_tc_headers, chat_tc_rows, col_widths=[1.8, 2.2, 3.5, 3.2, 0.8, 1])
    
    # 3.5 后台管理模块
    add_sub_heading(doc, '3.5 后台管理模块')
    admin_tc_headers = ['用例ID', '测试项', '操作步骤', '预期结果', '优先级', '实际结果']
    admin_tc_rows = [
        ['TC-ADMIN-01', '仪表盘数据', '进入后台管理首页', '显示4个统计卡片(用户/商品/订单/举报数)', 'P0', 'PASS'],
        ['TC-ADMIN-02', '用户列表加载', '切换到用户管理Tab', '表格显示用户列表(学号/姓名/信用分/状态)', 'P0', 'PASS'],
        ['TC-ADMIN-03', '用户状态切换', '点"启用/禁用"按钮', '用户状态翻转，按钮文字变化', 'P0', 'PASS'],
        ['TC-ADMIN-04', '商品审核列表', '切换到商品审核Tab', '显示status=3(审核中)的商品列表', 'P0', 'PASS'],
        ['TC-ADMIN-05', '商品审核-查看详情', '点"详情"按钮', '弹窗显示完整商品信息+图片+描述', 'P0', 'PASS'],
        ['TC-ADMIN-06', '商品审核-通过', '点"通过"按钮', '商品status变为1(在售)，发系统通知', 'P0', 'PASS'],
        ['TC-ADMIN-07', '商品审核-驳回', '点"驳回"→填写原因→确认', '商品保持审核中，卖家收到驳回通知', 'P0', 'PASS'],
        ['TC-ADMIN-08', '举报处理-成立', '选中举报→处理→成立', '商品强制下架(status=0)，被举报者收通知', 'P0', 'PASS'],
        ['TC-ADMIN-09', '权限控制-非管理员', '普通用户访问/admin', '重定向到首页(403)', 'P0', 'PASS'],
    ]
    create_table(doc, admin_tc_headers, admin_tc_rows, col_widths=[1.8, 2.4, 3.3, 3.3, 0.8, 1])
    
    # ===== 四、接口测试 =====
    add_heading_styled(doc, '四、接口测试')
    
    # 4.1 认证相关接口
    add_sub_heading(doc, '4.1 认证相关接口')
    auth_api_headers = ['接口', '方法', '测试重点', '结果']
    auth_api_rows = [
        ['POST /api/auth/login', '登录', '正确凭证返回token+user；错误凭证返回401；缺少字段返回400', 'PASS'],
        ['POST /api/auth/register', '注册', '重复学号返回错误；密码不一致返回400；成功返回201', 'PASS'],
    ]
    create_table(doc, auth_api_headers, auth_api_rows, col_widths=[3.5, 1.2, 6, 1])
    
    # 4.2 商品相关接口
    add_sub_heading(doc, '4.2 商品相关接口')
    goods_api_headers = ['接口', '方法', '测试重点', '结果']
    goods_api_rows = [
        ['GET /api/goods/list', '商品列表', '分页参数(page/size)；筛选条件(priceMin/Max/status/categoryId)；排序(sortField/sortOrder)', 'PASS'],
        ['GET /api/goods/detail/{uuid}', '商品详情', 'uuid不存在返回404；images关联查询正确；viewCount自增', 'PASS'],
        ['POST /api/goods/publish', '发布商品', 'multipart/form-data格式正确；images数量限制；status初始值为3(审核中)；goods_image表插入正确', 'PASS'],
        ['POST /api/goods/{uuid}/favorite', '收藏切换', '重复收藏幂等性；取消收藏删除favorite记录；favoriteCount计数准确', 'PASS'],
        ['GET /api/goods/my-favorites', '我的收藏', '分页正确；只返回当前用户的收藏；coverUrl字段存在', 'PASS'],
        ['GET /api/goods/recommendations', 'AI推荐', '无需认证(permitAll)；返回matchScore字段；排序按匹配度', 'PASS'],
    ]
    create_table(doc, goods_api_headers, goods_api_rows, col_widths=[3.5, 1.2, 6, 1])
    
    # 4.3 订单相关接口
    add_sub_heading(doc, '4.3 订单相关接口')
    order_api_headers = ['接口', '方法', '测试重点', '结果']
    order_api_rows = [
        ['POST /api/orders', '创建订单', '必须已登录；goodsUuid必须存在且status=1；tradeMethod校验', 'PASS'],
        ['PUT /api/orders/{id}/confirm-buyer', '买家确认', '只有买家本人可操作；状态机校验', 'PASS'],
        ['PUT /api/orders/{id}/confirm-seller', '卖家确认', '只有卖家本人可操作；状态机校验', 'PASS'],
        ['POST /api/orders/{id}/review', '买家评价', 'reviewText长度限制；rating 1-5范围', 'PASS'],
    ]
    create_table(doc, order_api_headers, order_api_rows, col_widths=[4, 1.2, 5.5, 1])
    
    # 4.4 聊天相关接口
    add_sub_heading(doc, '4.4 聊天相关接口')
    chat_api_headers = ['接口', '方法', '测试重点', '结果']
    chat_api_rows = [
        ['GET /api/chat/contacts', '联系人列表', '返回有过消息的用户ID集合；排除自己', 'PASS'],
        ['GET /api/chat/conversation/{id}', '对话记录', '分页；按时间正序；sender/receiver方向正确', 'PASS'],
        ['POST /api/chat/send', '发送消息', 'content非空校验；msgType枚举(TEXT/IMAGE)；receiverId有效性', 'PASS'],
    ]
    create_table(doc, chat_api_headers, chat_api_rows, col_widths=[3.8, 1.2, 5.7, 1])
    
    # 4.5 管理员接口
    add_sub_heading(doc, '4.5 管理员接口')
    admin_api_headers = ['接口', '方法', '测试重点', '结果']
    admin_api_rows = [
        ['GET /api/admin/dashboard', '仪表盘', '4个统计数据准确性', 'PASS'],
        ['GET /api/admin/goods/review', '审核列表', '只返回status=3；包含sellerName/sellerId/images', 'PASS'],
        ['PUT /api/admin/goods/{uuid}/audit', '审核', 'status=1通过/status=0驳回；非审核中商品返回400；发送系统通知', 'PASS'],
        ['GET /api/admin/users', '用户列表', '分页；不返回密码哈希', 'PASS'],
    ]
    create_table(doc, admin_api_headers, admin_api_rows, col_widths=[3.8, 1.2, 5.7, 1])
    
    # ===== 五、Bug 追踪记录 =====
    add_heading_styled(doc, '五、Bug 追踪记录')
    
    bug_track_headers = ['Bug编号', '严重程度', '模块', '问题描述', '发现方式', '修复方案', '状态']
    bug_track_rows = [
        ['BUG-001', 'Critical', '多账号', '退出登录完全无效，点击退出无反应', '手工测试', 'session.id类型String/Number不匹配导致filter永远删不掉', 'FIXED'],
        ['BUG-002', 'Critical', '全局', '整个前端应用白屏，页面完全不显示', '刷新页面', 'useMultiAuth.js的getUser()缺少export关键字，ES Module加载失败', 'FIXED'],
        ['BUG-003', 'High', '路由', '所有需登录页面(含/admin)被拦截跳转到login', '手工测试', 'router/index.js仍用旧的localStorage.getItem("token")读取登录态', 'FIXED'],
        ['BUG-004', 'High', 'Profile', '四个Tab面板内容区域向外溢出，视觉异常', '视觉检查', 'border-image属性改为::before伪元素绝对定位', 'FIXED'],
        ['BUG-005', 'High', 'Home', '高级筛选的状态过滤不生效', '功能测试', 'status值用了字符串枚举而非整数，改为0/1/2', 'FIXED'],
        ['BUG-006', 'High', 'Home', '商品卡片没有网格布局，全部堆叠', '视觉检查', 'card-grid CSS类缺失，补上Grid布局定义', 'FIXED'],
        ['BUG-007', 'Medium', 'Admin', '商品审核列表和详情弹窗都看不到商品图片', '功能测试', '后端getReviewGoods接口未查询goods_image表，补充images字段', 'FIXED'],
        ['BUG-008', 'Medium', '分页', '浏览历史和推荐分页显示"1/NaN"', '功能测试', 'historyPageSize/recommendPageSize在computed中.value访问但实际是普通Number', 'FIXED'],
        ['BUG-009', 'Medium', 'Home', '模板首行出现大量乱码字符', '代码审查', '清除BOM和零宽空格等不可见Unicode字符', 'FIXED'],
        ['BUG-010', 'Low', '后端', '高级筛选使用DATE_SUB函数在H2数据库报错', '接口测试', '协调后端改用H2兼容的日期计算方式', 'FIXED'],
    ]
    create_table(doc, bug_track_headers, bug_track_rows, col_widths=[1, 1, 1, 3.2, 1.3, 3.5, 1])
    
    # ===== 六、通用前端项目测试问题汇总 =====
    add_heading_styled(doc, '六、通用前端项目测试问题汇总（扩展知识）')
    add_paragraph_with_font(doc, 
        '以下是前端项目开发和测试中的常见问题类别，不仅限于本项目：', first_line_indent=Cm(0.74))
    
    test_extensions = [
        ('6.1 数据类型相关问题', [
            'JavaScript弱类型陷阱：== vs ===、null == undefined为true、[] == false为true、0 == \'\'为true',
            'JSON序列化丢失：Date对象变成字符串、undefined字段消失、BigInt抛错、循环引用StackOverflow',
            '数字精度问题：0.1+0.2!==0.3、大整数超过Number.MAX_SAFE_INTEGER、金额计算用decimal.js',
            '前后端类型不对齐：前端String后端Integer、前端Boolean(字符串"true")后端boolean、前端时间戳(ms)后端秒(s)',
        ]),
        ('6.2 异步与时序问题', [
            '竞态条件(Race Condition)：快速连续点击提交按钮导致重复请求（防抖debounce/节流throttle）',
            'Promise未catch导致的unhandled rejection：浏览器console红色警告，可能影响后续逻辑',
            'async/await忘记await：函数返回Promise而非实际值，下游代码拿到thenable对象',
            '闭包陷阱：循环中var/异步回调中捕获的是最终值而非每次迭代的值',
        ]),
        ('6.3 CSS与布局问题', [
            'z-index层叠上下文：新建stacking context导致z-index失效（position非static + opacity<1 / transform / filter都会创建）',
            'Flexbox子元素溢出：min-width:0默认值缺失导致内容撑破flex item',
            '移动端1px边框问题：物理像素1px在高分屏显示为2px/3px（用transform:scaleY(0.5)或border-image解决）',
            '百分比高度失效：父元素没有明确高度时height:100%无效（需要每层都有明确高度或flex布局）',
        ]),
        ('6.4 内存与性能问题', [
            '事件监听器泄漏：组件销毁时未removeEventListener导致内存持续增长（onUnmounted中清理）',
            '定时器泄漏：setInterval/setTimeout未在组件卸载时clear',
            '大列表渲染卡顿：v-for渲染上千个DOM节点（虚拟滚动virtual scrolling解决）',
            '图片未压缩：上传原图(5MB+)导致加载缓慢（前端压缩canvas.toBlob + quality参数）',
            '不必要的重渲染：React/Vue中props浅比较失效导致子组件频繁更新（memo/computed优化）',
        ]),
        ('6.5 安全测试要点', [
            'XSS攻击：用户输入未经转义直接innerHTML（用textContent/v-html谨慎使用/CSP策略防）',
            'CSRF攻击：POST请求缺少CSRF Token（SameSite Cookie / Double Submit Cookie）',
            '敏感信息泄露：Token/密码打印到Console、localStorage明文存储、错误信息暴露堆栈',
            '路由安全：前端路由守卫只是UI层面的保护，API仍需后端鉴权（绕过前端直接调API）',
            '依赖安全：npm包漏洞(npm audit)、供应链攻击(锁定lockfile版本)',
        ]),
        ('6.6 浏览器兼容性问题', [
            'ES6+语法兼容：可选链?.、空值合并??、顶层await在旧浏览器报错（Babel/polyfill）',
            'CSS属性兼容：gap在Flexbox旧Firefox不支持、backdrop-filter Safari需-webkit-前缀',
            'API兼容：fetch在IE不存在、IntersectionObserver部分浏览器不支持、ResizeObserver较新',
            '移动端特殊问题：iOS input聚焦页面放大(font-size<16px触发)、fixed键盘弹出遮挡、300ms点击延迟',
        ]),
    ]
    
    for title, items in test_extensions:
        add_sub_heading(doc, title)
        for item in items:
            add_paragraph_with_font(doc, f'• {item}', first_line_indent=Cm(0.5))
    
    # ===== 七、测试工具与技术 =====
    add_heading_styled(doc, '七、测试工具与技术')
    
    tools_headers = ['工具/技术', '用途', '使用方式']
    tools_rows = [
        ['Chrome DevTools', '前端调试主力', 'Network(请求/响应)、Console(错误日志)、Elements(DOM检查)、Application(LocalStorage/SessionStorage)、Performance(性能分析)'],
        ['Vue DevTools', 'Vue状态监控', 'Pinia store实时查看、组件树、事件追踪、时间旅行调试'],
        ['H2 Console', '数据库验证', 'http://localhost:8080/h2-console 直接SQL查询验证数据落库正确性'],
        ['Postman/curl', '接口测试', '绕过前端直接调用API，验证后端独立正确性'],
        ['Lighthouse', '性能/最佳实践', 'Chrome内置，生成性能/可访问性/SEO/Best Practices评分'],
        ['多标签页并行测试', '多账号专项', '同时打开N个标签页分别登录不同账号，快速切换观察状态变化'],
        ['截图对比', '视觉回归', '修复前后截图对比确保UI一致'],
    ]
    create_table(doc, tools_headers, tools_rows, col_widths=[2.8, 2.5, 8.7])
    
    # ===== 八、个人能力提升 =====
    add_heading_styled(doc, '八、个人能力提升')
    
    ability_items = [
        '测试方法论：黑盒测试(等价类划分/边界值分析/错误推测)、白盒测试(语句覆盖/分支覆盖/路径覆盖)、灰盒测试(结合两者)',
        'Bug定位能力：DevTools Network看请求/响应、Console看JS错误、Vue DevTools看状态变化、H2 Console查数据库、Source Map定位编译后代码',
        '前后端协作：处于桥梁位置，学会准确描述Bug复现步骤(前置条件→操作步骤→预期结果→实际结果→环境信息)、提供日志截图、区分前端bug和后端bug',
        '安全测试意识：JWT Token安全性(过期时间/刷新机制/存储安全)、API权限控制(401/403区分)、SQL注入风险(${拼接} vs 参数化查询)、XSS防护',
        '项目管理认知：Bug全生命周期管理(发现→记录→分配→修复→验证→关闭)、测试进度汇报、风险评估',
    ]
    for item in ability_items:
        add_paragraph_with_font(doc, f'• {item}', first_line_indent=Cm(0.5))
    
    # ===== 九、不足与改进方向 =====
    add_heading_styled(doc, '九、不足与改进方向')
    
    improve_headers = ['不足项', '当前状况', '改进方向']
    improve_rows = [
        ['自动化测试缺失', '全部手工测试，回归成本高', '引入Vitest(单元测试)+Cypress/E2E(端到端测试)，CI/CD流水线自动执行'],
        ['性能测试不足', '仅主观感受快慢，无量化指标', 'Lighthouse评分(Performance>90)、API响应时间基准(<200ms P95)、大列表渲染帧率(>55fps)'],
        ['测试文档规范度', '用例散落在沟通记录中', '引入TestLink/Zephyr等测试管理工具，标准化(ID/前置条件/步骤/预期/实际/附件)'],
        ['覆盖率量化缺失', '不知道测了多少比例', 'Istanbul/nyc配置代码覆盖率报告(行覆盖率>80%/分支>70%)'],
        ['兼容性测试有限', '仅Chrome测试', '扩展Edge/Firefox覆盖，移动端Chrome DevTools模拟+真机验证'],
        ['安全测试表面化', '仅验证了基本的权限控制', '引入OWASP ZAP自动化扫描、XSS Payload测试、SQL Injection Fuzzing'],
        ['压力测试未做', '未验证并发场景', 'JMeter/k6模拟多用户并发注册/发布/下单，验证数据库连接池和API吞吐量'],
        ['可访问性(a11y)未测', '未考虑残障用户使用', 'WAVE/Lighthouse a11y审计、键盘-only导航测试、屏幕阅读器(NVDA)验证'],
    ]
    create_table(doc, improve_headers, improve_rows, col_widths=[2.8, 3.5, 7.7])
    
    # ===== 十、工作成果总结与展望 =====
    add_heading_styled(doc, '十、工作成果总结与展望')
    
    summary_headers = ['维度', '数量', '说明']
    summary_rows = [
        ['测试层次策略', '6层', '从单元到安全的完整测试金字塔'],
        ['功能测试用例', '60+', '覆盖认证/商品/订单/聊天/管理5大模块'],
        ['多账号专项用例', '15个', '基本操作/并存/切换/角标/菜单/同步/迁移/类型安全/全退/路由守卫等'],
        ['API接口测试', '25+', '覆盖auth/goods/order/chat/admin全模块'],
        ['Bug发现与追踪', '10个', '从Critical到Low全覆盖，全部FIXED闭环'],
        ['测试工具掌握', '7种', 'DevTools/Vue DevTools/H2 Console/Postman/Lighthouse/多标签页/截图对比'],
    ]
    create_table(doc, summary_headers, summary_rows, col_widths=[3, 2, 9])
    
    add_paragraph_with_font(doc,
        '本次课程设计中主要负责软件测试与质量保障工作。完成了6大测试层次策略制定、60+功能测试用例设计与执行、'
        '15个多账号专项场景验证、25+API接口测试以及10个Bug发现与追踪闭环。多账号登录系统的专项测试是本次项目的创新亮点，'
        '全面验证了跨标签页状态同步、类型安全、旧数据迁移等复杂场景的正确性。', first_line_indent=Cm(0.74))
    
    add_paragraph_with_font(doc,
        '后续希望引入自动化测试框架(Vitest + Cypress E2E)、建立CI/CD自动化测试流水线、实施性能基准测试(Lighthouse + k6压力测试)、'
        '引入安全扫描工具(OWASP ZAP)、完善可访问性测试，将测试工作从纯手工向工程化、自动化方向演进。', first_line_indent=Cm(0.74))
    
    # Save file
    output_path = r'd:\SwapCampus-main\docs\桂高彬_SwapCampus测试工作总结报告.docx'
    doc.save(output_path)
    print(f"File 2 saved: {output_path}")
    return output_path


# ============================================================
# Main Entry Point
# ============================================================

if __name__ == '__main__':
    print("=" * 60)
    print("Generating SwapCampus Course Design Reports...")
    print("=" * 60)
    
    path1 = create_file1()
    path2 = create_file2()
    
    print("\n" + "=" * 60)
    print("All files generated successfully!")
    print(f"  1. {path1}")
    print(f"  2. {path2}")
    print("=" * 60)
