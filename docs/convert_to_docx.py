#!/usr/bin/env python3
"""
基于模板生成 D-02~D-05 Word 文档
- 保留模板全部格式（字体/行距/页边距/页眉页脚/样式）
- 在每个章节标题后的空段落中填入实际内容
"""
import os, copy
from lxml import etree
from docx import Document
from docx.oxml import OxmlElement
from docx.shared import Pt

WML_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
DOCS_DIR = '/Users/admin/Desktop/SwapCampus/docs'
TPL_DIR = os.path.join(DOCS_DIR, '产出模板')
OUT_DIR = DOCS_DIR

# ═══════════════════ 内容数据 ═══════════════════

D02 = {
'1.1  编写目的':
'本文档旨在完整定义 SwapCampus 校园闲置物品交易平台的软件需求，作为后续概要设计、详细设计、编码实现、测试验收各阶段的基线依据。预期读者包括：指导教师（评审需求完整性）、开发团队（指导设计与编码）、测试人员（编写测试用例与验收标准）、后续维护者（理解系统功能边界）。',
'1.2  项目背景':
'北京林业大学校内每年产生大量闲置物品（教材、电子产品、生活用品等），当前通过微信群进行散点式交易，存在信息碎片化、信用难追溯、交易效率低等问题。学校学生处希望建设一个面向本校师生的闲置物品交易平台，以提高校内资源流转率，同时为同学提供安全、便捷的交易体验。本项目是《软件工程（课程设计）》课程项目（选题编号 T-02），要求以 3-4 人团队，在 2 周集中实习期内，完成从立项到答辩的完整软件工程实践。',
'1.3  术语与缩略语':
'SRS — Software Requirements Specification，软件需求规格说明书\nJWT — JSON Web Token，无状态鉴权令牌\nMinIO — 开源对象存储服务\nSTOMP — Streaming Text Oriented Messaging Protocol，WebSocket 子协议\nC2C — Consumer to Consumer，个人对个人交易模式\n面交 — 买卖双方线下当面交易\n邮件柜 — 校内自助快递柜中转交易\n信用分 — 用户信誉评价（0-100分，默认80，交易完成+2，违规-10）\nCRUD — 数据库增删改查基本操作',
'1.4  参考资料':
'1. GB/T 8567—2006 计算机软件文档编制规范\n2. IEEE Std 830-1998 软件需求规格说明书推荐做法\n3. 《软件工程（课程设计）》课程设计任务书 v1.0, 2026-05\n4. T-02 SwapCampus 选题指南 v1.0, 2026-05\n5. Spring Boot 3.2 Reference Documentation\n6. Vue 3 Official Guide\n7. MyBatis-Plus Official Documentation',
'2.1  产品愿景与定位':
'SwapCampus 是一个面向北京林业大学师生的 C2C 校园闲置物品交易平台。用户可通过 Web 端发布闲置物品、浏览购买二手商品、与买卖双方实时沟通，并基于学号实名认证与信用积分体系建立可信交易社区。核心价值主张：让校园闲置物品流转更安全、更高效、更可信。',
'2.2  产品功能概要':
'F01 用户注册与登录 — 学号实名注册 + JWT 鉴权，支持买家/卖家/管理员三角色\nF02 个人主页与信用分 — 个人信息管理、信用积分展示（0-100分）\nF03 商品发布 — 多图上传（MinIO存储，≤9张）、分类选择、成色标注、价格设定\nF04 商品浏览与搜索 — 分类筛选、关键词全文搜索、按价格/时间排序\nF05 商品收藏 — 收藏/取消收藏、查看收藏列表\nF06 下单交易 — 选择交易方式（面交/邮件柜）、填写面交地点时间\nF07 交易确认 — 买卖双方各自确认收货，双方确认后信用分自动+2\nF08 订单评价 — 交易完成后1-5星评分和文字评价\nF09 站内即时通讯 — WebSocket STOMP 实时聊天，支持文字和图片消息\nF10 商品审核 — 管理员审核商品，可下架违规商品\nF11 用户管理 — 管理员查看用户列表，禁用/启用账号\nF12 举报处理 — 用户举报不良商品/用户 → 管理员审核处理或驳回\nF13 数据仪表盘 — 管理员统计总览（用户总数/商品总数/订单总数/待处理举报数）',
'2.3  用户特征':
'买家：频次约每周3-5次，具备基本计算机操作技能，核心关注点为搜得到、看得清、聊得通、买得放心。卖家：频次约每周1-3次，具备基本计算机操作技能，核心关注点为发布快、曝光高、交易顺利。管理员：每日使用，中等技术水平，核心关注点为维护社区秩序、及时处理违规、查看系统数据。',
'2.4  运行环境':
'服务端操作系统：Linux / macOS / Windows\nJava 运行环境：JDK 17+\n数据库：MySQL 8.0 (InnoDB引擎)\n缓存中间件：Redis 7\n对象存储：MinIO\n客户端浏览器：Google Chrome 90+ / Microsoft Edge 90+ / Apple Safari 15+\n网络环境：校园网或公共互联网',
'2.5  约束条件':
'技术约束：后端必须使用 Spring Boot 3.x，前端必须使用 Vue 3，数据库必须使用 MySQL 8。\n时间约束：编码实现阶段仅 4 天（D6-D9），须合理分配开发任务。\n团队约束：团队规模 3-4 人，角色可兼任。\n代码量约束：有效代码行数 ≥ 3000 行。\n测试约束：单元测试覆盖率 ≥ 60%，核心模块（Auth/Goods/Order）≥ 80%。\n合规约束：文档遵循 GB/T 8567—2006 规范，代码遵循项目编码规范。\n部署约束：须提供 Docker Compose 配置文件，实现一键部署。',
'2.6  假设与依赖':
'A01 假设所有用户拥有有效的学校学号，可通过学号进行实名认证。\nA02 假设用户具备基本的网页浏览器使用能力。\nA03 假设服务器可以连接互联网（用于 MinIO 图片服务的外部访问）。\nA04 假设校内邮件柜系统可以接入（选做功能的前置条件）。\nD01 系统依赖 MySQL 8.0 提供数据持久化服务。\nD02 系统依赖 MinIO 提供图片对象存储服务。\nD03 系统依赖 Redis 7 提供缓存和会话管理服务。',
'3.1  用例图':
'主用例图包含三大参与者（Actor）：\n买家 — 可执行用例：浏览商品、搜索商品、收藏商品、下单购买、即时聊天、确认收货、评价订单。\n卖家 — 可执行用例：发布商品、管理商品（编辑/下架）、即时聊天、确认交易。\n管理员 — 可执行用例：用户管理、商品审核、举报处理、查看数据仪表盘。\n所有参与者共享注册和登录用例。系统边界为 SwapCampus 平台。',
'3.2  用例详述（每个核心用例展开）':
'UC-01 用户注册：参与者为未注册用户，前置条件无。主流程：用户进入注册页面 → 输入学号/用户名/密码/真实姓名/手机号(选填)/邮箱(选填) → 系统校验学号和用户名唯一性 → BCrypt加密密码 → 创建User和Wallet记录（默认信用分80）→ 生成JWT Token → 自动登录跳转首页。备选流程：学号已存在 → 提示"该学号已注册"；用户名已存在 → 提示"该用户名已存在"。后置条件：用户创建成功，JWT Token写入前端localStorage。\n\nUC-02 用户登录：参与者为已注册用户，前置条件为已注册。主流程：用户输入用户名和密码 → 系统验证凭证 → 校验账号状态（status=1正常）→ 返回JWT Token和用户信息 → 跳转首页。异常流程：用户名或密码错误 → 提示"用户名或密码错误"；账号已被禁用 → 提示"账号已被禁用，请联系管理员"。\n\nUC-03 发布闲置商品：参与者为卖家（已登录用户），前置条件为已登录。主流程：用户点击"发布闲置"按钮 → 填写商品信息（标题/分类/价格/原价(选填)/成色1-10/交易方式/校区位置/描述）→ 上传商品图片（最多9张，单张≤5MB）→ 点击提交 → 系统生成商品UUID → 保存Goods记录 → 逐张上传图片到MinIO → 保存GoodsImage记录 → 商品状态设为"在售"。业务规则：图片最多9张，单张不超过5MB；标题1-50字；成色1-10（10为全新）。\n\nUC-04 浏览和搜索商品：参与者为所有用户（含未登录访客），前置条件无。主流程：用户进入首页 → 浏览商品卡片列表 → 可选按分类筛选（7大分类）→ 可选输入关键词搜索（MySQL FULLTEXT全文检索）→ 可选按价格/发布时间排序 → 点击商品卡片查看详情页。\n\nUC-05 下单购买：参与者为买家（已登录），前置条件为已登录、商品状态为"在售"、买家与卖家非同一人。主流程：用户在商品详情页点击"立即购买"→ 弹出下单对话框 → 选择交易方式（面交/邮件柜）→ 若选面交则填写面交地点和时间 → 确认下单 → 系统创建订单记录（status=0待确认）→ 商品状态自动变为"已售出"。异常流程：商品不存在→提示错误；购买自己的商品→提示"不能购买自己的商品"。\n\nUC-06 双方确认交易：参与者为买家+卖家，前置条件为订单状态为"待确认"。主流程：买卖双方线下完成面交/邮件柜取件后 → 买家在订单页点击"确认收货"→ 卖家在订单页点击"确认交易"→ 双方均确认后订单状态变为"已完成"→ 系统自动给买卖双方各增加2点信用分。备选流程：任一方可在确认前点击"取消订单"，订单变为"已取消"，商品恢复为"在售"状态。\n\nUC-07 站内即时通讯：参与者为买家+卖家（已登录），前置条件为已登录。主流程：用户在商品详情页点击"联系卖家"（或在聊天页选择联系人）→ 进入聊天界面 → 输入文字消息 → 点击发送 → 系统通过 WebSocket STOMP 协议将消息实时推送到接收方浏览器 → 接收方收到消息通知。性能要求：消息端到端延迟 ≤ 200ms。\n\nUC-08 管理员处理举报：参与者为管理员，前置条件为管理员已登录。主流程：管理员进入后台管理 → 点击"举报管理" → 查看举报列表（含举报人、原因、描述、被举报商品/用户）→ 审核举报内容 → 选择"处理"（下架商品/禁用用户）或"驳回"→ 系统记录处理结果、处理人和处理时间。',
'4.1  功能列表':
'FR-01 学号注册：用户输入学号/用户名/密码/姓名进行注册，系统校验学号和用户名唯一性后创建账号（BCrypt加密密码），默认信用分80。优先级：高。来源：选题指南M1。\nFR-02 JWT登录：已注册用户通过用户名+密码登录，系统验证后返回JWT Token（24h有效）。优先级：高。来源：选题指南M1。\nFR-03 信用积分体系：每个用户有关联的信用分（0-100），交易完成双方各+2分，违规被举报成立-10分，管理员可手动调整。优先级：中。来源：选题指南M1。\nFR-04 商品信息发布：卖家填写标题/分类/价格/成色/描述等字段，发布闲置商品。标题1-50字，成色1-10级。优先级：高。来源：选题指南M2。\nFR-05 多图上传：发布商品时可上传最多9张图片，支持jpg/png格式，单张不超过5MB，存储到MinIO对象存储。优先级：高。来源：选题指南M2。\nFR-06 商品编辑与下架：卖家可编辑自己发布的商品信息，或将商品下架（状态变为0）。优先级：高。来源：选题指南M2。\nFR-07 分类浏览：用户可按7大商品分类（教材教辅/电子数码/生活用品/运动户外/服饰美妆/图书音像/其他闲置）筛选浏览。优先级：高。来源：选题指南M3。\nFR-08 关键词搜索：用户输入关键词，系统使用MySQL FULLTEXT全文索引在商品标题和描述中搜索匹配结果。优先级：高。来源：选题指南M3。\nFR-09 排序筛选：搜索结果支持按价格升序/降序、按发布时间排序。优先级：中。来源：选题指南M3。\nFR-10 创建订单：买家选择商品后，选择交易方式（面交/邮件柜），填写面交信息，确认下单。下单后商品自动变为"已售出"。优先级：高。来源：选题指南M4。\nFR-11 双方确认收货：买卖双方各自在系统中确认交易完成，双方确认后订单完成，信用分自动增加。优先级：高。来源：选题指南M4。\nFR-12 订单评价：买家在交易完成后可对商品进行1-5星评分和文字评价。优先级：中。来源：选题指南M4。\nFR-13 收藏商品：用户可收藏感兴趣的商品，也可取消收藏，在个人中心查看收藏列表。优先级：中。来源：团队定义。\nFR-14 WebSocket实时聊天：买卖双方可通过站内即时通讯实时沟通，使用STOMP over WebSocket协议。优先级：高。来源：选题指南M5。\nFR-15 图文消息：聊天支持发送文字消息和图片消息。优先级：中。来源：选题指南M5。\nFR-16 用户管理：管理员可查看所有用户列表，对违规用户执行禁用/启用操作。优先级：高。来源：选题指南M6。\nFR-17 商品审核：管理员可审核商品，将违规商品下架。优先级：高。来源：选题指南M6。\nFR-18 举报处理：管理员查看用户提交的举报，审核后可选择"处理"或"驳回"。优先级：高。来源：选题指南M6。\nFR-19 数据仪表盘：管理员首页展示系统关键统计数据（用户总数/商品总数/订单总数/待处理举报数）。优先级：中。来源：团队定义。',
'4.2  功能详述':
'FR-01 用户注册：输入为学号、用户名、密码、真实姓名、手机号(选填)、邮箱(选填)。处理流程：(1)校验学号唯一性(userMapper.findByStudentId)(2)校验用户名唯一性(userMapper.findByUsername)(3)BCrypt加密密码(passwordEncoder.encode)(4)创建User记录(userMapper.insert)(5)创建Wallet记录(默认balance=0,points=0)(6)生成JWT Token。输出为TokenResponse对象(accessToken,userId,username,realName,role)。异常：学号重复→400"该学号已注册"；用户名重复→400"该用户名已存在"。\n\nFR-04 商品发布：输入为标题、分类ID、价格、原价(选填)、成色1-10、交易方式(FACE/LOCKER/BOTH)、校区位置、描述、图片文件数组。处理流程：(1)IdUtil.fastSimpleUUID()生成商品UUID(2)构建Goods实体，计算折扣价=price*discount/100(3)goodsMapper.insert(goods)(4)循环图片数组，逐个调用fileService.uploadImage()上传到MinIO(5)每次上传成功后goodsImageMapper.insert保存图片URL。输出为Goods实体对象。异常：图片超过9张→前端限制；单张超过5MB→400；参数校验失败→422。\n\nFR-14 即时通讯：输入为发送者ID、接收者ID、消息内容、消息类型(TEXT/IMAGE)、关联商品UUID(可选)。处理流程：(1)生成消息UUID(2)构建Message实体(isRead=0)(3)messageMapper.insert(message)保存到数据库(4)通过SimpMessagingTemplate.convertAndSendToUser()推送到接收者的/user/{receiverId}/queue/chat队列(5)前端STOMP订阅该队列实时接收。性能要求：端到端延迟≤200ms。异常处理：WebSocket连接断开时，消息仍保存到数据库，用户重连后通过REST API获取历史消息。',
'5  非功能性需求':
'性能要求：商品列表接口P95响应时间≤500ms；即时消息端到端延迟≤200ms；图片上传单张处理时间≤3s。可用性要求：开发阶段系统可在本地正常运行，Docker Compose一键启动所有服务。安全性要求：密码采用BCrypt加密存储；采用JWT无状态鉴权（24h过期）；所有SQL使用MyBatis-Plus参数化查询防止注入攻击；敏感信息（密码/Token）不输出到日志；管理员接口使用@PreAuthorize方法级权限控制。可维护性要求：关键公开方法须有Javadoc注释；Git提交遵循Conventional Commits规范。兼容性要求：支持Chrome/Edge/Safari最新2个主要版本。可扩展性要求：架构预留≥3个扩展点（个性化推荐系统、邮件柜接入、AI自动定价建议）。',
'6.1  用户接口（UI）':
'系统为 Web SPA（单页应用），共12个主要页面：\n首页(/) — 顶部搜索框 + 分类导航栏 + 商品卡片网格 + 底部分页器\n商品详情(/goods/:uuid) — 左侧图片轮播 + 右侧商品信息(价格/成色/交易方式/校区) + 卖家信息卡片 + 收藏/联系卖家/立即购买操作按钮\n登录(/login) — 居中卡片式登录表单（用户名+密码），渐变紫色背景\n注册(/register) — 居中卡片式注册表单（学号+用户名+密码+姓名+手机号+邮箱）\n发布商品(/publish) — 表单页面（标题/分类下拉/价格数字/成色滑块/交易方式单选/校区/描述文本域）+ Element Plus Upload图片上传组件\n即时聊天(/chat) — 左侧联系人列表 + 右侧聊天窗口（消息气泡 + 底部输入框）\n订单管理(/orders) — Tab切换"我购买的"/"我卖出的"，订单卡片列表含状态标签和操作按钮\n个人中心(/profile) — 头像+信用分+个人信息+编辑表单\n我的发布(/my-goods) — 表格列表展示，每行含编辑/下架操作按钮\n我的收藏(/favorites) — 收藏商品卡片网格\n后台管理(/admin) — 仪表盘统计卡片 + 用户管理表格 + 举报管理表格\n所有页面采用 Element Plus 组件库，中文本地化，响应式布局适配不同屏幕宽度。',
'6.2  硬件接口':
'N/A。本系统为纯软件系统，运行于标准 x86 服务器或个人计算机，无专用硬件设备接口。所有外部交互通过标准网络协议（HTTP/WebSocket/JDBC）完成。',
'6.3  软件接口（第三方 API）':
'MySQL 8.0：用途为数据持久化存储，通过 JDBC 协议连接（com.mysql.cj.jdbc.Driver），默认端口3306。\nRedis 7：用途为缓存服务，通过 RESP 协议连接，默认端口6379，集成方式为 Spring Data Redis。\nMinIO：用途为图片对象存储，通过 S3 兼容 API 连接（io.minio:minio 8.5.7），默认端口9000（API）和9001（控制台）。\n以上均为本地部署的开源中间件，无外部第三方云服务 API 依赖。',
'6.4  通信接口':
'REST API：协议为 HTTP/HTTPS，数据格式为 JSON，前后端通过26个 REST 端点进行数据交互（完整端点列表见概要设计说明书§8.1）。统一响应格式为 {"code":200,"message":"success","data":{...}}，错误码体系包括 200/400/401/403/404/422/500。\nWebSocket：协议为 STOMP over SockJS，端点路径 /ws/chat，用于即时通讯模块的实时双向消息推送。客户端通过 SockJS 连接，使用 STOMP 帧协议进行发布/订阅。',
'7  数据需求':
'初始用户规模：3-5人（灰度测试用户）\n种子商品数据：≥200条（覆盖7大分类）\n图片存储量预估：初期≤500MB（按每商品平均2张图、每张约1MB计算）\n数据保留期限：完整课程周期（14天实习期+答辩评审期）\n备份策略：每日手动执行 mysqldump 导出；SQL建表脚本纳入 Git 版本控制；MinIO 图片数据本地文件系统保留。\n数据安全：密码 BCrypt 加密；生产环境密钥通过环境变量注入。',
'附录 A  待解决问题清单 (TBD)':
'TBD-01 邮件柜接入：需确认学校邮件柜系统是否提供 API 接口，以及接口协议和认证方式。负责人：待分配。截止日期：D4 设计阶段结束前。\nTBD-02 信用分扣减规则：除违规-10分外，是否需要针对其他行为（如频繁取消订单、恶意差评）设计更多扣分规则。负责人：待分配。截止日期：D4。\nTBD-03 选做功能范围：选题指南中列出的选做功能（推荐系统、AI定价、签到任务、邮件柜接入）是否需要在本期实现。负责人：组长。截止日期：D4。',
'附录 B  需求评审会议纪要':
'（本条在需求评审会议后填写）\n评审日期：2026年___月___日\n与会人员：\n评审结论：\n主要问题与修改意见：\n决议事项：\n后续行动项：',
}

# ─── D-03 ───
D03 = {
'1.1  编写目的与读者':
'本文档描述 SwapCampus 系统的总体架构设计，回答"系统由哪些组件构成、组件之间如何协同、为什么采用这种划分方式"三个核心问题。预期读者为开发团队成员（指导详细设计和编码实现）以及指导教师（评审架构设计合理性）。文档采用 C4 模型进行架构描述，包含4+1视图中的逻辑视图、进程视图、部署视图和场景视图，并附有至少3份架构决策记录（ADR）。',
'1.2  与 SRS 的对应关系':
'本概要设计说明书的各章节与需求规格说明书（D-02 SRS）之间存在以下追溯关系：\n§2 总体设计（含架构图和技术选型）→ SRS §2 总体描述、§2.4 运行环境、§2.5 约束条件\n§3 系统结构（模块划分与依赖关系）→ SRS §4 功能性需求（19项功能的模块归属）\n§4 运行视图（顺序图和状态机）→ SRS §3 用例视图（8个核心用例的运行时体现）\n§6 数据视图（高层ER图）→ SRS §7 数据需求\n§8 接口设计（REST API 摘要和 WebSocket）→ SRS §6 外部接口需求\n§10 安全设计→ SRS §5 非功能性需求（安全性部分）',
'2.1  设计目标与原则':
'前后端分离原则：前端 Vue 3 SPA 与后端 Spring Boot REST API 完全解耦，通过 HTTP JSON 通信，可独立开发、测试和部署。无状态服务原则：后端不存储用户会话状态，所有鉴权信息通过 JWT Token 在客户端持有，便于系统水平扩展。分层架构原则：严格遵循 Controller → Service → Repository 三层架构，Controller 负责请求路由和参数校验，Service 负责业务逻辑和事务管理，Repository 负责数据访问。领域驱动设计轻量实践：按业务领域（user、goods、order、message、admin）划分顶层包结构，每个领域内部包含自己的 entity/dto/repository/service/controller，避免跨领域耦合。约定优于配置原则：充分利用 Spring Boot 自动配置和 MyBatis-Plus 代码生成能力，减少样板代码编写量。最小权限原则：公开API（商品浏览/搜索/分类/注册/登录）与认证API严格分离；管理员API通过@PreAuthorize("hasRole(\'ADMIN\')")进行方法级权限控制。可演进性原则：架构在以下方面预留扩展点：(1)推荐系统模块—通过独立的RecommendationService接口预留(2)邮件柜接入—通过TradeMethod枚举扩展(3)AI定价建议—通过PricingStrategy接口预留。',
'2.2  系统总体架构图':
'系统采用 C4 模型 L2 容器图风格描述。\n\n客户端层：Vue 3 单页应用（Element Plus UI组件库），运行于用户浏览器中。管理员功能的浏览器端逻辑与普通用户共享同一SPA，通过路由守卫和JWT角色信息区分功能可见性。\n\n服务层：Spring Boot 3.2 单体应用，内部分层如下：\n- Controller 层：处理HTTP请求路由和参数绑定（@RestController + @RequestMapping），包含7个Controller类。\n- Security Filter 层：JwtAuthenticationFilter 拦截所有请求，从 Authorization Header 提取并验证 JWT Token，将用户信息注入 SecurityContextHolder。\n- Service 层：业务逻辑核心，包含6个Service类，每个Service类通过@RequiredArgsConstructor注入所需的Mapper和其他Service。关键业务方法标注@Transactional确保事务一致性。\n- Repository 层：MyBatis-Plus BaseMapper 接口，通过@Mapper注解注册。复杂查询（搜索、对话记录）使用@Select注解编写自定义SQL。\n\n数据层：MySQL 8.0（主数据存储，InnoDB引擎）、Redis 7（缓存和Token黑名单）、MinIO（商品图片对象存储）。\n\n全部5个容器（frontend、backend、mysql、redis、minio）通过 Docker Compose 编排在 swapcampus-net 桥接网络中，一个 docker-compose up -d 命令即可启动全部服务。',
'2.3  技术选型与理由':
'后端框架 Spring Boot 3.2：课程推荐技术栈；生态成熟，社区活跃；自动配置（AutoConfiguration）大幅减少样板代码；内嵌Tomcat容器，无需单独部署应用服务器。\nORM框架 MyBatis-Plus 3.5.6：相比Spring Data JPA更加灵活，允许开发人员手写SQL以应对复杂查询场景；LambdaQueryWrapper提供类型安全的查询构建方式；内置分页插件和自动填充（created_at/updated_at）功能。\n鉴权方案 Spring Security 6 + JWT (jjwt 0.12.5)：Spring Security是Java生态最成熟的安全框架；JWT实现无状态鉴权，服务端无需维护Session，易于水平扩展；jjwt 0.12.x 是最新版本，支持现代化API。\n数据库 MySQL 8.0：课程推荐数据库；InnoDB引擎提供完整的事务支持（ACID）和行级锁；FULLTEXT全文索引支持中文分词搜索（ngram parser）；utf8mb4字符集支持emoji和CJK扩展字符。\n缓存 Redis 7：Spring Data Redis提供开箱即用的集成；适用于JWT Token黑名单、热点商品数据缓存等场景；高性能，单线程模型避免并发竞争。\n对象存储 MinIO：完全开源，Apache 2.0协议；兼容AWS S3 API，迁移成本低；支持本地部署，无需云服务账号；提供Web控制台（:9001）便于开发调试。\n即时通讯 WebSocket + STOMP：Spring框架原生支持STOMP over WebSocket；STOMP协议提供消息目标路由（/user/{id}/queue/chat），简化点对点消息推送实现；SockJS提供自动降级兼容。\n前端 Vue 3 + Element Plus：课程推荐前端框架；Composition API提供更好的代码组织和逻辑复用能力；Element Plus是国内最流行的Vue 3 UI组件库，中文文档完善。\n构建工具 Vite 5：Vue 3官方推荐构建工具；开发服务器秒级冷启动；ESBuild预构建依赖，速度远超Webpack。\n容器化 Docker Compose：课程要求提供一键部署方案；docker-compose.yml编排全部5个服务，环境一致性保证"在我机器上能跑"的问题不再存在。',
'3.1  模块划分':
'auth 认证模块：AuthController（POST /api/auth/register, /login），AuthService（注册业务逻辑，登录验证，JWT签发）。依赖：UserMapper、PasswordEncoder、JwtTokenProvider。\nuser 用户模块：UserController（GET /api/users/me, PUT /api/users/me, GET /api/users/{id}），UserService（个人信息查询与更新，信用分管理）。依赖：UserMapper。\ngoods 商品模块：GoodsController（GET /list, GET /detail/{uuid}, POST /publish, PUT /{uuid}, PUT /{uuid}/status, POST /{uuid}/favorite, GET /my-published），GoodsService（商品CRUD、搜索、收藏管理），FileService（MinIO图片上传）。依赖：GoodsMapper、GoodsImageMapper、FavoriteMapper、FileService。\norder 订单模块：OrderController（POST /, GET /{uuid}, PUT /{uuid}/buyer-confirm, PUT /{uuid}/seller-confirm, PUT /{uuid}/cancel, PUT /{uuid}/review, GET /buyer, GET /seller），OrderService（订单创建、确认收货、取消、评价、信用分更新）。依赖：OrderMapper、GoodsMapper、UserMapper。\nchat 聊天模块：ChatController（REST: GET /conversation/{contactId}, GET /contacts；WebSocket: @MessageMapping /chat/send），MessageService（消息发送、对话查询、已读标记）。依赖：MessageMapper、SimpMessagingTemplate。\nadmin 管理模块：AdminController（GET /dashboard, GET /users, PUT /users/{id}/status, PUT /goods/{uuid}/audit, GET /reports, PUT /reports/{id}/handle），所有方法标注 @PreAuthorize("hasRole(\'ADMIN\')")。依赖：UserMapper、GoodsMapper、OrderMapper、ReportMapper。\ncategory 分类模块：CategoryController（GET /），返回7大分类列表。依赖：CategoryMapper。\ncommon 公共模块：GlobalExceptionHandler（全局异常统一处理），CorsConfig（跨域配置），ApiResponse<T>（统一响应包装），PageQuery（分页参数），MybatisPlusConfig（分页插件和自动填充）。',
'3.2  模块间依赖':
'auth 模块被所有其他模块依赖 —— 前端发起的任何认证请求都经过 JwtAuthenticationFilter 解析Token并注入SecurityContext。user 模块被 goods 和 order 模块依赖 —— 商品发布需要获取卖家信息，订单创建需要获取买卖双方信息。goods 模块被 order 模块依赖 —— 创建订单时需要查询商品详情并更新商品状态。chat 模块相对独立，仅依赖 user 模块获取消息参与者的身份信息。admin 模块依赖所有业务模块 —— 仪表盘需要统计 users/goods/orders/reports 的全表计数，用户管理需要 UserService，举报处理需要 ReportMapper。category 模块完全独立，仅提供分类数据查询。各模块之间通过 Spring 依赖注入（@Autowired/@RequiredArgsConstructor）直接调用 Service 层方法，不使用RPC或消息队列。',
'4.1  关键流程顺序图（≥ 2 张）':
'顺序图一：用户登录认证流程\n参与者：用户浏览器 → AuthController → AuthService → UserMapper → JwtTokenProvider → 返回TokenResponse\n流程：(1)用户POST /api/auth/login发送{username,password}(2)AuthController接收请求，调用authService.login()(3)AuthService通过userMapper.findByUsername查询用户(4)校验用户非null且status=1(5)BCryptPasswordEncoder.matches()验证密码(6)jwtTokenProvider.generateToken(userId,username,role)签发JWT(7)返回TokenResponse{accessToken,userId,username,realName,role}。异常分支：用户不存在→400；密码错误→400；账号禁用→400。\n\n顺序图二：商品发布与图片上传流程\n参与者：用户浏览器 → GoodsController → GoodsService → FileService → MinIO → GoodsMapper → GoodsImageMapper → 返回Goods\n流程：(1)用户POST /api/goods/publish发送multipart/form-data(商品JSON+图片文件数组)(2)GoodsController解析@RequestPart提取数据(3)GoodsService.publishGoods()：生成UUID→构建Goods实体→goodsMapper.insert(goods)(4)循环调用fileService.uploadImage(file,uuid)：检查MinIO bucket→生成对象名→minioClient.putObject()上传→返回URL(5)每个上传成功后goodsImageMapper.insert保存图片记录(6)@Transactional确保全流程原子性：任一步失败则全部回滚(7)返回Goods实体。异常分支：图片超9张→前端限制；单张超5MB→400；MinIO连接失败→500。',
'4.2  状态机':
'订单（Order）状态机是系统中最关键的状态流转，涉及买卖双方和管理员三方的交互：\n\n初始状态：买家点击"立即购买"按钮。\n状态 0（待确认）：订单创建后的初始状态。此时买家尚未确认收货，卖家尚未确认交易。\n- 触发事件(买方取消)：买家或卖家点击"取消订单"→ 状态变为 -1（已取消），同时商品状态恢复为 1（在售）。\n- 触发事件(买方确认)：买家点击"确认收货"→ buyer_confirm = 1，检查 seller_confirm，若也为1则转入状态 2。\n- 触发事件(卖方确认)：卖家点击"确认交易"→ seller_confirm = 1，检查 buyer_confirm，若也为1则转入状态 2。\n状态 2（已完成）：买卖双方均已确认。此时系统自动执行：(1)给买家credit_score加2(上限100)；(2)给卖家credit_score加2(上限100)。\n- 触发事件(买家评价)：买家对订单进行评分（1-5星）和文字评价 → buyer_rating和buyer_review字段更新，订单仍保持状态 2。\n状态 -1（已取消）：订单被任一方取消后的终态。关联商品自动恢复为在售状态，买卖双方信用分不变。',
'5  部署视图':
'部署拓扑图（Docker Compose 编排）：\n\n宿主机（开发环境：macOS/Windows/Linux）\n├── swapcampus-frontend 容器\n│   Image: nginx:alpine（前端构建阶段为 node:20-alpine）\n│   Port: 3000:80\n│   职责：托管Vue 3编译后的静态文件；Nginx反向代理/api请求到backend:8080和/ws请求到backend:8080\n├── swapcampus-backend 容器\n│   Image: openjdk:17-jdk-slim（通过Dockerfile构建）\n│   Port: 8080:8080\n│   职责：运行Spring Boot JAR；通过JDBC连接mysql:3306；通过MinIO Client连接minio:9000\n├── swapcampus-mysql 容器\n│   Image: mysql:8.0\n│   Port: 3306:3306\n│   Volume: mysql_data:/var/lib/mysql\n│   职责：数据持久化；启动时自动执行db/schema.sql初始化数据库\n├── swapcampus-redis 容器\n│   Image: redis:7-alpine\n│   Port: 6379:6379\n│   Volume: redis_data:/data\n│   职责：缓存服务\n└── swapcampus-minio 容器\n    Image: minio/minio:latest\n    Port: 9000:9000 (API), 9001:9001 (Web Console)\n    Volume: minio_data:/data\n    职责：商品图片对象存储\n\n所有容器加入 swapcampus-net（bridge驱动）自定义网络，容器间通过服务名互相访问。启动顺序：mysql(healthcheck就绪)→redis→minio→backend→frontend。一键启动命令：make up 或 docker-compose up -d。',
'6  数据视图（参见 D-05 数据库设计）':
'核心实体及关系（高层ER图）：\n\nUser（用户）—— 1:N —— Goods（商品）：一个用户可以发布多个商品。\nUser（用户）—— 1:N —— Order as Buyer（购买订单）：一个用户可以购买多个商品，产生多条购买订单。\nUser（用户）—— 1:N —— Order as Seller（销售订单）：一个用户卖出的商品被购买，产生多条销售订单。\nGoods（商品）—— 1:N —— GoodsImage（商品图片）：一个商品可以上传多张展示图片（最多9张）。\nGoods（商品）—— 1:N —— Order（订单）：一个商品（在被售出前）可能产生一条订单记录。\nUser（用户）—— N:M —— Goods（商品） via Favorite（收藏）：用户可收藏多个商品，一个商品可被多个用户收藏。Favorite表通过UNIQUE(user_id,goods_uuid)保证不重复收藏。\nUser（用户）—— 1:N —— Message（消息）：用户作为发送方或接收方参与多条消息。\nUser（用户）—— 1:1 —— Wallet（钱包）：每个用户有且仅有一个钱包记录，存储余额和积分。\n\n系统共11张数据表：user, category, goods, goods_image, orders, favorite, message, wallet, report, notification。完整DDL和详细说明见《数据库设计说明书》（产出 D-05）。',
'7.1  ADR-001  示例：选择微服务还是单体？':
'状态：已采纳。\n上下文：SwapCampus系统包含6个功能模块（用户/商品/订单/聊天/管理/分类），团队规模3-4人，编码实现周期仅4天（D6-D9）。\n决策：采用 Spring Boot 单体应用架构，按业务领域进行包级别分层（com.swapcampus.{user,goods,order,message,admin}）。\n后果：正面——开发和调试效率高，IDE可一次性运行全部模块；部署简单，一个JAR文件+一个docker-compose.yml即可；团队成员学习成本低，无需掌握服务发现、负载均衡、分布式事务等微服务基础设施。负面——未来如果系统规模增长到数千用户，单体架构的性能瓶颈和耦合问题将显现，届时需要重构为微服务。\n替代方案：微服务架构（每个模块独立部署，通过REST/RPC通信）。不采纳理由——对于4天开发周期和3-4人团队来说，微服务的运维复杂度远超其收益；课程要求并未涉及大规模并发场景。',
'7.2  ADR-002  示例：选择 PostgreSQL 还是 MySQL？':
'状态：已采纳。\n上下文：系统需要关系型数据库提供持久化存储、事务支持（订单操作需要ACID保证）、外键约束（维护数据引用完整性）、全文搜索（商品关键词检索）。\n决策：选用 MySQL 8.0，全部表使用 InnoDB 存储引擎，字符集采用 utf8mb4。\n后果：正面——课程推荐技术栈，指导教师和组员都熟悉；InnoDB提供完整事务和行锁支持；FULLTEXT + ngram parser 原生支持中文分词搜索，无需额外部署 Elasticsearch；MySQL 8.0 的窗口函数和 CTE 满足复杂统计查询需求。负面——相比 PostgreSQL 缺少部分高级特性（如数组类型、JSONB索引、PostGIS地理查询），但对于校园二手交易场景这些特性并非必需。\n替代方案：(1)PostgreSQL 16——功能更丰富但课程推荐MySQL，且团队对其熟悉度较低；(2)MongoDB——文档模型灵活但无法保证ACID事务，订单场景下的事务一致性是关键需求。',
'7.3  ADR-003  示例：鉴权方案 JWT vs Session？':
'状态：已采纳。\n上下文：系统为前后端分离架构（Vue 3 + Spring Boot），需要一种鉴权方案来保护认证API和管理员API。\n决策：采用基于 JWT（JSON Web Token）的无状态鉴权方案。使用 jjwt 0.12.5 库，HS256 算法签名，Token 有效期设为 24 小时（86400000ms），密钥通过 JWT_SECRET 环境变量注入。\n后果：正面——服务端无需存储Session，天然支持水平扩展；前端将Token存储在localStorage，每次请求通过Authorization: Bearer <token>头发送；JWT payload中可携带userId、username、role等基本信息，减少数据库查询。负面——Token一旦签发就无法主动失效（除非引入Redis黑名单机制）；Token在前端localStorage存储存在XSS攻击风险（可通过HttpOnly Cookie缓解，但会增加架构复杂度）。\n替代方案：(1)基于Session的传统鉴权——需要服务端存储Session，违背无状态原则，不利于扩展；(2)OAuth 2.0——适合接入第三方登录的场景，但对于校内系统过于复杂，课程阶段不需要。',
'8.1  对外接口（OpenAPI 摘要）':
'以下列出 SwapCampus 系统暴露的全部 26 个 REST API 端点，按功能模块分组：\n\n认证模块（公开）：POST /api/auth/register——用户注册；POST /api/auth/login——用户登录。\n\n商品模块（公开+认证混合）：GET /api/goods/list——商品列表（支持?keyword=&categoryId=&sortBy=price|created_at&sortOrder=asc|desc&page=&size=）；GET /api/goods/detail/{uuid}——商品详情（含图片列表和当前用户收藏状态）；GET /api/category——商品分类列表。以下需登录：POST /api/goods/publish——发布商品（multipart/form-data）；PUT /api/goods/{uuid}——编辑商品；PUT /api/goods/{uuid}/status——修改商品状态（下架/删除）；POST /api/goods/{uuid}/favorite——收藏/取消收藏；GET /api/goods/my-published——我的发布列表。\n\n订单模块（全部需登录）：POST /api/orders——创建订单（?goodsUuid=&tradeMethod=&meetLocation=&meetTime=）；GET /api/orders/{uuid}——订单详情；PUT /api/orders/{uuid}/buyer-confirm——买家确认收货；PUT /api/orders/{uuid}/seller-confirm——卖家确认交易；PUT /api/orders/{uuid}/cancel——取消订单；PUT /api/orders/{uuid}/review——评价订单；GET /api/orders/buyer——我的购买列表；GET /api/orders/seller——我的销售列表。\n\n聊天模块（全部需登录）：GET /api/chat/conversation/{contactId}——与某联系人的对话记录；GET /api/chat/contacts——我的联系人列表。WebSocket端点：/ws/chat（STOMP over SockJS）。\n\n用户模块（需登录）：GET /api/users/me——当前用户信息；PUT /api/users/me——更新个人信息；GET /api/users/{id}——查看其他用户公开信息。\n\n管理模块（全部需登录+管理员角色）：GET /api/admin/dashboard——仪表盘统计数据；GET /api/admin/users——用户列表；PUT /api/admin/users/{id}/status——启用/禁用用户；PUT /api/admin/goods/{uuid}/audit——审核商品；GET /api/admin/reports——举报列表；PUT /api/admin/reports/{id}/handle——处理举报。\n\n统一响应格式：{"code": 200, "message": "success", "data": {...}}。错误码体系：200成功、400请求参数错误、401未登录/Token过期、403无权限（非管理员访问管理接口）、404资源不存在、422参数校验失败（MethodArgumentNotValidException）、500服务器内部错误。',
'8.2  内部接口（模块间）':
'WebSocket 内部接口：端点 /ws/chat，协议为 STOMP over SockJS。\n发送消息：客户端 SEND 到目标 /app/chat/send，消息体 {"receiverId": Long, "content": String, "msgType": "TEXT"|"IMAGE", "goodsUuid": String?}。\n接收消息：客户端 SUBSCRIBE 到 /user/{userId}/queue/chat，服务端通过 SimpMessagingTemplate.convertAndSendToUser() 推送。\n\n模块间 Java 调用：Controller → Service 通过 @RequiredArgsConstructor 依赖注入直接调用；Service → Mapper 同样通过构造器注入；跨模块 Service 调用（如 OrderService 调用 UserMapper 更新信用分）通过 Spring Bean 直接调用，不使用 RPC 或消息队列。\n\n共享数据契约：ApiResponse<T>（com.swapcampus.dto）是全部 Controller 方法的统一返回类型；PageQuery（com.swapcampus.dto）是全部列表查询接口的统一分页参数类型；JWT Token 的 claims 结构（userId, username, role）是 JwtTokenProvider 和 JwtAuthenticationFilter 之间的隐式契约。',
'9  错误处理与日志策略':
'错误处理策略：业务异常统一抛出 RuntimeException（或其子类），由 GlobalExceptionHandler（@RestControllerAdvice）集中捕获处理。RuntimeException → HTTP 400 + 中文错误消息（如"该学号已注册"）；BadCredentialsException → HTTP 401；AccessDeniedException → HTTP 403；MethodArgumentNotValidException → HTTP 422 + 字段级错误信息拼接；Exception（兜底）→ HTTP 500 + "服务器内部错误"（记录完整堆栈到日志但不暴露给前端）。\n\n日志策略：Controller 层使用 INFO 级别（记录请求URL、响应状态码）；Service 层使用 DEBUG 级别（记录关键业务步骤和参数）；Security 层使用 INFO 级别（记录认证成功/失败事件，但不记录Token原文）。\n\n敏感信息保护：密码字段（user.password）在任何日志输出中均为 null（因API返回前已 set null）；JWT Token 原文不出现在日志中；MinIO AccessKey/SecretKey 通过环境变量注入，不在配置文件中明文存储。\n\n关键告警阈值（预留）：连续登录失败 ≥ 5次/分钟；订单取消率 ≥ 30%；数据库连接池耗尽。',
'10  安全设计':
'传输层安全：生产环境使用 HTTPS（TLS 1.2+）加密全部HTTP通信；本地开发环境使用 HTTP（localhost）。鉴权与认证：JWT (jjwt 0.12.5) + HS256算法签名；Token有效期24小时，过期后前端自动清除localStorage并跳转登录页；每个请求经JwtAuthenticationFilter解析Token并注入SecurityContext。\n密码安全：用户密码使用 BCryptPasswordEncoder（strength=10）加密存储；登录验证时使用 passwordEncoder.matches(raw, encoded) 进行常量时间比较。SQL注入防护：全部数据库查询使用 MyBatis-Plus 的 #{} 参数化查询语法，杜绝字符串拼接SQL。跨域安全：通过 CorsConfig 配置 allowedOriginPatterns 白名单；开发阶段允许所有来源，生产环境限制为前端域名。权限控制：管理员API使用 @PreAuthorize("hasRole(\'ADMIN\')") 方法级注解；SecurityConfig 区分公开URL（/api/auth/**, /api/goods/list/**, /api/category/**）和需认证URL。敏感数据保护：user.password 在序列化返回前通过 setPassword(null) 清除；数据库密码、JWT Secret、MinIO密钥均通过环境变量注入（${VAR_NAME}），不在代码或配置文件中硬编码。',
}

# ─── D-04 ───
D04 = {
'1.1  编写目的':
'本文档描述 SwapCampus 系统各模块内部的类结构、关键方法的处理逻辑与算法、模块间调用关系和数据流，作为编码实现阶段的直接依据。要求任何一位具有中等 Java/Vue 开发经验的团队成员都能根据本文档独立完成对应模块的编码工作，无需依赖口头沟通或阅读他人代码。',
'1.2  本文档与概要设计的衔接':
'§3.1 商品模块详细设计 → D-03 §3.1 goods 模块、D-03 §4.1 发布商品顺序图\n§3.2 订单模块详细设计 → D-03 §3.1 order 模块、D-03 §4.2 订单状态机\n§3.3 聊天模块详细设计 → D-03 §3.1 chat 模块、D-03 §8.2 WebSocket 内部接口\n§3.4 认证模块详细设计 → D-03 §8.1 认证模块 API、D-03 §10 安全设计\n§4 关键算法描述 → D-03 §4 运行视图\n§5 类与时序图 → D-03 §3 系统结构、D-03 §4 运行视图\n§6 异常处理与并发设计 → D-03 §9 错误处理策略\n§7 单元测试设计要点 → SRS §5 非功能性需求（可维护性）',
'2.1  代码规范':
'Java 后端代码规范（基于阿里巴巴Java开发手册）：缩进使用4个空格（禁止Tab）；类名采用 PascalCase（如 GoodsController、OrderService）；方法名和变量名采用 camelCase（如 findByUuid、goodsMapper）；常量采用 UPPER_SNAKE_CASE（如 MAX_IMAGE_COUNT）；包名全部小写，按业务领域分层（com.swapcampus.controller/.service/.entity/.dto/.repository/.config/.security）；每个公开方法须有 Javadoc 注释说明参数、返回值和异常。\n\nVue 3 前端代码规范：缩进使用2个空格；组件文件名采用 kebab-case（如 main-layout.vue、goods-detail.vue），组件内 name 属性采用 PascalCase；变量和方法采用 camelCase；CSS 类名采用 kebab-case。\n\n通用规范：禁止提交包含调试输出的代码（console.log、System.out.println）；禁止在代码中硬编码密钥或密码（须通过环境变量注入）；Git 提交信息遵循 Conventional Commits 格式（feat:/fix:/docs:/refactor:/chore:/test:）。',
'2.2  公共数据结构':
'ApiResponse<T>（com.swapcampus.dto.ApiResponse）：系统全局统一响应包装类。字段：Integer code（业务状态码），String message（提示信息），T data（业务数据，泛型）。工厂方法：success(T data) → code=200+"success"，success(String msg, T data) → code=200+自定义消息，error(int code, String msg) → 自定义错误码和消息。使用场景：所有 Controller 方法的返回类型统一使用 ApiResponse<T>。\n\nPageQuery（com.swapcampus.dto.PageQuery）：全局统一分页查询参数类。字段：Integer page=1（当前页码），Integer size=10（每页条数），String keyword（搜索关键词），Long categoryId（分类筛选），String sortBy="created_at"（排序字段），String sortOrder="desc"（升序/降序）。使用场景：所有列表查询接口（商品列表、订单列表、用户列表、举报列表）统一使用此参数类。\n\nTokenResponse（com.swapcampus.dto.TokenResponse）：JWT 登录/注册成功后返回的响应对象。字段：String accessToken（JWT Token字符串），String tokenType="Bearer"，Long expiresIn=86400000（24小时毫秒数），Long userId，String username，String realName，String avatarUrl，Integer role（0普通/1管理员）。使用场景：AuthService.login()和AuthService.register()的返回值。',
'2.3  公共工具类':
'IdUtil（cn.hutool.core.util.IdUtil）：Hutool 工具包内置类。使用 IdUtil.fastSimpleUUID() 方法生成不带连字符的32位UUID字符串，用于商品UUID、订单UUID、消息UUID的唯一标识生成。\n\nBCryptPasswordEncoder（org.springframework.security.crypto.bcrypt）：Spring Security 内置密码编码器。encode(rawPassword)用于注册时加密密码；matches(rawPassword, encodedPassword)用于登录时验证密码，采用常量时间比较防止时序攻击。\n\nJwtTokenProvider（com.swapcampus.security.JwtTokenProvider）：自定义JWT工具类，封装 jjwt 0.12.5 库。核心方法：generateToken(userId, username, role) 签发Token；validateToken(token) 验证有效性；getUserIdFromToken(token) / getUsernameFromToken(token) / getRoleFromToken(token) 解析Token。\n\nGlobalExceptionHandler（com.swapcampus.exception.GlobalExceptionHandler）：全局统一异常处理切面（@RestControllerAdvice）。将 RuntimeException → 400、BadCredentialsException → 401、AccessDeniedException → 403、MethodArgumentNotValidException → 422、Exception → 500 统一转换为 ApiResponse 错误格式。',
'3.1  模块 A：__________':
'模块 A：商品模块（goods）\n\n包结构：com.swapcampus.entity.{Goods, GoodsImage, Favorite}（实体）、com.swapcampus.dto.{GoodsRequest, PageQuery}（DTO）、com.swapcampus.repository.{GoodsMapper, GoodsImageMapper, FavoriteMapper}（数据访问）、com.swapcampus.service.{GoodsService, FileService}（业务逻辑）、com.swapcampus.controller.GoodsController（REST控制器）。\n\n核心类 GoodsService：\n- publishGoods(GoodsRequest request, Long sellerId, MultipartFile[] images): @Transactional。流程：(1)IdUtil.fastSimpleUUID()生成UUID (2)构建Goods实体(discountPrice=price*discount/100) (3)goodsMapper.insert(goods) (4)循环images调用fileService.uploadImage上传MinIO (5)每次上传后goodsImageMapper.insert保存URL (6)返回Goods。\n- searchGoods(PageQuery query): 调用goodsMapper.searchGoods()，通过MyBatis动态SQL+MySQL FULLTEXT索引实现关键词搜索和分类筛选，支持按price或created_at排序，MyBatis-Plus Page对象分页。\n- getByUuid(String uuid): 查询商品详情，同时更新viewCount+1。\n- updateGoods(String uuid, GoodsRequest, Long sellerId, MultipartFile[] images): 校验卖家属权后更新商品字段，若有新图片则先删除旧图再上传。\n- toggleFavorite(Long userId, String goodsUuid): 检查是否已收藏→有则删除(favoriteCount-1)，无则新增(favoriteCount+1)，通过favorite表的UNIQUE(user_id,goods_uuid)约束防重复。\n\n核心类 FileService：\n- uploadImage(MultipartFile file, String goodsUuid): 检查MinIO bucket是否存在→不存在则创建→提取文件扩展名→生成对象名"goods/{uuid}/{newUuid}.{ext}"→minioClient.putObject()上传→返回完整URL。\n\n核心类 GoodsMapper（MyBatis-Plus BaseMapper扩展）：\n- findByUuid(@Param("uuid")): @Select查询单条\n- searchGoods(Page, keyword, categoryId, sortBy, sortOrder): 动态SQL（MyBatis <script><if>），MATCH...AGAINST全文检索\n- findBySellerId(Page, sellerId): 卖家商品分页列表',
'3.2  模块 B：__________':
'模块 B：订单模块（order）\n\n包结构：com.swapcampus.entity.Order（实体）、com.swapcampus.repository.OrderMapper（数据访问）、com.swapcampus.service.OrderService（业务逻辑）、com.swapcampus.controller.OrderController（REST控制器）。\n\n核心类 OrderService：\n- createOrder(Long buyerId, String goodsUuid, String tradeMethod, String meetLocation, String meetTime): @Transactional。流程：(1)goodsMapper.findByUuid查商品 (2)三重校验：goods!=null、goods.status==1(在售)、!goods.sellerId.equals(buyerId)(不能自买) (3)IdUtil.fastSimpleUUID()生成订单UUID (4)构建Order(amount=goods.price,status=0待确认) (5)orderMapper.insert(order) (6)goods.status=2(已售出),goodsMapper.updateById(goods) (7)返回Order。\n- buyerConfirm(String orderUuid, Long userId) / sellerConfirm(String orderUuid, Long userId): @Transactional。确认方法：(1)orderMapper.findByUuid查订单 (2)校验操作者是buyer/seller (3)设置buyerConfirm/sellerConfirm=1 (4)检查对方是否也已确认→若是则order.status=2(已完成) (5)addCreditScore给双方各+2分 (6)返回Order。\n- cancelOrder(String orderUuid, Long userId): @Transactional。流程：(1)校验订单存在且操作者是buyer或seller (2)order.status=-1 (3)恢复商品goods.status=1。\n- reviewOrder(String orderUuid, Long buyerId, Integer rating, String review): 仅买家可评价已完成订单（status=2），设置buyerRating和buyerReview。\n\n信用分更新算法 CreditScoreUpdate(userId, delta)：userMapper.selectById查用户→newScore=Math.min(100,Math.max(0,creditScore+delta))→userMapper.updateById。O(1)复杂度，delta范围[-100,100]。',
'3.3  模块 C：__________':
'模块 C：聊天模块（chat）\n\n双通道设计：WebSocket（实时推送）+ REST API（历史记录查询）。\n\nWebSocket 配置（WebSocketConfig）：注册STOMP端点/ws/chat，启用SockJS降级；配置SimpleBroker（/topic, /queue, /user前缀）；设置用户目标前缀/user实现点对点消息。\n\n消息发送流程（ChatController.sendMessage @MessageMapping("/chat/send")）：(1)从Principal获取senderId (2)从@Payload提取receiverId/content/msgType/goodsUuid (3)调用messageService.sendMessage()保存到数据库 (4)messagingTemplate.convertAndSendToUser(String.valueOf(receiverId),"/queue/chat", message)推送。\n\n核心类 MessageService：\n- sendMessage(senderId, receiverId, content, msgType, goodsUuid): 生成UUID→构建Message(isRead=0)→messageMapper.insert→返回Message。\n- getConversation(userId1, userId2, page, size): 分页查询两人之间的双向消息记录，ORDER BY created_at ASC。\n- markAsRead(senderId, receiverId): 批量更新对方发来的未读消息为已读。\n\n模块 D：认证模块（auth）\n\nJWT认证过滤器链：HTTP请求 → JwtAuthenticationFilter(OncePerRequestFilter) → extractToken(从Authorization: Bearer xxx提取) → jwtTokenProvider.validateToken(验证签名+过期时间) → 提取claims(userId,username,role) → 构建UsernamePasswordAuthenticationToken → SecurityContextHolder.setAuthentication。\n\nJwtTokenProvider.generateToken()：使用jjwt 0.12.x builder API，claims包含userId/username/role，subject为username，有效期24h，签名算法HS256。secret通过@Value("${jwt.secret}")从环境变量注入。\n\nAuthService.register()：校验学号唯一性→校验用户名唯一性→BCrypt加密密码→创建User+Wallet→生成JWT→返回TokenResponse。\nAuthService.login()：userMapper.findByUsername查用户→校验status!=0→BCrypt.matches验证密码→生成JWT→返回TokenResponse。',
'4.1  算法 1（如：派评算法 / 推荐算法 / 调度算法）':
'算法名称：信用分增量计算 CreditScoreUpdate(userId, delta)\n输入参数：userId — 用户主键ID（Long），delta — 分数增量（Integer，正数为加分，负数为扣分）\n输出：更新后的信用分值（Integer，范围[0, 100]）\n\n伪代码：\nfunction CreditScoreUpdate(userId, delta):\n    user = userMapper.selectById(userId)\n    if user == null: throw RuntimeException("用户不存在")\n    newScore = user.creditScore + delta\n    if newScore < 0:   newScore = 0\n    if newScore > 100: newScore = 100\n    user.creditScore = newScore\n    userMapper.updateById(user)\n    return newScore\n\n复杂度分析：时间复杂度 O(1) — 单条 SELECT + 单条 UPDATE；空间复杂度 O(1)。\n边界条件：(1)delta为正且newScore>100时夹紧到100 (2)delta为负且newScore<0时夹紧到0 (3)userId不存在→抛异常。\n使用场景：交易完成双方+2（OrderService.buyerConfirm/sellerConfirm）；管理员手动调整（AdminController.toggleUserStatus未来扩展）；违规举报成立-10（AdminController.handleReport未来扩展）。',
'4.2  算法 2':
'算法名称：商品全文搜索 SearchGoods(keyword, categoryId, sortBy, sortOrder, page, size)\n输入参数：keyword（搜索关键词字符串），categoryId（分类筛选，可为null），sortBy（排序字段，"price"或"created_at"），sortOrder（"asc"或"desc"），page（页码），size（每页条数）\n输出：MyBatis-Plus Page<Goods> 分页结果，包含 records 列表、total 总数、current 当前页\n\n伪代码：\nfunction SearchGoods(keyword, categoryId, sortBy, sortOrder, page, size):\n    pageObj = new Page<Goods>(page, size)\n    if keyword != null and keyword != "":\n        // 使用MySQL FULLTEXT布尔模式搜索\n        sqlCondition = "MATCH(title, description) AGAINST(#{keyword} IN BOOLEAN MODE)"\n    if categoryId != null:\n        sqlCondition += " AND category_id = #{categoryId}"\n    sqlCondition += " AND status = 1"  // 仅搜索在售商品\n    sqlOrder = ORDER BY {sortBy} {sortOrder}\n    return goodsMapper.searchGoods(pageObj, keyword, categoryId, sortBy, sortOrder)\n\n复杂度分析：理论上 MySQL FULLTEXT 索引查询时间复杂度为 O(log N)（N为商品总数），实际在课程数据规模（≤500条商品）下所有查询均 < 10ms。\n边界条件：(1)keyword为空时退化为分类筛选+分页查询，使用category_id索引 (2)分类为null时不添加分类条件 (3)无匹配结果时返回空列表，total=0。',
'5.1  核心类图':
'Entity 层核心类：User（id, studentId, username, password, realName, avatarUrl, phone, email, role, creditScore, status）关联Wallet（一对一）、Goods（一对多，作为卖家）、Order（一对多，作为买家或卖家）。Goods（id, uuid, sellerId→User, categoryId→Category, title, description, price, originalPrice, conditionLevel, isBargain, tradeMethod, campusLocation, viewCount, favoriteCount, status）关联GoodsImage（一对多）、Order（一对一）。Order（id, uuid, goodsUuid→Goods, buyerId→User, sellerId→User, amount, tradeMethod, meetLocation, meetTime, buyerConfirm, sellerConfirm, status, buyerRating, buyerReview）关联评价信息。Message（id, uuid, senderId→User, receiverId→User, goodsUuid, content, msgType, isRead）关联收发双方。\n\nService 层依赖关系：GoodsService → GoodsMapper、GoodsImageMapper、FavoriteMapper、FileService。OrderService → OrderMapper、GoodsMapper、UserMapper。MessageService → MessageMapper。AuthService → UserMapper、PasswordEncoder、JwtTokenProvider。AdminController → UserMapper、GoodsMapper、OrderMapper、ReportMapper。\n\nController→Service 依赖：全部通过 @RequiredArgsConstructor（Lombok）构造器注入，Spring 自动管理 Bean 生命周期。',
'5.2  关键时序图':
'时序图一：下单到交易完成全流程\n参与者：买家浏览器、卖家浏览器、GoodsController、OrderController、OrderService、GoodsMapper、UserMapper。\n流程：(1)买家GET /api/goods/detail/{uuid}浏览商品 (2)买家POST /api/orders下单 → OrderController.create→OrderService.createOrder：校验商品→生成UUID→orderMapper.insert→goods.status=2→返回Order (3)买卖双方通过WebSocket Chat沟通面交细节 (4)面交完成后买家PUT /api/orders/{uuid}/buyer-confirm→OrderService.buyerConfirm：confirm=1，检查双方→信用分+2 (5)卖家PUT /api/orders/{uuid}/seller-confirm→同理 (6)双方确认后订单自动完成。\n\n时序图二：用户注册登录全流程\n参与者：浏览器、AuthController、AuthService、UserMapper、PasswordEncoder、JwtTokenProvider。\n流程：(1)注册POST /api/auth/register→AuthService.register：校验学号/用户名唯一性→BCrypt加密→创建User+Wallet→JWT签发→返回TokenResponse (2)登录POST /api/auth/login→AuthService.login：findByUsername→BCrypt.matches验证→JWT签发→返回TokenResponse (3)前端将Token存入localStorage，后续请求自动附加Authorization Header。',
'6  异常处理与并发设计':
'异常处理：全部业务异常通过 GlobalExceptionHandler（@RestControllerAdvice）统一拦截。RuntimeException → 400 + message；MethodArgumentNotValidException → 422 + 字段级错误拼接；BadCredentialsException → 401；AccessDeniedException → 403；Exception（兜底）→ 500 + 日志记录。关键业务方法使用 @Transactional(rollbackFor = Exception.class) 确保数据库回滚。\n\n并发控制：(1)订单创建：下单前校验 goods.status==1（在售），创建订单后立即 goods.status=2（已售出）。由于 UPDATE goods SET status=2 WHERE id=? AND status=1 是原子操作，InnoDB行锁自动防止并发重复下单。(2)双方确认：buyerConfirm 和 sellerConfirm 各自独立 UPDATE，不依赖读取时的状态值，因此无需乐观锁版本号。(3)消息发送：每条消息独立 INSERT，UUID唯一约束保证不重复。(4)收藏操作：favorite 表 UNIQUE(user_id, goods_uuid) 约束防止重复收藏，toggleFavorite 中的 INSERT/DELETE 在数据库层面保证正确性。\n\n幂等性设计：所有创建操作（注册、发布商品、下单、发送消息）均使用服务端生成的UUID作为业务主键，数据库 UNIQUE 约束天然保证幂等——重复请求不会产生重复数据。',
'7  单元测试设计要点':
'各模块测试策略、关键用例和Mock边界：\n\nAuthService 单元测试：正常注册（Mock UserMapper返回null表示学号未注册→验证UserMapper.insert被调用且password已加密）；学号重复（Mock findByStudentId返回已有User→期望抛RuntimeException）；用户名重复（Mock findByUsername返回已有User→期望抛RuntimeException）；登录成功（Mock findByUsername返回User→Mock passwordEncoder.matches返回true→验证返回TokenResponse非null）；密码错误（Mock matches返回false→期望抛RuntimeException）。\n\nGoodsService 单元测试：正常发布（Mock goodsMapper.insert+goodsImageMapper.insert+fileService.uploadImage→验证返回Goods非null且UUID非空）；搜索有结果（Mock goodsMapper.searchGoods返回含3条记录的Page→验证records.size()==3）；搜索无结果（Mock返回空Page→验证records为空）；下架（Mock findByUuid返回Goods(status=1)→调用changeStatus(status=0)→验证goodsMapper.updateById被调用）。\n\nOrderService 单元测试：正常下单（Mock goodsMapper.findByUuid返回在售商品→Mock orderMapper.insert→验证商品status变为2）；购买自己商品（sellerId==buyerId→期望抛异常）；双方确认（Mock findByUuid返回Order→连续调用buyerConfirm+sellerConfirm→验证status变为2且addCreditScore各被调用一次）；取消订单（验证status=-1且商品恢复status=1）。\n\nJwtTokenProvider 纯逻辑测试（无需Mock）：生成Token→验证非空；解析Token→验证userId/username/role正确；过期Token→验证validateToken返回false；篡改Token→验证validateToken返回false。\n\nGlobalExceptionHandler 单元测试：使用 MockMvc 模拟Controller抛出各类异常→验证HTTP状态码和响应JSON中的code字段正确。\n\n覆盖率目标：整体行覆盖率 ≥ 60%（JaCoCo统计）；核心模块（AuthService、GoodsService、OrderService）行覆盖率 ≥ 80%；分支覆盖率 ≥ 50%。',
}

# ─── D-05 ───
D05 = {
'1.1  编写目的':
'本文档定义 SwapCampus 校园闲置物品交易平台的数据库设计，涵盖概念模型（ER图）、逻辑模型（关系模式与范式分析）、物理模型（完整DDL脚本与数据字典）、索引策略、安全与备份策略、数据生命周期管理。预期读者为后端开发人员（指导数据库相关编码）和数据库管理员（指导运维操作）。',
'1.2  数据库选型与版本':
'DBMS 选型：MySQL 8.0.21+。存储引擎：InnoDB（全部11张表统一使用）。字符集与排序规则：utf8mb4 + utf8mb4_unicode_ci。服务器时区：Asia/Shanghai（UTC+8）。\n\n选型理由（详见 ADR-002）：(1)课程任务书和选题指南明确推荐 MySQL 8 (2)InnoDB引擎完整支持ACID事务、行级锁和外键约束——订单场景的事务一致性是核心需求 (3)MySQL 8.0 内置 FULLTEXT 全文索引配合 ngram parser 支持中文分词，无需额外部署 Elasticsearch (4)团队在前期课程中已积累 MySQL 使用经验。',
'2  概念模型（ER 图）':
'全局 ER 图描述系统11张表的实体及关系：\n\n核心实体：(1)User（用户）——属性：id(PK), student_id(AK), username(AK), password, real_name, avatar_url, phone, email, role, credit_score, status, created_at, updated_at (2)Goods（商品）——属性：id(PK), uuid(AK), seller_id(FK→User), category_id(FK→Category), title, description, price, original_price, condition_level, is_bargain, trade_method, campus_location, view_count, favorite_count, status, created_at, updated_at (3)Order（订单）——属性：id(PK), uuid(AK), goods_uuid(FK→Goods), buyer_id(FK→User), seller_id(FK→User), amount, trade_method, meet_location, meet_time, buyer_confirm, seller_confirm, status, buyer_rating, buyer_review, completed_at, created_at, updated_at (4)Message（消息）——属性：id(PK), uuid(AK), sender_id(FK→User), receiver_id(FK→User), goods_uuid, content, msg_type, is_read, created_at (5)其他实体：Category, GoodsImage, Favorite, Wallet, Report, Notification。\n\n主要关系：(1)User 1:N Goods ——一个用户可以发布多个商品（通过seller_id外键）(2)User 1:N Order(as buyer/seller)——用户作为买家(buyer_id)或卖家(seller_id)参与多条订单 (3)Goods 1:N GoodsImage——一个商品拥有多张展示图片 (4)User N:M Goods via Favorite——用户收藏商品的多对多关系，通过Favorite中间表实现，含UNIQUE(user_id,goods_uuid)约束 (5)User 1:N Message——用户作为发送方(sender_id)或接收方(receiver_id)收发消息 (6)User 1:1 Wallet——每个用户有唯一钱包记录（wallet.user_id UNIQUE约束）。\n\n交易域子图：User（买家）→ Order ← Goods，Order同时关联买卖双方和商品。\n社交域子图：User ↔ Message ↔ User（双向聊天），User → Favorite → Goods（收藏关系），Goods → GoodsImage（图片）。',
'3.1  关系模式':
'以下使用标准关系模式表示法 R(属性列表, 主键, 外键) 描述全部11张表：\n\nuser(id, student_id, username, password, real_name, avatar_url, phone, email, role, credit_score, status, created_at, updated_at)。PK: id。AK: student_id, username。\n\ncategory(id, name, icon, sort_order, status, created_at)。PK: id。\n\ngoods(id, uuid, seller_id, category_id, title, description, price, original_price, condition_level, is_bargain, trade_method, campus_location, view_count, favorite_count, status, created_at, updated_at)。PK: id。AK: uuid。FK: seller_id→user(id), category_id→category(id)。\n\ngoods_image(id, goods_uuid, url, sort_order, created_at)。PK: id。FK: goods_uuid→goods(uuid)。\n\norders(id, uuid, goods_uuid, buyer_id, seller_id, amount, trade_method, meet_location, meet_time, buyer_confirm, seller_confirm, status, buyer_rating, buyer_review, completed_at, created_at, updated_at)。PK: id。AK: uuid。FK: goods_uuid→goods(uuid), buyer_id→user(id), seller_id→user(id)。\n\nfavorite(id, user_id, goods_uuid, created_at)。PK: id。FK: user_id→user(id), goods_uuid→goods(uuid)。UK: (user_id, goods_uuid)。\n\nmessage(id, uuid, sender_id, receiver_id, goods_uuid, content, msg_type, is_read, created_at)。PK: id。AK: uuid。FK: sender_id→user(id), receiver_id→user(id)。\n\nwallet(id, user_id, balance, points, created_at, updated_at)。PK: id。FK: user_id→user(id)。UK: user_id。\n\nreport(id, reporter_id, goods_uuid, reported_user_id, reason, description, status, handler_id, handle_remark, created_at, handled_at)。PK: id。FK: reporter_id→user(id)。\n\nnotification(id, user_id, title, content, is_read, created_at)。PK: id。FK: user_id→user(id)。',
'3.2  范式说明与反范式取舍':
'范式验证：全部11张表均满足第三范式（3NF）要求——每个非主属性完全函数依赖于主键，不存在传递函数依赖。以 orders 表为例：所有属性（amount、status、trade_method等）都由主键 id 唯一确定，buyer_id 和 seller_id 通过外键关联 user 表而非在 orders 中冗余存储用户的姓名等信息。\n\n唯一的反范式设计取舍：goods 表中的 favorite_count 字段。该字段的值可以由 SELECT COUNT(*) FROM favorite WHERE goods_uuid=? 实时计算得出。但考虑到商品列表是系统最高频的查询场景（首页加载需展示收藏数），每次列表查询都对 favorite 表做子查询或JOIN 会显著增加查询开销。因此选择在 goods 表中冗余存储 favorite_count，并在 GoodsService.toggleFavorite() 方法中同步维护（收藏+1，取消收藏-1），通过 @Transactional 保证数据一致性。\n\n其他设计说明：category 表独立存储分类信息而非在 goods 表中存储分类名称字符串，既满足3NF（避免数据冗余），也便于未来增加/修改分类。',
'4.1  数据字典':
'数据字典完整记录每张表的每个字段的元数据，以下列出4张核心表的详细信息（全部11张表的数据字典见 db/schema.sql 中的注释）：\n\nuser 表：id — BIGINT — PK/AUTO_INCREMENT — 用户唯一标识。student_id — VARCHAR(20) — NOT NULL/UNIQUE — 学号。username — VARCHAR(64) — NOT NULL/UNIQUE — 登录用户名。password — VARCHAR(255) — NOT NULL — BCrypt加密密文（敏感字段，API返回前setNull）。real_name — VARCHAR(64) — NOT NULL — 真实姓名。avatar_url — VARCHAR(512) — NULL — 头像URL。phone — VARCHAR(20) — NULL — 手机号。email — VARCHAR(128) — NULL — 邮箱。role — TINYINT — NOT NULL/DEFAULT 0 — 角色（0普通用户/1管理员）。credit_score — INT — NOT NULL/DEFAULT 80 — 信用分。status — TINYINT — NOT NULL/DEFAULT 1 — 状态（0禁用/1正常/2待审核）。created_at — DATETIME — NOT NULL/DEFAULT CURRENT_TIMESTAMP。updated_at — DATETIME — NOT NULL/ON UPDATE。\n\ngoods 表：id — BIGINT — PK/AUTO_INCREMENT。uuid — VARCHAR(36) — NOT NULL/UNIQUE — 业务主键（对外暴露）。seller_id — BIGINT — NOT NULL/FK→user(id) — 卖家。category_id — BIGINT — NULL/FK→category(id) — 商品分类。title — VARCHAR(255) — NOT NULL — 商品标题。description — TEXT — NULL — 商品描述。price — DECIMAL(10,2) — NOT NULL — 售价。original_price — DECIMAL(10,2) — NULL — 原价。condition_level — TINYINT — NOT NULL/DEFAULT 5 — 成色(1-10)。is_bargain — TINYINT — NOT NULL/DEFAULT 0 — 是否接受议价。trade_method — VARCHAR(32) — NOT NULL/DEFAULT \'FACE\' — 交易方式(FACE/LOCKER/BOTH)。campus_location — VARCHAR(128) — NULL — 交易校区。view_count — INT — NOT NULL/DEFAULT 0 — 浏览量。favorite_count — INT — NOT NULL/DEFAULT 0 — 收藏数（反范式冗余）。status — TINYINT — NOT NULL/DEFAULT 1 — 状态(-1删除/0下架/1在售/2售出/3审核中)。created_at — DATETIME — NOT NULL。updated_at — DATETIME — NOT NULL。\n\norders 表：id — BIGINT — PK/AUTO_INCREMENT。uuid — VARCHAR(36) — NOT NULL/UNIQUE — 订单号。goods_uuid — VARCHAR(36) — NOT NULL/FK→goods(uuid)。buyer_id — BIGINT — NOT NULL/FK→user(id)。seller_id — BIGINT — NOT NULL/FK→user(id)。amount — DECIMAL(10,2) — NOT NULL — 成交金额。trade_method — VARCHAR(32) — NOT NULL。meet_location — VARCHAR(256) — NULL。meet_time — DATETIME — NULL。buyer_confirm — TINYINT — NOT NULL/DEFAULT 0。seller_confirm — TINYINT — NOT NULL/DEFAULT 0。status — TINYINT — NOT NULL/DEFAULT 0 — 状态(-1取消/0待确认/1进行中/2完成)。buyer_rating — TINYINT — NULL — 评分1-5。buyer_review — TEXT — NULL — 评价内容。completed_at — DATETIME — NULL。created_at — DATETIME — NOT NULL。updated_at — DATETIME — NOT NULL。\n\nmessage 表：id — BIGINT — PK/AUTO_INCREMENT。uuid — VARCHAR(36) — NOT NULL/UNIQUE。sender_id — BIGINT — NOT NULL/FK→user(id)。receiver_id — BIGINT — NOT NULL/FK→user(id)。goods_uuid — VARCHAR(36) — NULL — 关联商品。content — TEXT — NOT NULL — 消息内容。msg_type — VARCHAR(16) — NOT NULL/DEFAULT \'TEXT\' — 消息类型(TEXT/IMAGE/SYSTEM)。is_read — TINYINT — NOT NULL/DEFAULT 0。created_at — DATETIME — NOT NULL。',
'4.2  DDL 脚本':
'完整 DDL 脚本位于项目根目录下的 db/schema.sql 文件中，包含以下内容：\n\n(1) 数据库创建：DROP DATABASE IF EXISTS swapcampus; CREATE DATABASE swapcampus CHARACTER SET utf8mb4 COLLATE utf8mb4_unicode_ci;\n\n(2) 11张表的 CREATE TABLE 语句，按外键依赖顺序排列（先创建被引用表再创建引用表）：user → category → goods → goods_image → orders → favorite → message → wallet → report → notification。每张表含完整的字段定义（类型/约束/默认值/注释）、PRIMARY KEY（AUTO_INCREMENT）、FOREIGN KEY（REFERENCES + 级联规则）、INDEX（普通索引/唯一索引/全文索引/复合索引）。\n\n(3) 种子数据：7条商品分类记录（教材教辅/电子数码/生活用品/运动户外/服饰美妆/图书音像/其他闲置，sort_order1-7）。1个管理员账号（admin/admin123，BCrypt加密，role=1）。3个测试用户（zhangsan/lisi/wangwu，密码123456，role=0，credit_score分别为85/90/75）。3条测试钱包记录（zhangsan余额0积分100，lisi余额100积分200，wangwu余额50积分50）。\n\n(4) Docker部署集成：schema.sql 在 docker-compose.yml 中挂载为 MySQL 容器的 /docker-entrypoint-initdb.d/01-schema.sql，容器首次启动时自动执行建库建表和插入种子数据。',
'5.1  索引清单':
'系统为全部11张表设计了共计20+个索引，覆盖所有高频查询路径：\n\nuser表：(1)PRIMARY(id)聚簇索引 (2)idx_student_id(student_id)普通索引——用于登录时学号查询 (3)idx_username(username)普通索引——用于登录时用户名查询 (4)idx_status(status)普通索引——用于管理员筛选用户。\n\ngoods表：(1)PRIMARY(id)聚簇索引 (2)idx_uuid(uuid)普通索引——最高频查询，商品详情/订单创建均通过UUID查询 (3)idx_seller(seller_id)普通索引——卖家查看自己的商品列表 (4)idx_category(category_id)普通索引——分类筛选 (5)idx_status(status)普通索引——只查询在售商品(status=1) (6)ft_title_desc(title,description)FULLTEXT索引+ngram parser——支持中文关键词全文检索。\n\ngoods_image表：idx_goods_uuid(goods_uuid)普通索引——按商品UUID查询所有图片。\n\norders表：(1)PRIMARY(id)聚簇索引 (2)idx_uuid(uuid)普通索引——通过订单号查询 (3)idx_buyer(buyer_id)普通索引——买家查看自己的购买订单 (4)idx_seller(seller_id)普通索引——卖家查看自己的销售订单 (5)idx_status(status)普通索引——按状态筛选（待确认/已完成/已取消）。\n\nfavorite表：uk_user_goods(user_id,goods_uuid)UNIQUE索引——保证同一用户不重复收藏同一商品，同时支持"我的收藏"列表查询。\n\nmessage表：(1)PRIMARY(id) (2)idx_sender(sender_id) (3)idx_receiver(receiver_id) (4)idx_goods(goods_uuid)关联商品查询 (5)idx_created(created_at)时间排序。对话查询同时使用sender_id和receiver_id。\n\nnotification表：idx_user_read(user_id,is_read)复合索引——查询某用户的未读通知。\n\nwallet表：user_id UNIQUE索引——保证一个用户一个钱包。\n\nreport表：idx_status(status)——管理员按状态筛选举报。',
'5.2  性能预估':
'课程阶段数据规模预估（非生产环境，仅开发和演示阶段）：用户总数 ≤ 10 人（3-4名组员+5-6名灰度测试用户）。商品总数 ≤ 500 条（含200条种子数据+约300条测试发布）。订单总数 ≤ 100 条（课程周期2周内的测试交易）。消息总数 ≤ 2000 条（测试聊天记录）。单表最大行数为 message 表约 2000 行。并发查询数（QPS）≤ 10（单人/少人测试场景）。\n\n在此数据规模下，所有数据库查询均在毫秒级完成，不存在性能瓶颈。MySQL InnoDB Buffer Pool 足以将全部热数据缓存于内存中。\n\n生产环境 Top 5 慢查询预留治理思路：(1)商品全文搜索 → FULLTEXT索引已覆盖，数据量增长后考虑引入Elasticsearch (2)订单列表按买家+状态查询 → idx_buyer+idx_status可能需要复合索引 (3)消息对话记录 → 当前idx_sender+idx_receiver+idx_created可能需调整为(sender_id,receiver_id,created_at)复合索引 (4)管理员用户列表 → 数据量小，idx_status+分页即足够 (5)举报列表 → idx_status+分页。',
'6.1  权限设计':
'数据库账号与权限设计（最小权限原则）：\n\n本地开发环境：使用 MySQL root 账号（拥有 ALL PRIVILEGES），连接仅允许 localhost（127.0.0.1），密码通过环境变量注入，不硬编码。用途：开发阶段的建表、数据修改、调试。\n\n生产环境（预留，课程阶段无需实施）：(1)swapcampus_app 账号——拥有 SELECT, INSERT, UPDATE, DELETE 权限，仅限 swapcampus 数据库内的所有表，用于 Spring Boot 应用的数据访问连接。连接来源限制为应用服务器 IP。(2)swapcampus_readonly 账号——拥有 SELECT 权限，用于数据分析和报表查询。连接来源限制为内网 IP。(3)swapcampus_admin 账号——拥有 ALL PRIVILEGES，仅限 localhost，用于紧急数据库维护操作。\n\n权限授予示例（生产环境）：GRANT SELECT, INSERT, UPDATE, DELETE ON swapcampus.* TO \'swapcampus_app\'@\'app_host\' IDENTIFIED BY \'strong_password\'; GRANT SELECT ON swapcampus.* TO \'swapcampus_readonly\'@\'%\' IDENTIFIED BY \'readonly_password\';',
'6.2  敏感字段处理':
'user.password：存储层——使用 BCrypt（strength=10，Spring Security默认）加密存储，数据库中的值为$2a$10$...格式的密文，即使数据库泄露攻击者也无法还原明文密码。应用层——API返回User对象前调用 user.setPassword(null) 清除密码字段，确保不会意外序列化到前端。日志层——所有日志输出中不包含password字段。\n\nuser.phone 和 user.email：仅在用户本人查看自己的个人中心（GET /api/users/me）和订单详情（买卖双方需联系）时返回完整值。在其他用户查看公开信息（GET /api/users/{id}）时可选择性脱敏（未来扩展）。\n\nJWT Secret 和 MinIO 密钥：不在 application.yml 中硬编码，通过 ${JWT_SECRET} 和 ${MINIO_SECRET_KEY} 环境变量占位符引用，实际值在 docker-compose.yml 或 .env 文件中注入。.env 文件已加入 .gitignore，不纳入版本控制。',
'6.3  审计与备份':
'课程阶段备份策略：备份频率为每日手动执行一次完整数据导出。备份方法：mysqldump -u root -p swapcampus > backup_YYYYMMDD.sql。备份文件保留方式：纳入 Git 版本控制（db/schema.sql 为最新版本），历史备份文件存放在本地磁盘。恢复测试：删除数据库后重新执行 mysql < db/schema.sql 验证可完整重建。\n\n生产环境备份策略（预留，课程阶段无需实施）：备份频率为每日凌晨3:00自动全量备份（通过 cron job 或 GitHub Actions scheduled workflow）。备份保留周期为30天滚动（超过30天的自动删除）。备份存储为异地云存储（如阿里云OSS/AWS S3）+ 本地保留最近7天。恢复演练：每季度执行一次完整的备份恢复演练，验证 RTO（恢复时间目标）≤ 30分钟。\n\n审计日志（预留）：关键操作（管理员禁用用户、管理员处理举报、管理员下架商品）应记录操作人、操作时间、操作类型、操作对象到 audit_log 表，便于追溯。课程阶段数据量小，暂不实现。',
'7  数据生命周期':
'各表数据保留策略：\n\nuser 表：课程期间保留全部账号。课程结束后可删除或匿名化处理（清除phone/email等个人信息，保留统计数据）。\n\ngoods 表：在售(status=1)、已售出(status=2)、审核中(status=3)的商品记录保留。下架(status=0)和删除(status=-1)的商品为软删除，物理数据保留至课程结束。\n\norders 表：已完成(status=2)和已取消(status=-1)的订单保留作为交易历史记录，取消的订单不物理删除。\n\nmessage 表：全部消息记录保留至课程结束，不主动清理。聊天记录可用于演示和测试。\n\nfavorite 表：用户取消收藏时物理删除（DELETE FROM favorite WHERE user_id=? AND goods_uuid=?），不保留取消收藏的历史记录。\n\nreport 表：已处理和已驳回的举报保留至课程结束作为管理记录。\n\nnotification 表：通知标记 is_read=1 后不物理删除，课程结束后可统一清理。\n\n总体原则：课程项目为期2周+评审，数据量小（总量不超过几千行），无需设计复杂的归档或定时清理任务。所有数据在课程周期内保留即可，课程结束后数据库可整体删除或归档。',
}


# ═══════════════ 转换核心 ═══════════════

def fill_doc(tmpl_name, content_map):
    tmpl_path = os.path.join(TPL_DIR, tmpl_name)
    out_path = os.path.join(OUT_DIR, tmpl_name.replace('_模板', ''))

    doc = Document(tmpl_path)

    # 构建段落索引：{章节标题文本 → 段落序号}
    title_index = {}
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if t:
            title_index[t] = i

    # 按段落索引从大到小排序（从文档底部往上处理），避免插入内容后索引偏移
    sorted_sections = []
    for section_title, content in content_map.items():
        # 找到章节标题段落
        if section_title in title_index:
            idx = title_index[section_title]
        else:
            # 模糊匹配
            matched = None
            for t in title_index:
                if section_title in t or t in section_title:
                    matched = t
                    break
            if not matched:
                continue
            idx = title_index[matched]
        sorted_sections.append((idx, section_title, content))

    # 按索引降序排列（从文档底部向上处理）
    sorted_sections.sort(key=lambda x: x[0], reverse=True)

    # 遍历内容映射，从文档底部向上处理，避免索引偏移
    for _, section_title, content in sorted_sections:
        if section_title not in title_index:
            continue
        section_idx = title_index[section_title]

        # 找到标题后的第一个空段落（无runs或runs全为空）
        insert_target = None
        for j in range(section_idx + 1, len(doc.paragraphs)):
            p = doc.paragraphs[j]
            text = p.text.strip()
            runs_text = ''.join(r.text for r in p.runs)
            if not text and not runs_text:
                insert_target = j
                break
            # 如果遇到了下一个有文字的段落（非空），说明没有空段落可填充
            if text:
                break

        if insert_target is None:
            continue

        # 在目标空段落及其后插入内容
        insert_content(doc, insert_target, content)

    doc.save(out_path)
    print(f'✅ {os.path.basename(out_path)}')


def insert_content(doc, para_idx, content):
    """在指定段落位置插入多行内容，第一行替换原段落"""
    from docx.oxml import OxmlElement
    import copy

    target_p = doc.paragraphs[para_idx]
    body = target_p._element.getparent()
    target_elem = target_p._element

    # 获取原段落的格式参考（段落属性pPr和run属性rPr）
    pPr = target_elem.find(f'{{{WML_NS}}}pPr')
    ref_rPr = None
    for r in target_p.runs:
        rPr_elem = r._element.find(f'{{{WML_NS}}}rPr')
        if rPr_elem is not None:
            ref_rPr = copy.deepcopy(rPr_elem)
            break

    lines = [l for l in content.split('\n') if l.strip()]
    if not lines:
        return

    # 第一行：写到原占位段落
    for r in target_p.runs:
        r._element.getparent().remove(r._element)
    new_r = OxmlElement('w:r')
    if ref_rPr is not None:
        new_r.append(copy.deepcopy(ref_rPr))
    new_t = OxmlElement('w:t')
    new_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    new_t.text = lines[0]
    new_r.append(new_t)
    target_elem.append(new_r)

    # 后续行：在target_p之后插入新段落
    body_children = list(body)
    insert_pos = body_children.index(target_elem)
    for line in lines[1:]:
        insert_pos += 1
        new_p = OxmlElement('w:p')
        if pPr is not None:
            new_p.append(copy.deepcopy(pPr))
        new_r2 = OxmlElement('w:r')
        if ref_rPr is not None:
            new_r2.append(copy.deepcopy(ref_rPr))
        new_t2 = OxmlElement('w:t')
        new_t2.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        new_t2.text = line
        new_r2.append(new_t2)
        new_p.append(new_r2)
        body.insert(insert_pos, new_p)


# ═══════════════ 执行 ═══════════════
if __name__ == '__main__':
    print('基于模板生成 D-02~D-05 Word 文档...\n')

    fill_doc('D-02_需求规格说明书_SRS_模板.docx', D02)
    fill_doc('D-03_概要设计说明书_模板.docx', D03)
    fill_doc('D-04_详细设计说明书_模板.docx', D04)
    fill_doc('D-05_数据库设计说明书_模板.docx', D05)

    print('\n全部完成！')
