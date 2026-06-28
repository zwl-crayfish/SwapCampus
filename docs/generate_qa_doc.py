#!/usr/bin/env python3
"""生成 SwapCampus 答辩QA问答 Word 文档"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ============================================================
# 全局样式设置
# ============================================================
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(4)

# 设置页边距
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# 自定义标题样式
def set_heading_style(heading_style, size, color_hex):
    heading_style.font.name = '微软雅黑'
    heading_style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    heading_style.font.size = Pt(size)
    heading_style.font.color.rgb = RGBColor.from_string(color_hex)
    heading_style.font.bold = True
    heading_style.paragraph_format.space_before = Pt(18)
    heading_style.paragraph_format.space_after = Pt(8)

set_heading_style(doc.styles['Heading 1'], 18, '1A5276')
set_heading_style(doc.styles['Heading 2'], 14, '2C3E50')
set_heading_style(doc.styles['Heading 3'], 12, '34495E')

def add_code_block(doc, text):
    """添加代码样式段落"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
    # 灰色背景效果通过缩进实现
    return p

def add_note(doc, text, note_type="info"):
    """添加提示框"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    prefix = {"info": "ℹ️ ", "warn": "⚠️ ", "success": "✅ ", "cross": "❌ "}
    run = p.add_run(prefix.get(note_type, "") + text)
    run.font.size = Pt(9.5)
    run.font.name = '微软雅黑'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if note_type == "warn":
        run.font.color.rgb = RGBColor(0xE6, 0x7E, 0x22)
    elif note_type == "cross":
        run.font.color.rgb = RGBColor(0xE7, 0x4C, 0x3C)
    elif note_type == "success":
        run.font.color.rgb = RGBColor(0x27, 0xAE, 0x60)
    return p

def add_table_with_data(doc, headers, rows, col_widths=None):
    """添加格式化表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.size = Pt(9.5)
    # 数据行
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
    doc.add_paragraph()  # 表后间距
    return table

# ============================================================
# 封面
# ============================================================
for _ in range(4):
    doc.add_paragraph()

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run('SwapCampus 校园闲置物品交易平台')
run.font.size = Pt(28)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1A, 0x52, 0x76)
run.font.name = '微软雅黑'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

subtitle_p = doc.add_paragraph()
subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle_p.add_run('答辩 Q&A 问答汇编')
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
run.font.name = '微软雅黑'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

doc.add_paragraph()

info_p = doc.add_paragraph()
info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
info_lines = [
    '北京林业大学 · 信息学院',
    '《软件工程（课程设计）》2026春 · T-02',
    f'生成日期：{datetime.date.today().strftime("%Y年%m月%d日")}',
]
for line in info_lines:
    run = info_p.add_run(line + '\n')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)
    run.font.name = '微软雅黑'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

doc.add_page_break()

# ============================================================
# 目录页
# ============================================================
doc.add_heading('目  录', level=1)
toc_items = [
    ('一', '技术栈总览'),
    ('二', '核心模块与完整度'),
    ('三', '商品状态流转机制'),
    ('四', '高级筛选与成色字段'),
    ('五', '商品分类的数据来源'),
    ('六', '聊天系统通知'),
    ('七', '信用分机制'),
    ('八', '管理员权限'),
    ('九', '面交地点与时间'),
    ('十', '浏览量机制与 AI 功能'),
    ('十一', '并发与一致性'),
    ('十二', '实时聊天技术实现'),
    ('十三', '安全与权限'),
    ('十四', '审核与举报流程'),
    ('十五', '性能与数据存储'),
    ('十六', '总结：设计与实现的差距'),
]
for num, title in toc_items:
    p = doc.add_paragraph()
    run = p.add_run(f'{num}、{title}')
    run.font.size = Pt(12)
    run.font.name = '微软雅黑'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

doc.add_page_break()

# ============================================================
# 一、技术栈总览
# ============================================================
doc.add_heading('一、技术栈总览', level=1)

doc.add_paragraph(
    'SwapCampus 采用前后端分离的经典分层架构，通过 Docker Compose 实现一键部署。'
    '全栈由 5 个容器（MySQL、Redis、MinIO、Spring Boot 后端、Vue 3 前端）'
    '组成，通过桥接网络互通，使用 .env 文件管理敏感配置。'
)

doc.add_heading('1.1 各层技术选型', level=2)
add_table_with_data(doc,
    ['层级', '技术栈', '版本/说明'],
    [
        ['前端', 'Vue 3 + Element Plus + Axios + Pinia\n+ SockJS + STOMP.js', 'Vue 3.4.21 / Element Plus 2.6.1'],
        ['后端', 'Spring Boot + MyBatis-Plus\n+ Spring Security + WebSocket STOMP', 'Boot 3.2.5 / MP 3.5.6 / JDK 17'],
        ['数据库', 'MySQL 8 (InnoDB) + Redis 7', 'utf8mb4 / Asia/Shanghai'],
        ['鉴权', 'JWT (jjwt 0.12.5) + BCrypt', '无状态 Token / 24h 过期'],
        ['即时通讯', 'WebSocket STOMP (SockJS 兜底)', '消息持久化到 MySQL'],
        ['对象存储', '本地文件系统 (./uploads/)', 'MinIO 容器已配置但代码未对接'],
        ['容器化', 'Docker + Docker Compose', '5 容器 / 桥接网络'],
    ]
)

doc.add_heading('1.2 项目结构', level=2)
add_code_block(doc, '''SwapCampus/
├── backend/          # Spring Boot 3.2 + MyBatis-Plus
│   ├── config/       # Security / WebSocket / MinIO / CORS / MyBatis-Plus 配置
│   ├── controller/   # 8 个控制器 (Auth, Goods, Order, Chat, User, Admin, Category, Report)
│   ├── dto/          # 7 个数据传输对象
│   ├── entity/       # 10 个实体类
│   ├── exception/    # 全局异常处理
│   ├── repository/   # 9 个 MyBatis-Plus Mapper 接口
│   ├── security/     # JWT Token 工具 + Spring Security 配置 + 认证过滤器
│   └── service/      # 7 个业务服务
├── frontend/         # Vue 3 SPA + Element Plus
│   ├── api/          # Axios 封装 (含拦截器，自动附加 JWT Token)
│   ├── components/   # MainLayout 布局组件
│   ├── router/       # Vue Router 路由配置
│   ├── store/        # Pinia 用户状态管理
│   └── views/        # 12 个页面视图
├── db/schema.sql     # 完整 DDL (11张表) + 种子数据
└── docker-compose.yml''')

doc.add_page_break()

# ============================================================
# 二、核心模块与完整度
# ============================================================
doc.add_heading('二、核心模块与完整度', level=1)

doc.add_paragraph(
    '系统围绕"用户-商品-订单-消息-举报"五条主线设计了 8 个核心模块，'
    '覆盖从商品发布到交易完成再到信用积累的完整闭环。'
)

add_table_with_data(doc,
    ['序号', '模块', '对应 Controller', '主要功能', '完整度'],
    [
        ['1', '用户认证', 'AuthController', '学号实名注册 / 登录 / JWT 签发 / BCrypt 加密', '✅ 完整'],
        ['2', '商品管理', 'GoodsController', '发布(多图) / 编辑 / 上下架 / 删除 / 收藏', '✅ 完整'],
        ['3', '商品检索', 'GoodsController.list', '分类浏览 / 关键词搜索 / 排序 / 分页', '⚠️ 基本可用'],
        ['4', '交易流程', 'OrderController', '下单 → 双确认 → 取消 → 评价 → 信用分联动', '✅ 最完整'],
        ['5', '即时通讯', 'ChatController', 'WebSocket STOMP 实时推送 + REST 历史拉取', '✅ 完整'],
        ['6', '后台管理', 'AdminController', '仪表盘 / 用户管理 / 商品审核 / 举报处理', '✅ 完整'],
        ['7', '举报治理', 'ReportController', '提交举报 → 处理 → 下架 + 扣信用分(事务)', '✅ 完整'],
        ['8', '信用体系', '(内嵌于 Order/Report)', '初始 80 / 交易 +2 / 举报 -10 / 上下限 0~100', '✅ 完整'],
    ]
)

doc.add_heading('2.1 最完整的模块：交易流程', level=2)
doc.add_paragraph(
    '交易模块（OrderService + OrderController）是代码质量最高、逻辑最闭环的模块。'
    '关键设计亮点：'
)
bullets = [
    '下单时金额从数据库商品价格读取（goods.getPrice()），而非信任用户提交参数，防止价格篡改',
    '商品状态校验（必须是"在售"且不能购买自己的商品）',
    '创建订单与更新商品状态在同一 @Transactional 事务中，保证原子性',
    '买卖双方独立确认（buyerConfirm / sellerConfirm），任一方先确认后等待对方',
    '首次双方确认完成后自动各 +2 信用分（通过检查对方确认状态避免重复加分）',
    '取消订单后自动恢复商品为"在售"状态',
    '评价仅在订单完成后（status=2）才允许提交',
]
for b in bullets:
    p = doc.add_paragraph(b, style='List Bullet')

doc.add_page_break()

# ============================================================
# 三、商品状态流转
# ============================================================
doc.add_heading('三、商品状态流转机制', level=1)

doc.add_heading('3.1 状态码定义', level=2)
add_table_with_data(doc,
    ['状态码', '含义', '含义说明'],
    [
        ['-1', '已删除', '卖家标记删除（软删除，数据保留在数据库）'],
        ['0', '下架', '卖家手动下架 或 举报核实后被强制下架'],
        ['1', '在售', '正常展示在列表页，可被搜索和购买'],
        ['2', '已售出', '有人下单后自动变更，不可再被购买'],
        ['3', '审核中', '数据库定义了此状态，但当前代码发布时直接设为 1（在售）'],
    ]
)

doc.add_heading('3.2 实际流转路径', level=2)
add_code_block(doc, '''┌─────────────────────────────────────────────────────────────────┐
│                                                                 │
│  发布商品 ──→ status=1 (在售)                                    │
│                  │                                              │
│     ┌────────────┼────────────┬────────────┐                    │
│     ↓            ↓            ↓            ↓                    │
│  有人下单    卖家手动     管理员审核    卖家删除                  │
│     ↓            ↓            ↓            ↓                    │
│ status=2     status=0     status=1/0   status=-1               │
│ (已售出)     (下架)       (通过/驳回)   (已删除)                 │
│     │                                                            │
│     ├── 买家确认 ──┬── 双方都确认 → status=2 (完成)              │
│     │              │                +2 信用分（买卖双方）         │
│     │              │                                             │
│     ├── 卖家确认 ──┘                                             │
│     │                                                            │
│     └── 任一方取消 → status=-1 (订单取消)                        │
│                      → 商品恢复 status=1 (重新上架)               │
│                                                                 │
└─────────────────────────────────────────────────────────────────┘''')

add_note(doc,
    '当前 GoodsService.publishGoods() 第 92 行：status(1) 并注释"直接上架"。'
    'status=3（审核中）仅在 Schema DDL 中定义，发布流程未经过审核环节。'
    '管理员可通过 /api/admin/goods/{uuid}/audit 事后手动审核。',
    'warn'
)

doc.add_page_break()

# ============================================================
# 四、高级筛选
# ============================================================
doc.add_heading('四、高级筛选与成色字段', level=1)

doc.add_heading('4.1 当前筛选维度', level=2)
add_table_with_data(doc,
    ['维度', '前端控件', '后端对应字段', 'SQL 实现方式'],
    [
        ['分类筛选', '横向分类标签按钮（7个分类 + "全部"）', 'goods.category_id', '= #{categoryId} 精确匹配'],
        ['关键词搜索', '顶部搜索框', 'goods.title, goods.description', "LIKE CONCAT('%', #{keyword}, '%')"],
        ['排序方式', '下拉选择（最新发布 / 价格排序）', 'created_at / price', 'MyBatis <choose> 白名单限定'],
        ['排序方向', '升降序按钮', 'ASC / DESC', 'sortOrder 参数'],
        ['分页', '底部分页组件', '—', 'MyBatis-Plus Page 对象，默认 12 条/页'],
    ]
)

doc.add_heading('4.2 成色 (condition_level) 字段', level=2)
doc.add_paragraph(
    'condition_level 是 TINYINT 类型字段（1-10），定义在 goods 表中。'
    '仅在发布/编辑页面作为滑块控件使用，用于卖家标注商品成色：'
)
add_table_with_data(doc,
    ['值', '含义'],
    [
        ['1', '废品'],
        ['3', '可用'],
        ['5', '良好'],
        ['7', '较新'],
        ['10', '全新'],
    ]
)
add_note(doc,
    '列表页当前没有成色筛选功能。PageQuery DTO 中没有 conditionLevel 过滤参数，'
    'searchGoods SQL 中也没有相关 WHERE 条件。成色字段仅在商品详情页展示（如"全新"角标：conditionLevel >= 9）。',
    'warn'
)

doc.add_page_break()

# ============================================================
# 五、商品分类
# ============================================================
doc.add_heading('五、商品分类的数据来源', level=1)

doc.add_paragraph('分类数据不是写死在前端的，完整的数据流如下：')

add_code_block(doc, '''数据库 category 表                    ← 7 条种子数据
      │  (id, name, icon, sort_order, status)
      │
      ▼
后端 GET /api/category (公开接口)
      │  CategoryController → CategoryMapper.selectList()
      │
      ▼
前端 onMounted() 调用 categoryApi.list()
      │  Home.vue  → 渲染分类标签
      │  Publish.vue → 渲染下拉选择
      │
      ▼
页面展示：教材教辅 | 电子数码 | 生活用品 | 运动户外 | 服饰美妆 | 图书音像 | 其他闲置''')

doc.add_paragraph(
    '管理员可以通过直接修改 category 表来增删改分类（名称、图标、排序、启用/禁用），'
    '前端无需改动代码即可生效。种子数据在 schema.sql 和 DataInitializer.java 两处同步维护。'
)

doc.add_page_break()

# ============================================================
# 六、聊天系统通知
# ============================================================
doc.add_heading('六、聊天"系统通知"', level=1)

doc.add_paragraph(
    '消息表 message.msg_type 支持三种类型：TEXT、IMAGE、SYSTEM。'
    '但当前代码中 MessageService.sendMessage() 只是原样保存客户端传来的 msgType，'
    '没有任何服务端逻辑自动生成 SYSTEM 类型消息。'
)

doc.add_paragraph(
    '另外存在独立的 notification 表（系统通知表），字段包括 user_id、title、content、'
    'is_read 等，但同样没有对应的 Service 或 Controller 来创建通知。'
)

add_note(doc,
    '结论：系统通知的表结构（message.msg_type=SYSTEM 和独立的 notification 表）'
    '已预留，但实际业务代码中并未实现自动推送。'
    '典型的应该触发系统通知但未实现的场景包括："您的商品被收藏了"、"订单状态变更"、'
    '"商品审核结果"、"举报处理结果"等。',
    'warn'
)

doc.add_page_break()

# ============================================================
# 七、信用分
# ============================================================
doc.add_heading('七、信用分机制', level=1)

doc.add_heading('7.1 基本参数', level=2)
add_table_with_data(doc,
    ['参数', '值', '代码位置'],
    [
        ['初始值', '80 分', 'schema.sql: DEFAULT 80 / AuthService.java:49 creditScore(80)'],
        ['上限', '100 分', 'OrderService.java:186 Math.min(100, ...)'],
        ['下限', '0 分', 'ReportService.java:99 Math.max(0, ...)'],
        ['加分（交易完成）', '双方各 +2', 'OrderService.java:85-86 addCreditScore(userId, 2)'],
        ['扣分（举报核实）', '-10 分', 'ReportService.java:99 current - 10'],
    ]
)

doc.add_heading('7.2 触发时机', level=2)
add_code_block(doc, '''加分触发（OrderService.buyerConfirm / sellerConfirm）:
  买家点"确认收货" → 检查卖家是否也已确认
  卖家点"确认"     → 检查买家是否也已确认
  首次双方都确认   → 交易完成 + 各加 2 分（上限 100）

扣分触发（ReportService.applyReportPenalty）:
  管理员"处理"举报 → status=1（已处理）
  → 被举报商品强制下架 (goods.status=0)
  → 被举报用户信用分 -10（下限 0）''')

doc.add_paragraph(
    '演示中看到的 87、92、77 等分值不是默认值——那是种子数据为测试用户设定的初始值'
    '（张三 85、李四 90、王五 75），模拟不同信用等级的用户。'
)

doc.add_page_break()

# ============================================================
# 八、管理员权限
# ============================================================
doc.add_heading('八、管理员权限', level=1)

doc.add_heading('8.1 为什么切换到 admin 账号', level=2)
doc.add_paragraph(
    '演示中切换到 admin 账号是为了展示后台管理功能——普通用户看不到管理后台入口，'
    '也无法访问 /api/admin/** 接口。这是系统权限分离设计的体现。'
)

doc.add_heading('8.2 权限对比', level=2)
add_table_with_data(doc,
    ['功能', '普通用户 (role=0)', '管理员 (role=1)'],
    [
        ['浏览商品 / 搜索 / 收藏', '✅', '✅'],
        ['发布商品 / 编辑 / 上下架', '✅', '✅'],
        ['下单购买 / 确认收货', '✅', '✅'],
        ['实时聊天', '✅', '✅'],
        ['数据仪表盘（用户数/商品数/订单数/举报数）', '❌', '✅'],
        ['用户管理（列表/详情/禁用/启用）', '❌', '✅'],
        ['商品审核（通过/驳回）', '❌', '✅'],
        ['举报管理（列表/处理/驳回）', '❌', '✅'],
    ]
)

doc.add_heading('8.3 权限控制实现', level=2)
doc.add_paragraph('采用双层防护确保管理员接口不被越权访问：')
add_code_block(doc, '''第一层：SecurityConfig.java（URL 级别）
  .requestMatchers("/api/admin/**").hasRole("ADMIN")

第二层：AdminController.java（类级别）
  @RestController
  @RequestMapping("/api/admin")
  @PreAuthorize("hasRole('ADMIN')")    ← 类上注解，所有方法继承
  public class AdminController { ... }

JWT Token 中 role=1 → 解析为 ROLE_ADMIN 权限
JWT Token 中 role=0 → 解析为 ROLE_USER 权限 → 访问 /api/admin/** 返回 403''')

doc.add_page_break()

# ============================================================
# 九、面交地点时间
# ============================================================
doc.add_heading('九、面交地点与时间', level=1)

add_table_with_data(doc,
    ['问题', '答案'],
    [
        ['存在哪里？', 'orders 表的 meet_location (VARCHAR 256) 和 meet_time (DATETIME) 两个字段'],
        ['谁填写？', '买家在下单时填写，通过 OrderController.create() 的 @RequestParam 传入'],
        ['谁可以看？', '该订单的买卖双方都可以在订单详情中看到，用于协调面交'],
        ['是否必填？', 'required=false，两个字段均为可选参数'],
        ['前端展示', 'Orders.vue 订单列表的 detail-item 区域，交易方式为面交(FACE)时显示地点和时间'],
    ]
)

doc.add_page_break()

# ============================================================
# 十、浏览量 & AI
# ============================================================
doc.add_heading('十、浏览量机制与 AI 功能', level=1)

doc.add_heading('10.1 浏览量', level=2)
doc.add_paragraph(
    '是的，每次打开商品详情页（调用 GET /api/goods/detail/{uuid}），'
    'GoodsService.getByUuid() 方法中会执行 viewCount + 1 并更新数据库：'
)
add_code_block(doc, '''// GoodsService.java:53-54
goods.setViewCount(goods.getViewCount() + 1);
goodsMapper.updateById(goods);''')
add_note(doc,
    '无去重、无限流——同一用户反复刷新会持续累加浏览量。没有按 session/IP/用户ID 去重的逻辑。',
    'warn'
)

doc.add_heading('10.2 AI 自动分类与 AI 定价建议', level=2)
add_note(doc,
    '❌ 本项目未实现任何 AI 功能。AI 自动分类和 AI 定价建议在当前版本的代码中均不存在。',
    'cross'
)

doc.add_paragraph(
    '经过全项目关键词检索（AI / LLM / OpenAI / ChatGPT / 定价 / 自动分类 / classify），确认：'
)
bullets_ai = [
    '后端无任何 HTTP 调用外部 AI API（如 OpenAI、百度文心、阿里通义等）的代码',
    '前端无任何 AI 相关交互界面或功能入口',
    'pom.xml 中无任何 AI SDK 依赖',
    '设计文档（D-03 概要设计说明书）中"AI 定价"仅出现在"预留扩展点"列表——'
    '原话是"预留 ≥3 个扩展点（推荐系统、邮件柜接入、AI 定价）"，这是规划中的功能，未实际开发',
    '答辩文档（D-11）中的"AI 工具使用申报"指的是团队在开发过程中使用 AI 辅助'
    '（写文档、分析日志、排查错误），而非产品功能特性',
]
for b in bullets_ai:
    doc.add_paragraph(b, style='List Bullet')

doc.add_paragraph(
    '因此，关于"调用的哪个大模型/API""喂给模型的输入是什么""定价区间 ¥38~¥3213 怎么算出来的"'
    '"如何防止 LLM 一本正经地报离谱价""AI 接口超时限流怎么兜底""用户能否不采纳 AI 建议"'
    '——这些问题在当前版本中均不适用，这是未实现的规划功能。'
)

doc.add_page_break()

# ============================================================
# 十一、并发与一致性
# ============================================================
doc.add_heading('十一、并发与一致性', level=1)

doc.add_heading('11.1 超卖防护', level=2)
add_note(doc,
    '❌ 当前无并发防护。两人同时点"立即购买"理论上可能超卖。',
    'cross'
)
doc.add_paragraph(
    'OrderService.createOrder() 的逻辑是：查询商品 → 检查 status != 1 → 插入订单 → 更新 status=2。'
    '虽然标注了 @Transactional，但没有使用任何锁机制：'
)
add_code_block(doc, '''// 当前代码（有并发风险）:
Goods goods = goodsMapper.findByUuid(goodsUuid);  // 普通 SELECT，无锁
if (goods.getStatus() != 1) {
    throw new RuntimeException("商品已下架或已售出");
}
// ... 插入订单 ...
goods.setStatus(2);
goodsMapper.updateById(goods);''')
doc.add_paragraph(
    '两个并发请求可能同时通过 status != 1 的检查，导致两人都成功下单。'
    '正确的做法是使用 SELECT ... FOR UPDATE 悲观锁或添加 version 字段做乐观锁。'
    '答辩 Q&A 中承认这是后续优化点："并发条件下的乐观锁控制"。'
)

doc.add_heading('11.2 超时自动确认', level=2)
add_note(doc, '❌ 没有超时自动确认或取消机制。', 'cross')
doc.add_paragraph(
    '订单创建后 status=0（待确认），完全依赖买卖双方手动点击"确认收货"按钮。'
    '没有定时任务（@Scheduled）、没有消息队列延迟消费、没有任何超时处理逻辑。'
    '答辩文档提到"未来可以加超时自动取消机制——下单 24 小时内未完成确认则自动取消"。'
)

doc.add_heading('11.3 取消后自动上架', level=2)
add_note(doc, '✅ 已实现。', 'success')
doc.add_paragraph(
    '取消订单时（OrderService.cancelOrder()），在同一个 @Transactional 事务中自动将商品状态恢复为 1（在售）：'
)
add_code_block(doc, '''// OrderService.java:128-133
// 恢复商品为上架状态
Goods goods = goodsMapper.findByUuid(order.getGoodsUuid());
if (goods != null) {
    goods.setStatus(1);
    goodsMapper.updateById(goods);
}''')

doc.add_page_break()

# ============================================================
# 十二、实时聊天
# ============================================================
doc.add_heading('十二、实时聊天技术实现', level=1)

doc.add_heading('12.1 技术方案', level=2)
add_table_with_data(doc,
    ['方面', '实现方式', '详细说明'],
    [
        ['通信协议', 'WebSocket + STOMP', '非轮询。前端 @stomp/stompjs + sockjs-client'],
        ['消息持久化', '✅ MySQL message 表', '每条消息 INSERT 入库（含 UUID 去重标识）'],
        ['实时推送', 'SimpMessagingTemplate\n.convertAndSendToUser()', '点对点推送到 /user/{id}/queue/chat'],
        ['历史消息', 'GET /api/chat/conversation/{id}', '分页拉取两人间的完整对话记录'],
        ['已读标记', 'message.is_read 字段', '拉取对话时自动标记对方发来的消息为已读'],
        ['联系人列表', 'GET /api/chat/contacts', 'DISTINCT 查询所有有过对话的联系人'],
        ['HTTP 兜底', 'POST /api/chat/send (REST)', 'WebSocket 不可用时降级为 HTTP 发送'],
    ]
)

doc.add_heading('12.2 WebSocket 配置', level=2)
add_code_block(doc, '''// WebSocketConfig.java
@Configuration
@EnableWebSocketMessageBroker
public class WebSocketConfig implements WebSocketMessageBrokerConfigurer {
    @Override
    public void configureMessageBroker(MessageBrokerRegistry config) {
        config.enableSimpleBroker("/topic", "/queue", "/user");
        config.setApplicationDestinationPrefixes("/app");
        config.setUserDestinationPrefix("/user");
    }

    @Override
    public void registerStompEndpoints(StompEndpointRegistry registry) {
        registry.addEndpoint("/ws/chat")
                .setAllowedOriginPatterns("*")
                .withSockJS();                    // SockJS 兜底
    }
}''')

doc.add_heading('12.3 离线消息处理', level=2)
doc.add_paragraph(
    '消息已持久化到 MySQL，对方离线时消息不会丢失。重新上线后：'
)
bullets_chat = [
    '通过 GET /api/chat/conversation/{contactId} 拉取未读历史消息',
    '实时的 WebSocket 推送仅对当前在线且 WebSocket 连接活跃的用户生效',
    'is_read 标记：发送时默认 0（未读），接收方拉取对话后自动批量标记为 1（已读）',
]
for b in bullets_chat:
    doc.add_paragraph(b, style='List Bullet')

doc.add_heading('12.4 多端登录同步', level=2)
add_note(doc,
    '未做多端同步处理。SimpMessagingTemplate.convertAndSendToUser() 按用户 ID 推送，'
    '但同一用户多端 WebSocket 连接时只有当前活跃连接能收到实时推送。'
    '消息在 DB 中持久化，多端均可以拉取历史记录，所以不会丢失消息，但无法实时多端同步。',
    'warn'
)

doc.add_page_break()

# ============================================================
# 十三、安全与权限
# ============================================================
doc.add_heading('十三、安全与权限', level=1)

doc.add_heading('13.1 鉴权流程', level=2)
add_code_block(doc, '''┌──────────┐     ①POST /api/auth/login       ┌──────────┐
│  前端     │ ────────────────────────────────→ │  后端     │
│ (Vue 3)  │ ←──────────────────────────────── │ (Spring) │
└──────────┘   ②返回 JWT Token + 用户信息        └──────────┘
     │
     │ ③Token 存入 localStorage
     │ ④每次请求自动附加 Authorization: Bearer <token>
     │
     ▼
┌──────────────────────────────────────────────────────────────┐
│ JwtAuthenticationFilter (OncePerRequestFilter)               │
│   → extractToken(): 从 Authorization 头提取 Bearer Token     │
│   → jwtTokenProvider.validateToken(): 验证签名+过期          │
│   → 解析 userId / username / role → 设置 SecurityContext     │
│   → role=1 → ROLE_ADMIN / role=0 → ROLE_USER                │
└──────────────────────────────────────────────────────────────┘''')

doc.add_heading('13.2 防止价格篡改', level=2)
doc.add_paragraph(
    '关键防护在 OrderService.createOrder() 第 54 行——订单金额取自数据库而非用户提交：'
)
add_code_block(doc, '''// 正确做法：从数据库读取商品价格
.amount(goods.getPrice())   // ← 卖家设定的价格，不可被买家篡改

// 如果错误地写成：
// .amount(request.getAmount())  ← 用户可通过抓包改参数，价格变为 0''')
doc.add_paragraph(
    '配合 @Valid 参数校验 + 全局异常处理（GlobalExceptionHandler），'
    '对用户提交的参数进行合法性检查（非空、类型校验等）。'
)

doc.add_heading('13.3 Admin 接口防护', level=2)
add_code_block(doc, '''第一层 — SecurityConfig (URL 级别):
  .requestMatchers("/api/admin/**").hasRole("ADMIN")

第二层 — AdminController (类级别注解):
  @RestController
  @PreAuthorize("hasRole('ADMIN')")  // 方法执行前再次校验
  public class AdminController { ... }

第三层 — JWT Token 角色嵌入:
  Token 中包含 userId + username + role
  普通用户的 role=0 → ROLE_USER → 访问 admin 接口返回 403 Forbidden
  管理员的 role=1 → ROLE_ADMIN → 允许访问

全局异常处理:
  AccessDeniedException → 403 "无权限访问"''')

doc.add_heading('13.4 密码存储', level=2)
add_table_with_data(doc,
    ['项目', '说明'],
    [
        ['加密方式', 'BCrypt 加盐哈希 (org.springframework.security.crypto.bcrypt)'],
        ['Salt', 'BCrypt 自动生成随机盐，内嵌于哈希结果中'],
        ['哈希格式', '$2a$10$<22位salt+hash> — 10 为 cost factor (2^10 = 1024 轮)'],
        ['加密时机', '注册时 (AuthService.register) 调用 passwordEncoder.encode()'],
        ['验证时机', '登录时 (AuthService.login) 调用 passwordEncoder.matches()'],
        ['明文存储', '❌ 从未存储明文密码'],
    ]
)

doc.add_page_break()

# ============================================================
# 十四、审核与举报
# ============================================================
doc.add_heading('十四、审核与举报流程', level=1)

doc.add_heading('14.1 商品审核', level=2)
add_note(doc,
    '商品"审核中"状态（status=3）已在数据库 DDL 中定义，但当前发布流程直接上架（status=1），'
    '不走审核环节。管理员可事后通过管理后台手动审核。这是事前审核缺失、事后管理兜底的现状。',
    'warn'
)
add_code_block(doc, '''// GoodsService.publishGoods() — 当前行为:
goods.setStatus(1);  // 直接上架，不走审核

// AdminController.auditGoods() — 事后审核:
PUT /api/admin/goods/{uuid}/audit?status=1  → 通过（上架）
PUT /api/admin/goods/{uuid}/audit?status=0  → 驳回（下架）''')

doc.add_heading('14.2 举报处理流程', level=2)
add_code_block(doc, '''┌─────────────────────────────────────────────────────────────┐
│                                                             │
│  用户提交举报 (ReportService.createReport)                    │
│    ├── 校验：被举报商品或用户至少填一项                        │
│    ├── 校验：被举报对象确实存在                               │
│    └── 写入 report 表，status=0 (待处理)                     │
│                         │                                    │
│                         ▼                                    │
│  管理员在后台查看举报列表                                     │
│    ├── 查看到所有举报（分页）                                 │
│    └── 每条举报显示：原因、描述、状态、时间                    │
│                         │                                    │
│              ┌──────────┴──────────┐                         │
│              ▼                     ▼                         │
│         点击"处理"            点击"驳回"                      │
│      status → 1 (已处理)    status → 2 (已驳回)              │
│              │                     │                         │
│              ▼                     ▼                         │
│      applyReportPenalty()    无处罚                          │
│        ├── 被举报商品 → status=0 (下架)                       │
│        └── 被举报用户 → credit_score -10 (下限0)              │
│                                                             │
│  记录：handler_id, handle_remark, handled_at                 │
│                                                             │
└─────────────────────────────────────────────────────────────┘''')

doc.add_paragraph('整个处理过程在 ReportService.handleReport() 的 @Transactional 事务中完成，保证举报状态更新、商品下架、信用分扣减的原子性。')

doc.add_page_break()

# ============================================================
# 十五、性能与数据
# ============================================================
doc.add_heading('十五、性能与数据存储', level=1)

doc.add_heading('15.1 分页策略', level=2)
add_table_with_data(doc,
    ['页面', '每页条数', '分页实现'],
    [
        ['首页商品列表', '12 条/页', 'MyBatis-Plus Page + LIMIT/OFFSET'],
        ['我的发布 / 收藏', '50 条/页', '同上'],
        ['聊天对话记录', '20 条/页', '同上'],
        ['联系人列表', '50 条/页', '同上'],
        ['管理后台-用户', '50 条/页', '同上'],
        ['管理后台-举报', '50 条/页', '同上'],
    ]
)

doc.add_heading('15.2 索引设计', level=2)
add_table_with_data(doc,
    ['表名', '索引名', '索引类型', '字段', '用途'],
    [
        ['goods', 'idx_uuid', '普通索引', 'uuid', '商品详情查询（高频）'],
        ['goods', 'idx_seller', '普通索引', 'seller_id', '卖家商品列表'],
        ['goods', 'idx_category', '普通索引', 'category_id', '分类筛选'],
        ['goods', 'idx_status', '普通索引', 'status', '状态筛选（在售商品）'],
        ['goods', 'idx_title', '普通索引', 'title', '标题查询'],
        ['goods', 'ft_title_desc', 'FULLTEXT', 'title, description', '全文搜索（已建未用）'],
        ['user', 'idx_student_id', '普通索引', 'student_id', '学号唯一校验'],
        ['user', 'idx_username', '普通索引', 'username', '用户名唯一校验 + 登录查询'],
        ['orders', 'idx_buyer', '普通索引', 'buyer_id', '买家订单列表'],
        ['orders', 'idx_seller', '普通索引', 'seller_id', '卖家订单列表'],
        ['orders', 'idx_status', '普通索引', 'status', '订单状态筛选'],
        ['message', 'idx_sender', '普通索引', 'sender_id', '发件人查询'],
        ['message', 'idx_receiver', '普通索引', 'receiver_id', '收件人查询'],
        ['message', 'idx_created', '普通索引', 'created_at', '消息时间排序'],
        ['favorite', 'uk_user_goods', 'UNIQUE', '(user_id, goods_uuid)', '防重复收藏 + 快速查询'],
    ]
)

doc.add_heading('15.3 搜索实现', level=2)
add_note(doc,
    '当前实际使用的是 SQL LIKE 模糊匹配，而非 MySQL FULLTEXT 全文索引。'
    '设计文档中描述了 MATCH...AGAINST 方案但代码未采用。',
    'warn'
)
add_table_with_data(doc,
    ['方案', '当前实现（代码）', '设计文档描述'],
    [
        ['SQL 方式', "LIKE CONCAT('%', #{keyword}, '%')", 'MATCH(title, description) AGAINST(#{keyword} IN BOOLEAN MODE)'],
        ['中文分词', '❌ 无（LIKE 只能做子串匹配）', '未实现'],
        ['搜索引擎', '❌ 无（未接入 ES/Solr）', '未涉及'],
        ['索引利用', '使用 idx_title 普通索引', 'FULLTEXT 索引已建但未被使用'],
    ]
)

doc.add_heading('15.4 图片存储', level=2)
add_note(doc,
    '实际存储方式为本地文件系统，而非 MinIO 对象存储。'
    'docker-compose 中配置了 MinIO 容器，但代码中 FileService 没有调用 MinIO SDK。',
    'warn'
)
add_table_with_data(doc,
    ['方面', '当前实现'],
    [
        ['存储位置', '本地文件系统：./uploads/goods/{goodsUuid}/{filename}'],
        ['访问 URL', '/uploads/goods/{goodsUuid}/{filename}（通过 Spring 静态资源映射）'],
        ['文件名生成', 'UUID (hutool IdUtil.fastSimpleUUID) + 原始扩展名'],
        ['前端限制', '最多 9 张图片，格式 jpg/png，单张不超过 5MB（仅 Element Plus 组件提示）'],
        ['后端校验', '❌ 无大小限制、无 MIME 类型校验、无病毒扫描、无图片压缩'],
        ['CDN', '❌ 未使用'],
        ['MinIO 状态', '容器已启动（端口 9000/9001），但代码未对接'],
    ]
)

doc.add_page_break()

# ============================================================
# 十六、总结
# ============================================================
doc.add_heading('十六、总结：设计与实现的差距', level=1)

doc.add_paragraph(
    '以下汇总了在源码审查中发现的"设计文档描述 vs 实际代码实现"之间的主要差异，'
    '供答辩时如实说明：'
)

add_table_with_data(doc,
    ['维度', '设计文档描述', '实际代码实现', '影响评估'],
    [
        ['对象存储', 'MinIO（docker-compose 已配置）', '本地文件系统 (FileService)', '演示可用，生产需迁移'],
        ['全文搜索', 'MySQL FULLTEXT + BOOLEAN MODE', 'SQL LIKE 模糊匹配', '中文搜索精度低，性能随数据增长下降'],
        ['商品审核', 'status=3 审核中（发布→审核→上架）', '发布直接 status=1 在售', '事前审核缺失'],
        ['AI 定价/分类', '预留扩展点（概要设计）', '完全未实现', '非当前版本功能'],
        ['并发控制', '—（未涉及）', '无锁（SELECT 后 UPDATE 存在竞态）', '超卖风险'],
        ['超时自动取消', '—（未涉及）', '无定时任务', '订单可能永久挂起'],
        ['系统通知', 'notification 表 + message.msg_type=SYSTEM', '表结构存在但无推送代码', '用户无法收到状态变更通知'],
        ['中文分词', '—（未涉及）', '未实现', '搜索体验受限'],
        ['后端图片校验', '—（未涉及）', '无大小/格式/MIME 校验', '安全风险'],
    ]
)

doc.add_heading('16.1 项目亮点', level=2)
bullets_good = [
    '交易模块最完整：下单校验 → 双确认状态机 → 信用分联动 → 取消回滚，全事务保护',
    '安全防护到位：JWT 无状态认证、BCrypt 加盐、价格取 DB 防篡改、Admin 双层权限控制',
    '实时聊天：WebSocket STOMP 推送 + MySQL 持久化 + REST 历史拉取，离线消息不丢失',
    '举报治理闭环：举报→处理→下架+扣分，同一事务保证原子性',
    '数据库设计规范：外键约束、唯一索引、FULLTEXT 索引、utf8mb4 编码、完整的 ER 关系',
    '部署方案齐全：Docker Compose（开发/演示）+ K8s Kustomize（生产）双方案',
    '种子数据丰富：202 条商品 + 4 个用户 + 7 个分类，开箱即用',
]
for b in bullets_good:
    doc.add_paragraph(b, style='List Bullet')

doc.add_heading('16.2 已知待改进项', level=2)
bullets_improve = [
    '并发控制：加乐观锁（version 字段）或 SELECT ... FOR UPDATE 防止超卖',
    '超时机制：加 @Scheduled 定时任务自动取消超时未确认的订单',
    '搜索升级：启用 MySQL FULLTEXT 或接入 Elasticsearch + 中文分词（IK Analyzer）',
    '图片存储：对接 MinIO SDK 替代本地文件系统；加后端图片大小/格式/MIME 校验',
    '审核流程：支持"发布→审核中→管理员审核→上架"的完整事前审核链路',
    '系统通知：基于 Spring Event 实现关键业务事件的消息/通知自动推送',
    'WebSocket 鉴权：在 STOMP CONNECT 阶段校验 JWT Token',
    '测试覆盖：补充 GoodsService 和 OrderService 的单元测试 + API 集成测试',
]
for b in bullets_improve:
    doc.add_paragraph(b, style='List Bullet')

# ============================================================
# 保存
# ============================================================
output_path = '/Users/admin/Desktop/SwapCampus/docs/SwapCampus_答辩QA问答汇编.docx'
doc.save(output_path)
print(f'文档已生成: {output_path}')
