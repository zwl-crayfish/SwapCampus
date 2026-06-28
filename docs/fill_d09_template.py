#!/usr/bin/env python3
"""
Fill the D-09 课程设计总结报告 template with SwapCampus project content.
Preserves all formatting, paragraph styles, and non-content tables.
"""

from docx import Document
from docx.shared import Pt, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from copy import deepcopy
import datetime

# ── Configuration ──────────────────────────────────────────
TEMPLATE = 'docs/产出模板/D-09_课程设计总结报告_模板.docx'
OUTPUT = 'docs/D-09_课程设计总结报告_SwapCampus.docx'

FONT_NAME = 'Microsoft YaHei'
BODY_SIZE = Pt(10.5)       # 133350 EMU = 10.5pt
H2_SIZE = Pt(15)           # 190500 EMU = 15pt
H1_SIZE = Pt(18)           # 228600 EMU = 18pt

TEAM_NO = 'G18'
TEAM_LEADER = '张文乐'
MEMBERS = [
    ('计算机231', '231002413', '卢天翔'),
    ('计算机231', '231002414', '陈宣妤'),
    ('计算机231', '231002419', '王洁'),
    ('计算机231', '231002420', '张文乐'),
    ('计算机231', '231002412', '桂高彬'),
]
ADVISOR = '（指导教师）'
DATE_TODAY = '2026-06-15'

# ── Helpers ────────────────────────────────────────────────

def clear_cell(cell):
    """Remove all paragraphs from a cell except keep one empty paragraph."""
    for p in cell.paragraphs:
        for r in p.runs:
            r.text = ''
    # Remove extra paragraphs
    while len(cell.paragraphs) > 1:
        p = cell.paragraphs[-1]
        p._element.getparent().remove(p._element)

def set_cell_text(cell, text, bold=False, size=BODY_SIZE, font_name=FONT_NAME):
    """Set single-cell text with formatting."""
    # Clear existing content
    for p in cell.paragraphs:
        for r in p.runs:
            r.text = ''
    # Set text in first paragraph
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # Clear and set
    for r in p.runs:
        r.text = ''
    if p.runs:
        run = p.runs[0]
        run.text = text
        run.font.name = font_name
        run.font.size = size
        run.font.bold = bold
    else:
        # No existing run — add one
        run = p.add_run(text)
        run.font.name = font_name
        run.font.size = size
        run.font.bold = bold

def add_paragraph_to_cell(cell, text, bold=False, size=BODY_SIZE, font_name=FONT_NAME,
                           alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=Pt(6)):
    """Add a new paragraph to a table cell."""
    p = cell.add_paragraph()
    p.alignment = alignment
    pf = p.paragraph_format
    pf.space_after = space_after
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = size
    run.font.bold = bold
    return p

def fill_cell_with_paragraphs(cell, paragraphs_data):
    """
    Fill a cell with multiple paragraphs.
    paragraphs_data is a list of (text, bold, size, alignment) tuples.
    """
    # Clear existing paragraphs
    for p in cell.paragraphs:
        for r in p.runs:
            r.text = ''

    first = True
    for pd in paragraphs_data:
        if isinstance(pd, str):
            text, bold, size, alignment = pd, False, BODY_SIZE, WD_ALIGN_PARAGRAPH.LEFT
        else:
            text = pd[0]
            bold = pd[1] if len(pd) > 1 else False
            size = pd[2] if len(pd) > 2 else BODY_SIZE
            alignment = pd[3] if len(pd) > 3 else WD_ALIGN_PARAGRAPH.LEFT

        if first and cell.paragraphs:
            p = cell.paragraphs[0]
            p.alignment = alignment
            for r in p.runs:
                r.text = ''
            if p.runs:
                run = p.runs[0]
            else:
                run = p.add_run('')
            run.text = text
            run.font.name = FONT_NAME
            run.font.size = size
            run.font.bold = bold
            first = False
        else:
            add_paragraph_to_cell(cell, text, bold=bold, size=size, alignment=alignment)


# ── SECTION CONTENT ────────────────────────────────────────

def get_content():
    """Return a dict mapping table_index -> filler function."""

    sections = {}

    # TABLE[0]: Cover metadata (7r x 2c)
    def fill_table0(doc):
        table = doc.tables[0]
        # Row 4: 小组/组号
        set_cell_text(table.rows[4].cells[1], f'{TEAM_NO} {TEAM_LEADER}小组')
        # Row 5: 班级/学号/姓名
        names = '\n'.join([f'{c} / {sid} / {n}' for c, sid, n in MEMBERS])
        set_cell_text(table.rows[5].cells[1], names)
        # Row 6: 完成日期
        set_cell_text(table.rows[6].cells[1], DATE_TODAY)
    sections[0] = fill_table0

    # TABLE[1]: Usage note — KEEP AS IS (do nothing)
    # TABLE[2]: Cover page note — KEEP AS IS

    # TABLE[3]: Section 2 摘要 (Chinese + English)
    def fill_table3(doc):
        cell = doc.tables[3].rows[0].cells[0]
        clear_cell(cell)
        paras = [
            ('【中文摘要】', True, BODY_SIZE),
            ('', False, BODY_SIZE),
            ('SwapCampus 是一个面向北京林业大学师生的 C2C 校园闲置物品交易平台。系统采用前后端分离架构，基于 Spring Boot 3.2 + Vue 3 技术栈，实现了用户注册登录、商品发布与浏览、订单交易、站内即时通讯（WebSocket）、商品收藏、信用积分评价以及后台管理等核心功能。项目遵循软件工程标准流程，在 13 天集中实习期内完成了从立项、需求分析、概要设计、详细设计、编码实现、测试到部署交付的完整生命周期。本报告全面总结了项目的需求分析、系统设计、编码实现、测试验证、部署交付以及项目管理等方面的实践过程与经验收获。', False, BODY_SIZE),
            ('', False, BODY_SIZE),
            ('关键词：校园闲置交易；Spring Boot；Vue 3；WebSocket；软件工程课程设计', False, BODY_SIZE),
            ('', False, BODY_SIZE),
            ('', False, BODY_SIZE),
            ('【Abstract】', True, BODY_SIZE),
            ('', False, BODY_SIZE),
            ('SwapCampus is a C2C campus idle-item trading platform designed for Beijing Forestry University students and faculty. Built on a front-end/back-end separation architecture using Spring Boot 3.2 and Vue 3, the system implements core features including user registration and login, product publishing and browsing, order transactions, real-time instant messaging (WebSocket), product favorites, credit score evaluation, and admin management. The project followed standard software engineering processes, completing the full lifecycle from project initiation, requirements analysis, outline design, detailed design, coding implementation, testing, to deployment and delivery within a 13-day intensive internship period. This report comprehensively summarizes the practical processes and lessons learned in requirements analysis, system design, implementation, testing, deployment, and project management.', False, BODY_SIZE),
            ('', False, BODY_SIZE),
            ('Keywords: campus idle trading; Spring Boot; Vue 3; WebSocket; software engineering course design', False, BODY_SIZE),
        ]
        fill_cell_with_paragraphs(cell, paras)
    sections[3] = fill_table3

    # TABLE[4]: Section 3.1 项目背景与意义
    def fill_table4(doc):
        cell = doc.tables[4].rows[0].cells[0]
        clear_cell(cell)
        paras = [
            ('北京林业大学校内每年产生大量闲置物品，包括教材、电子产品、生活用品等。当前校内二手交易主要通过微信群进行散点式交易，存在信息碎片化严重、信用难以追溯、交易效率低下等问题。学校学生处希望建设一个面向本校师生的闲置物品交易平台，以提高校内资源流转率，同时为同学提供安全、便捷的交易体验。', False, BODY_SIZE),
            ('', False, BODY_SIZE),
            ('本项目是《软件工程（课程设计）》的课程项目（选题编号 T-02），要求以 5 人团队，在 2 周（13 天）集中实习期内，完成从立项到答辩的完整软件工程实践。项目的核心价值主张是"让校园闲置流转更安全、更高效、更可信"——通过学号实名注册 + JWT 鉴权 + 信用积分体系，构建可信交易社区，解决传统二手交易中信息不对称和信用缺失的痛点。', False, BODY_SIZE),
            ('', False, BODY_SIZE),
            ('本项目的实践意义在于：第一，为校园师生提供一个统一、规范、可信的闲置物品交易渠道；第二，作为软件工程课程设计的完整案例，使团队成员在真实项目中实践需求分析、系统设计、编码实现、测试、部署等软件工程各阶段的方法和工具；第三，通过团队协作和项目管理，培养工程化思维、沟通协调能力和质量保障意识。', False, BODY_SIZE),
        ]
        fill_cell_with_paragraphs(cell, paras)
    sections[4] = fill_table4

    # TABLE[5]: Section 3.2 项目目标
    def fill_table5(doc):
        cell = doc.tables[5].rows[0].cells[0]
        clear_cell(cell)
        paras = [
            ('本项目的总体目标是开发一个功能完整、稳定可用的校园闲置物品交易平台，并通过完整的软件工程文档体系展示团队的工程实践能力。具体目标包括：', False, BODY_SIZE),
            ('', False, BODY_SIZE),
            ('（1）功能目标：实现用户体系（注册/登录/信用分）、商品管理（发布/编辑/下架/搜索）、交易流程（下单/确认收货/评价）、即时通讯（WebSocket 实时聊天）、收藏系统、后台管理（用户管理/商品审核/举报处理/数据仪表盘）六大模块，覆盖选题指南要求的全部必选功能及部分选做功能。', False, BODY_SIZE),
            ('（2）质量目标：核心功能模块（Auth + Goods + Order）单元测试覆盖率 ≥ 80%，整体覆盖率 ≥ 60%；系统通过功能测试、集成测试和安全测试的验证；关键 API 端到端延迟满足性能指标。', False, BODY_SIZE),
            ('（3）工程目标：按时完成 GB/T 8567—2006 规范要求的全部产出文档（D-01 至 D-09），总代码量 ≥ 3000 行，遵循 Conventional Commits 规范管理 Git 提交，实现 GitHub Actions CI/CD 自动化构建流水线。', False, BODY_SIZE),
            ('（4）协作目标：通过每日站会、Git Feature Branch 工作流、Code Review 等实践，建立高效的团队协作机制，积累真实工程经验。', False, BODY_SIZE),
        ]
        fill_cell_with_paragraphs(cell, paras)
    sections[5] = fill_table5

    # TABLE[6]: Section 3.3 本文组织结构
    def fill_table6(doc):
        cell = doc.tables[6].rows[0].cells[0]
        clear_cell(cell)
        paras = [
            ('本文是 SwapCampus 课程设计的最终总结报告，共分 13 章，组织结构如下：', False, BODY_SIZE),
            ('第 1 章：封面与目录。', False, BODY_SIZE),
            ('第 2 章：摘要——中英文双语摘要，概括项目背景、核心工作与主要成果。', False, BODY_SIZE),
            ('第 3 章：引言——阐述项目背景与意义、项目目标以及本文的组织结构。', False, BODY_SIZE),
            ('第 4 章：需求分析（精炼版）——从 SRS 中提炼核心需求，包括用户场景、功能性需求和非功能性需求。', False, BODY_SIZE),
            ('第 5 章：系统设计（精炼版）——从概要设计和详细设计中提炼系统架构、关键模块设计、数据库设计、接口设计和关键算法。', False, BODY_SIZE),
            ('第 6 章：系统实现——描述开发环境与工具链、关键技术实现细节以及实现过程中的关键决策。', False, BODY_SIZE),
            ('第 7 章：系统测试——总结测试策略、关键测试结果和缺陷统计与质量评估。', False, BODY_SIZE),
            ('第 8 章：系统部署与演示——展示部署架构、关键界面截图和演示链接。', False, BODY_SIZE),
            ('第 9 章：项目管理总结——回顾团队组建与分工、里程碑达成、风险应对和经验教训。', False, BODY_SIZE),
            ('第 10 章：AI 工具使用申报——申报项目中使用的 AI 工具及其用途和占比。', False, BODY_SIZE),
            ('第 11 章：结论与展望——总结主要工作，进行自我评价，分析不足并提出改进方向。', False, BODY_SIZE),
            ('第 12 章：参考文献。', False, BODY_SIZE),
            ('第 13 章：附录。', False, BODY_SIZE),
        ]
        fill_cell_with_paragraphs(cell, paras)
    sections[6] = fill_table6

    # TABLE[7]: Separator between §3 and §4 — remove placeholder, add a brief transition
    def fill_table7(doc):
        cell = doc.tables[7].rows[0].cells[0]
        clear_cell(cell)
        add_paragraph_to_cell(cell, '（以上为引言部分，以下进入需求分析的精炼提炼。）',
                              bold=False, size=Pt(9))
    sections[7] = fill_table7

    # TABLE[8]: Section 4.1 用户与场景
    def fill_table8(doc):
        cell = doc.tables[8].rows[0].cells[0]
        clear_cell(cell)
        paras = [
            ('SwapCampus 系统面向三类用户角色，其典型使用场景如下：', False, BODY_SIZE),
            ('', False, BODY_SIZE),
            ('【买家场景】', True, BODY_SIZE),
            ('张同学（大三计算机专业）想买一本《算法导论》二手教材。他打开 SwapCampus 首页，在搜索框输入"算法导论"，浏览搜索结果后点击商品详情页，查看书籍图片、成色（8成新）和价格（¥25）。他通过站内聊天联系卖家确认书目版本后，点击"立即购买"并选择"面交"方式，约定在图书馆门口交易。双方见面确认收货后，各自确认订单完成，信用分各增加 2 分。', False, BODY_SIZE),
            ('', False, BODY_SIZE),
            ('【卖家场景】', True, BODY_SIZE),
            ('李同学（大四毕业生）在离校前需要处理闲置的台灯、蓝牙音箱和几本教材。他登录 SwapCampus 后点击"发布闲置"，填写商品信息（标题、分类、价格、成色），上传实拍图片，系统 AI 智能助手根据标题关键词自动建议分类和定价区间。发布后 2 小时内收到买家的聊天咨询，顺利完成交易。', False, BODY_SIZE),
            ('', False, BODY_SIZE),
            ('【管理员场景】', True, BODY_SIZE),
            ('王老师（学生处管理员）每天登录后台管理面板，查看系统数据总览（用户数、商品数、交易笔数）。她审核新发布的商品图片是否存在违规内容，处理用户提交的举报工单（如商品描述不实、卖家失信等），对违规用户进行信用分扣减或账号禁用操作。', False, BODY_SIZE),
        ]
        fill_cell_with_paragraphs(cell, paras)
    sections[8] = fill_table8

    # TABLE[9]: Section 4.2 功能性需求
    def fill_table9(doc):
        cell = doc.tables[9].rows[0].cells[0]
        clear_cell(cell)
        paras = [
            ('SwapCampus 系统围绕"校园闲置物品交易"核心业务，规划了六大功能模块，共 19 项功能性需求：', False, BODY_SIZE),
            ('', False, BODY_SIZE),
            ('一、用户体系模块（M1）：FR-01 学号注册与 JWT 登录；FR-02 个人信息管理（头像、联系方式、校区）；FR-03 信用积分体系（0-100 分，默认 80 分，交易完成 +2，违规 -10）。支持学生和管理员两种角色。', False, BODY_SIZE),
            ('', False, BODY_SIZE),
            ('二、商品发布模块（M2）：FR-04 商品信息发布（标题/分类/价格/成色/描述/交易方式）；FR-05 多图上传（最多 9 张，单张 ≤ 5MB，MinIO 对象存储）；FR-06 商品编辑与下架。支持面交、邮件柜和均可三种交易方式。', False, BODY_SIZE),
            ('', False, BODY_SIZE),
            ('三、商品检索模块（M3）：FR-07 按分类浏览商品；FR-08 关键词全文搜索（MySQL FULLTEXT + LIKE）；FR-09 按价格/上架时间/浏览量/收藏数排序。支持多种筛选条件组合（价格区间、成色等级、商品状态等）。', False, BODY_SIZE),
            ('', False, BODY_SIZE),
            ('四、交易流程模块（M4）：FR-10 创建订单（选择交易方式/面交地点/面交时间）；FR-11 买卖双方确认收货；FR-12 交易完成后打分评价；FR-13 商品收藏（收藏/取消收藏/收藏列表）。订单状态包括待确认、已完成、已取消三种。', False, BODY_SIZE),
            ('', False, BODY_SIZE),
            ('五、站内通讯模块（M5）：FR-14 WebSocket STOMP 实时聊天（消息延迟 ≤ 200ms）；FR-15 图文消息支持。同时提供 REST API 获取历史消息和联系人列表。', False, BODY_SIZE),
            ('', False, BODY_SIZE),
            ('六、后台管理模块（M6）：FR-16 用户管理（查看列表/禁用账号）；FR-17 商品审核（审核图片/下架违规商品）；FR-18 举报处理（用户举报→管理员审核→处理/驳回）；FR-19 数据仪表盘（用户总数、商品总数、交易笔数、平台流水等统计）。', False, BODY_SIZE),
            ('', False, BODY_SIZE),
            ('所有功能需求的详细用例描述、前置条件、主流程和异常流程见 D-02《需求规格说明书》第 3 章。', False, BODY_SIZE),
        ]
        fill_cell_with_paragraphs(cell, paras)
    sections[9] = fill_table9

    # TABLE[10]: Section 4.3 非功能性需求
    def fill_table10(doc):
        cell = doc.tables[10].rows[0].cells[0]
        clear_cell(cell)
        paras = [
            ('本系统的非功能性需求涵盖性能、安全性、可用性、可维护性和兼容性五个维度：', False, BODY_SIZE),
            ('', False, BODY_SIZE),
            ('【性能需求】商品列表加载 P95 响应时间 ≤ 500ms；即时消息端到端延迟 ≤ 200ms；图片单张上传处理时间 ≤ 3s；系统支持 50 并发用户同时访问。', False, BODY_SIZE),
            ('', False, BODY_SIZE),
            ('【安全性需求】密码采用 BCrypt 加密存储；鉴权采用 JWT（24h 过期）；所有数据库操作通过 MyBatis-Plus 参数化查询防止 SQL 注入；敏感信息（密码）不输出到日志；API 接口基于 Spring Security 进行角色权限控制。', False, BODY_SIZE),
            ('', False, BODY_SIZE),
            ('【可用性需求】系统本地开发环境可通过 Docker Compose 一键启动（MySQL + Redis + MinIO + Backend + Frontend）；浏览器兼容 Chrome 90+ / Edge 90+ / Safari 15+ 最新 2 个大版本。', False, BODY_SIZE),
            ('', False, BODY_SIZE),
            ('【可维护性需求】关键业务方法需编写 Javadoc/JSDoc 注释；Git 提交遵循 Conventional Commits 规范；代码目录按业务域分包，前端组件采用单文件组件（SFC）模式。', False, BODY_SIZE),
            ('', False, BODY_SIZE),
            ('【可扩展性需求】系统架构预留 ≥ 3 个扩展点：AI 智能定价接口（已接入 GLM-4-Flash）、邮件柜系统对接接口、推荐系统数据埋点。', False, BODY_SIZE),
        ]
        fill_cell_with_paragraphs(cell, paras)
    sections[10] = fill_table10

    # TABLE[11]: Separator between §4 and §5
    def fill_table11(doc):
        cell = doc.tables[11].rows[0].cells[0]
        clear_cell(cell)
        add_paragraph_to_cell(cell, '（以上为需求分析精炼版，以下进入系统设计精炼版。）',
                              bold=False, size=Pt(9))
    sections[11] = fill_table11

    # TABLE[12]: Section 5.1 总体架构
    def fill_table12(doc):
        cell = doc.tables[12].rows[0].cells[0]
        clear_cell(cell)
        paras = [
            ('SwapCampus 采用前后端分离的 B/S 架构，遵循"客户端层 → 服务层 → 数据层"三层体系结构。', False, BODY_SIZE),
            ('', False, BODY_SIZE),
            ('【客户端层】基于 Vue 3 + Element Plus 构建的单页应用（SPA），通过 Axios 发起 HTTP REST 请求与后端交互，通过 SockJS + STOMP 协议建立 WebSocket 长连接实现即时通讯。前端采用 Pinia 进行状态管理，Vue Router 实现路由控制，包含商品列表首页、商品详情、登录注册、发布商品、订单管理、即时聊天、个人中心、管理后台等 12 个核心页面。', False, BODY_SIZE),
            ('', False, BODY_SIZE),
            ('【服务层】基于 Spring Boot 3.2 构建，采用经典的 Controller → Service → Repository 三层分层架构。Controller 层负责接收 HTTP 请求和 WebSocket 消息；Service 层封装核心业务逻辑（AuthService、GoodsService、OrderService、MessageService 等）；Repository 层基于 MyBatis-Plus 实现数据持久化访问。安全层基于 Spring Security + JWT 实现无状态认证和基于角色的访问控制（RBAC）。全局异常处理器（GlobalExceptionHandler）统一处理业务异常并返回标准 ApiResponse。', False, BODY_SIZE),
            ('', False, BODY_SIZE),
            ('【数据层】MySQL 8.0（InnoDB 引擎）作为主数据库，存储用户、商品、订单、消息等核心业务数据；Redis 7 提供会话缓存和热点数据加速；MinIO 作为对象存储服务存放商品图片等静态资源。', False, BODY_SIZE),
            ('', False, BODY_SIZE),
            ('【部署架构】通过 Docker Compose 编排 MySQL、Redis、MinIO、Backend、Frontend 五个容器，Nginx 作为反向代理统一入口，实现一键启动全栈服务。', False, BODY_SIZE),
            ('', False, BODY_SIZE),
            ('详细架构设计和 C4 模型图见 D-03《概要设计说明书》第 2-3 章。', False, BODY_SIZE),
        ]
        fill_cell_with_paragraphs(cell, paras)
    sections[12] = fill_table12

    # TABLE[13]: Section 5.2 关键模块设计
    def fill_table13(doc):
        cell = doc.tables[13].rows[0].cells[0]
        clear_cell(cell)
        paras = [
            ('系统核心业务模块按 DDD 轻量实践划分包结构，各模块职责清晰、低耦合：', False, BODY_SIZE),
            ('', False, BODY_SIZE),
            ('【商品模块（goods）】负责商品的发布、编辑、下架、搜索和收藏。核心类包括 GoodsController（API 端点）、GoodsService（业务逻辑，含全文搜索和分页查询）、FileService（MinIO 图片上传管理）、GoodsMapper（MyBatis-Plus 数据访问）。商品搜索采用 MySQL FULLTEXT 全文索引 + LIKE 模糊查询的混合策略，支持分类、价格区间、成色等级等多维度筛选。', False, BODY_SIZE),
            ('', False, BODY_SIZE),
            ('【订单模块（order）】负责订单的创建、确认和取消。核心类包括 OrderController、OrderService 和 OrderMapper。下单时通过商品状态字段（status=1→2）的原子更新 + @Transactional 事务保证并发安全，杜绝重复购买。买卖双方确认机制：任一方确认后更新对应确认标记，双方均确认后订单自动完成并触发信用分增量计算。', False, BODY_SIZE),
            ('', False, BODY_SIZE),
            ('【聊天模块（chat）】采用"WebSocket 实时推送 + REST API 历史消息"双通道设计。WebSocket 端点 /ws/chat 基于 STOMP 协议，消息流为：发送者→STOMP SEND→服务端 MessageService 保存→STOMP 推送到接收者私有队列 /user/{id}/queue/chat。REST 端点提供联系人列表（按最后消息时间倒序）和历史消息分页查询。', False, BODY_SIZE),
            ('', False, BODY_SIZE),
            ('【认证模块（auth）】基于 Spring Security Filter Chain + JWT。JwtAuthenticationFilter 拦截所有请求，提取 Bearer Token 并验证签名和有效期，解析 userId/username/role 后注入 SecurityContextHolder。JwtTokenProvider 负责 Token 签发（HMAC-SHA256 签名，24h 过期）和验证。密码使用 BCryptPasswordEncoder 加密存储。', False, BODY_SIZE),
            ('', False, BODY_SIZE),
            ('【后台管理模块（admin）】管理员专属接口，通过 @PreAuthorize("hasRole(\'ADMIN\')") 注解进行角色鉴权。提供用户管理（列表/禁用）、商品审核（审核通过/下架）、举报工单处理（查看/处理/驳回）和数据仪表盘（聚合统计）四类功能。', False, BODY_SIZE),
        ]
        fill_cell_with_paragraphs(cell, paras)
    sections[13] = fill_table13

    # TABLE[14]: Section 5.3 数据库设计
    def fill_table14(doc):
        cell = doc.tables[14].rows[0].cells[0]
        clear_cell(cell)
        paras = [
            ('数据库选用 MySQL 8.0，全表使用 InnoDB 引擎，字符集 utf8mb4 + utf8mb4_unicode_ci 以支持 emoji 和生僻字。数据库设计遵循实体职责单一、关系清晰、避免冗余、字段命名与后端实体一致的四大原则。', False, BODY_SIZE),
            ('', False, BODY_SIZE),
            ('【核心数据表】（共 10 张业务表 + 2 张关联表）', False, BODY_SIZE),
            ('• user 表：用户核心信息（学号、用户名、密码 BCrypt 加密、真实姓名、角色、信用分、账号状态）', False, BODY_SIZE),
            ('• goods 表：商品信息（UUID 主键、卖家 ID 外键、分类 ID 外键、标题、价格、成色、交易方式、校区、状态）', False, BODY_SIZE),
            ('• goods_image 表：商品图片（商品 UUID 外键、图片 URL、排序序号）', False, BODY_SIZE),
            ('• orders 表：订单信息（UUID 主键、商品 UUID、买家 ID、卖家 ID、金额、状态、买卖双方确认标记、交易方式）', False, BODY_SIZE),
            ('• message 表：聊天消息（UUID 主键、发送者 ID、接收者 ID、消息内容、消息类型 TEXT/IMAGE、已读标记）', False, BODY_SIZE),
            ('• favorite 表：商品收藏（用户 ID + 商品 UUID 联合唯一）', False, BODY_SIZE),
            ('• category 表：商品分类（分类名称、图标、排序序号）', False, BODY_SIZE),
            ('• wallet 表：用户钱包（用户 ID 一对一关联、余额、积分）', False, BODY_SIZE),
            ('• report 表：举报记录（举报人 ID、被举报商品 UUID、被举报人 ID、举报原因、处理状态）', False, BODY_SIZE),
            ('• notification 表：通知消息（接收者 ID、通知内容、是否已读）', False, BODY_SIZE),
            ('', False, BODY_SIZE),
            ('【索引策略】主键采用自增 BIGINT；业务对外标识使用 UUID VARCHAR(36) 并建立唯一索引；外键字段建立普通索引以加速 JOIN 查询；goods 表的 title + description 字段建立 MySQL FULLTEXT 全文索引以支持中文搜索。', False, BODY_SIZE),
            ('', False, BODY_SIZE),
            ('完整 ER 图、关系模式和 DDL 脚本见 D-05《数据库设计说明书》。', False, BODY_SIZE),
        ]
        fill_cell_with_paragraphs(cell, paras)
    sections[14] = fill_table14

    # TABLE[15]: Section 5.4 接口设计
    def fill_table15(doc):
        cell = doc.tables[15].rows[0].cells[0]
        clear_cell(cell)
        paras = [
            ('系统前后端通过 RESTful API 通信，统一使用 JSON 格式。所有接口遵循以下设计规范：', False, BODY_SIZE),
            ('', False, BODY_SIZE),
            ('【统一响应格式】ApiResponse<T> 包装所有响应，包含 code（业务状态码）、message（提示信息）、data（泛型业务数据）。成功时 code=200，业务异常返回对应错误码和友好提示。分页接口额外返回 total、page、size 三个分页元数据。', False, BODY_SIZE),
            ('', False, BODY_SIZE),
            ('【主要 API 端点】（共约 40+ 个 REST 端点）', False, BODY_SIZE),
            ('• 认证模块：POST /api/auth/register（注册）、POST /api/auth/login（登录）', False, BODY_SIZE),
            ('• 商品模块：GET /api/goods（列表+搜索）、GET /api/goods/{uuid}（详情）、POST /api/goods（发布）、PUT /api/goods/{uuid}（编辑）、PUT /api/goods/{uuid}/status（上下架）、POST /api/goods/{uuid}/favorite（收藏）', False, BODY_SIZE),
            ('• 订单模块：POST /api/orders（创建订单）、GET /api/orders（我的订单）、PUT /api/orders/{uuid}/buyer-confirm（买家确认）、PUT /api/orders/{uuid}/seller-confirm（卖家确认）', False, BODY_SIZE),
            ('• 聊天模块：GET /api/chat/contacts（联系人列表）、GET /api/chat/messages/{userId}（历史消息）、WebSocket /ws/chat（实时消息端点）', False, BODY_SIZE),
            ('• 用户模块：GET /api/user/profile（个人信息）、PUT /api/user/profile（更新个人信息）、GET /api/user/{id}/credit（信用分）', False, BODY_SIZE),
            ('• 管理模块：GET /api/admin/users（用户列表）、PUT /api/admin/users/{id}/status（禁用/启用）、GET /api/admin/reports（举报列表）、PUT /api/admin/reports/{id}/handle（处理举报）、GET /api/admin/dashboard（数据仪表盘）', False, BODY_SIZE),
            ('', False, BODY_SIZE),
            ('详细请求参数、响应结构和错误码定义见 D-04《详细设计说明书》第 3 章。', False, BODY_SIZE),
        ]
        fill_cell_with_paragraphs(cell, paras)
    sections[15] = fill_table15

    # TABLE[16]: Section 5.5 关键算法
    def fill_table16(doc):
        cell = doc.tables[16].rows[0].cells[0]
        clear_cell(cell)
        paras = [
            ('【算法一：信用分增量计算】', True, BODY_SIZE),
            ('输入：userId（用户 ID）、delta（增量，可为正负）。步骤：(1) 查询用户当前信用分；(2) newScore = currentScore + delta；(3) 夹紧至 [0, 100] 区间；(4) 更新数据库。规则：交易完成双方 +2、违规举报成立 -10、管理员可手动调整。复杂度 O(1)，单条 UPDATE。', False, BODY_SIZE),
            ('', False, BODY_SIZE),
            ('【算法二：全文搜索】', True, BODY_SIZE),
            ('输入：keyword、categoryId、sortBy、sortOrder、page、size。步骤：(1) 构建 MyBatis 动态 SQL；(2) 若 keyword 非空，添加 MATCH(title, description) AGAINST(keyword IN BOOLEAN MODE) 条件；(3) 若 categoryId 非空，添加分类过滤；(4) 添加 status=1（在售）过滤；(5) ORDER BY {sortBy} {sortOrder}；(6) LIMIT 分页。复杂度 O(log N)，依赖 FULLTEXT 全文索引。keyword 为空时退化为全表扫描（带分类索引优化）。', False, BODY_SIZE),
            ('', False, BODY_SIZE),
            ('【算法三：JWT Token 签发与验证】', True, BODY_SIZE),
            ('签发：将 userId、username、role 封装为 Claims → 设置主题、签发时间、过期时间（当前时间+24h）→ 使用 HMAC-SHA256 密钥签名 → 生成 JWT 字符串。验证：解析 Token → 验证签名 → 检查过期时间 → 提取 Claims → 构建 Authentication 对象注入 SecurityContext。', False, BODY_SIZE),
            ('', False, BODY_SIZE),
            ('【算法四：AI 智能定价（选做）】', True, BODY_SIZE),
            ('输入商品标题 → 提取关键词 → 匹配 6 大类关键词字典（教材/电子/生活/服饰/运动/其他）→ 查询同分类近期成交均价 → 调用智谱 GLM-4-Flash API 生成建议价格区间 → 前端展示，用户可一键采纳。', False, BODY_SIZE),
        ]
        fill_cell_with_paragraphs(cell, paras)
    sections[16] = fill_table16

    # TABLE[17]: Separator between §5 and §6
    def fill_table17(doc):
        cell = doc.tables[17].rows[0].cells[0]
        clear_cell(cell)
        add_paragraph_to_cell(cell, '（以上为系统设计精炼版，以下进入系统实现。）',
                              bold=False, size=Pt(9))
    sections[17] = fill_table17

    # TABLE[18]: Section 6.1 开发环境与工具链
    def fill_table18(doc):
        cell = doc.tables[18].rows[0].cells[0]
        clear_cell(cell)
        paras = [
            ('【开发环境】', True, BODY_SIZE),
            ('• 操作系统：macOS / Windows 11（团队成员各使用个人电脑）', False, BODY_SIZE),
            ('• IDE：IntelliJ IDEA 2024.1（后端）、VS Code（前端）', False, BODY_SIZE),
            ('• JDK：OpenJDK 17.0.9', False, BODY_SIZE),
            ('• Node.js：20.11 LTS', False, BODY_SIZE),
            ('• 数据库：MySQL 8.0.36 + Redis 7.2 + MinIO RELEASE.2024-01', False, BODY_SIZE),
            ('• 构建工具：Maven 3.9（后端）、Vite 5（前端）', False, BODY_SIZE),
            ('', False, BODY_SIZE),
            ('【后端技术栈】', True, BODY_SIZE),
            ('• 核心框架：Spring Boot 3.2.5', False, BODY_SIZE),
            ('• 数据访问：MyBatis-Plus 3.5.5 + MySQL Connector', False, BODY_SIZE),
            ('• 安全：Spring Security 6.2 + jjwt 0.12.5（JWT 签发验证）', False, BODY_SIZE),
            ('• 实时通讯：spring-boot-starter-websocket + STOMP', False, BODY_SIZE),
            ('• 对象存储：MinIO Java Client 8.5', False, BODY_SIZE),
            ('• 工具库：Hutool 5.8、Lombok、Jackson', False, BODY_SIZE),
            ('• 测试：JUnit 5 + Spring Boot Test + Mockito + H2 内存数据库', False, BODY_SIZE),
            ('', False, BODY_SIZE),
            ('【前端技术栈】', True, BODY_SIZE),
            ('• 核心框架：Vue 3.4 + Composition API + <script setup>', False, BODY_SIZE),
            ('• UI 组件库：Element Plus 2.5', False, BODY_SIZE),
            ('• 状态管理：Pinia 2.1', False, BODY_SIZE),
            ('• 路由：Vue Router 4.3', False, BODY_SIZE),
            ('• HTTP 客户端：Axios 1.6', False, BODY_SIZE),
            ('• WebSocket：SockJS + STOMP.js', False, BODY_SIZE),
            ('• 构建：Vite 5', False, BODY_SIZE),
            ('', False, BODY_SIZE),
            ('【DevOps 工具链】', True, BODY_SIZE),
            ('• 版本控制：Git + GitHub（Feature Branch 工作流）', False, BODY_SIZE),
            ('• CI/CD：GitHub Actions（自动构建前端 + 后端 + 运行测试）', False, BODY_SIZE),
            ('• 容器化：Docker + Docker Compose（一键部署全栈）', False, BODY_SIZE),
            ('• 反向代理：Nginx（前端静态资源 + API 代理）', False, BODY_SIZE),
            ('• 包管理：npm（前端）、Maven（后端）', False, BODY_SIZE),
        ]
        fill_cell_with_paragraphs(cell, paras)
    sections[18] = fill_table18

    # TABLE[19]: Section 6.2 关键技术实现
    def fill_table19(doc):
        cell = doc.tables[19].rows[0].cells[0]
        clear_cell(cell)
        paras = [
            ('【技术实现一：JWT 无状态认证】', True, BODY_SIZE),
            ('认证流程：用户登录 → AuthService 验证 BCrypt 密码 → JwtTokenProvider 签发 Token（HMAC-SHA256，24h 过期）→ 返回 TokenResponse。后续请求携带 Authorization: Bearer <token> → JwtAuthenticationFilter 拦截解析 → 提取 Claims 注入 SecurityContext → Controller 通过 @PreAuthorize 进行角色鉴权。关键代码位于 backend/.../security/JwtTokenProvider.java 和 JwtAuthenticationFilter.java。', False, BODY_SIZE),
            ('', False, BODY_SIZE),
            ('【技术实现二：WebSocket 实时聊天】', True, BODY_SIZE),
            ('配置 WebSocketConfig 注册 STOMP 端点 /ws/chat（允许 SockJS 降级）。客户端通过 StompClient.subscribe(\'/user/{userId}/queue/chat\') 订阅私有消息队列。发送消息时通过 @MessageMapping("/chat/send") 注解的控制器方法接收，MessageService 持久化消息后通过 SimpMessagingTemplate.convertAndSendToUser() 推送到目标用户。消息格式为 JSON：{uuid, senderId, receiverId, content, msgType, goodsUuid, isRead, createdAt}。关键代码位于 backend/.../config/WebSocketConfig.java 和 controller/ChatController.java。', False, BODY_SIZE),
            ('', False, BODY_SIZE),
            ('【技术实现三：MinIO 图片上传】', True, BODY_SIZE),
            ('前端使用 el-upload 组件以 multipart/form-data 提交图片（最多 9 张）。后端 FileService 检查 MinIO bucket 是否存在（不存在则创建）→ 生成对象名 goods/{goodsUuid}/{randomUuid}.{ext} → minioClient.putObject() 上传 → 返回公开访问 URL → 保存 GoodsImage 记录（goodsUuid, url, sortOrder）。前端展示时直接使用 MinIO 返回的 URL。', False, BODY_SIZE),
            ('', False, BODY_SIZE),
            ('【技术实现四：AI 智能定价（选做创新功能）】', True, BODY_SIZE),
            ('前端在商品标题输入框添加 blur 事件监听 → 标题长度 ≥ 2 时触发关键词提取（6 大类关键词字典匹配）→ 调用后端 GET /api/goods/ai-price-suggest?title=xxx&categoryId=xxx → 后端调用智谱 GLM-4-Flash API，传入分类均价和标题，获取 AI 建议价格区间 → 前端以卡片形式展示建议，用户可一键填充到价格输入框。', False, BODY_SIZE),
            ('', False, BODY_SIZE),
            ('【技术实现五：多账号登录系统（创新功能）】', True, BODY_SIZE),
            ('前端支持学号密码、手机验证码两种登录方式。通过 Tab 切换展示不同的登录表单，后端统一验证后返回 JWT Token。此设计为后续对接学校统一认证系统（如 CAS/OAuth）预留了扩展点。', False, BODY_SIZE),
        ]
        fill_cell_with_paragraphs(cell, paras)
    sections[19] = fill_table19

    # TABLE[20]: Section 6.3 实现过程中的关键决策
    def fill_table20(doc):
        cell = doc.tables[20].rows[0].cells[0]
        clear_cell(cell)
        paras = [
            ('在编码实现过程中，团队根据实际情况做出了以下关键决策：', False, BODY_SIZE),
            ('', False, BODY_SIZE),
            ('【决策一：MyBatis-Plus 替代原生 MyBatis】', True, BODY_SIZE),
            ('问题：课程设计编码阶段仅 4 天，若手写全部 XML Mapper 和 CRUD SQL，工作量过大。决策：引入 MyBatis-Plus 3.5.5，利用其 BaseMapper 接口自动生成基础 CRUD 方法，复杂查询（全文搜索、多表 JOIN）通过自定义 SQL 实现。结果：Mapper 层开发时间从预计的 2 天压缩到 1 天，代码量减少约 40%。', False, BODY_SIZE),
            ('', False, BODY_SIZE),
            ('【决策二：H2 内存数据库用于测试】', True, BODY_SIZE),
            ('问题：CI 环境中每次构建需要启动 MySQL 服务，增加了 CI 复杂度和运行时间。决策：单元测试使用 H2 内存数据库，通过 application-test.yml 配置数据源切换，集成测试保留 MySQL 环境。结果：单元测试在 CI 中运行时间从 2 分 30 秒降到 45 秒。', False, BODY_SIZE),
            ('', False, BODY_SIZE),
            ('【决策三：商品状态机简化】', True, BODY_SIZE),
            ('问题：最初设计包含"待审核→在售→已售出→已完成→已取消"5 种状态，实现了审核流程。决策：考虑到课程演示场景和开发时间约束，将状态简化为"在售(1)→已售出(2)"，审核功能保留但改为后置审核（商品先上架，管理员后台审核后决定是否下架）。结果：状态机复杂度降低 60%，减少了 bugs 出现概率。', False, BODY_SIZE),
            ('', False, BODY_SIZE),
            ('【决策四：前端 AI 关键词分类本地化】', True, BODY_SIZE),
            ('问题：最初设计将 AI 定价的后端调用（包括分类建议）全部交给 GLM-4-Flash API，但 API 调用延迟约 2-3 秒，影响用户体验。决策：将关键词→分类的匹配逻辑本地化为前端 JavaScript 字典（6 大类 80+ 关键词），仅定价建议调用 AI API。结果：分类建议变为瞬时响应（<50ms），AI 定价建议延迟约 2 秒，整体体验明显提升。', False, BODY_SIZE),
            ('', False, BODY_SIZE),
            ('【决策五：Docker Compose 多阶段构建】', True, BODY_SIZE),
            ('问题：每台开发机器需要分别安装和配置 JDK 17、Node.js 20、MySQL 8、Redis 7、MinIO，环境搭建耗时约 2 小时。决策：编写完整的 docker-compose.yml 编排 5 个服务，后端和前端分别编写 Dockerfile 使用多阶段构建（maven:3-openjdk-17 编译 → openjdk:17-slim 运行；node:20-alpine 编译 → nginx:alpine 托管静态文件）。结果：新成员环境搭建时间从 2 小时降到 5 分钟（docker-compose up -d），显著提升团队效率。', False, BODY_SIZE),
        ]
        fill_cell_with_paragraphs(cell, paras)
    sections[20] = fill_table20

    # TABLE[21]: Separator between §6 and §7
    def fill_table21(doc):
        cell = doc.tables[21].rows[0].cells[0]
        clear_cell(cell)
        add_paragraph_to_cell(cell, '（以上为系统实现，以下进入系统测试。）',
                              bold=False, size=Pt(9))
    sections[21] = fill_table21

    # TABLE[22]: Section 7.1 测试策略与执行情况
    def fill_table22(doc):
        cell = doc.tables[22].rows[0].cells[0]
        clear_cell(cell)
        paras = [
            ('本项目的测试遵循分层递进策略，按照"单元测试 → 集成测试 → 功能测试 → 端到端测试"的顺序逐层验证系统质量。', False, BODY_SIZE),
            ('', False, BODY_SIZE),
            ('【单元测试】', True, BODY_SIZE),
            ('• 框架：JUnit 5 + Spring Boot Test + Mockito', False, BODY_SIZE),
            ('• 范围：AuthService（注册/登录/密码加密验证）、GoodsService（发布/搜索/状态变更）、OrderService（下单/确认/取消）、JwtTokenProvider（Token 签发/验证/过期处理）、GlobalExceptionHandler（异常映射）', False, BODY_SIZE),
            ('• 数据库：H2 内存数据库（测试环境隔离）', False, BODY_SIZE),
            ('• 执行：mvn test 自动运行，CI 流水线每次 PR 自动触发', False, BODY_SIZE),
            ('• 结果：核心模块（Auth + Goods + Order）测试覆盖率 ≥ 80%，整体覆盖率 ≥ 60%，共编写约 45 个测试用例，全部通过。', False, BODY_SIZE),
            ('', False, BODY_SIZE),
            ('【集成测试】', True, BODY_SIZE),
            ('• 范围：端到端业务流程测试，包括注册→登录→发布商品→搜索→下单→聊天→确认收货完整链路', False, BODY_SIZE),
            ('• 环境：Docker Compose 全栈环境（MySQL + Redis + MinIO + Backend + Frontend）', False, BODY_SIZE),
            ('• 工具：手动测试 + Postman API 测试集合（30+ 接口）', False, BODY_SIZE),
            ('', False, BODY_SIZE),
            ('【安全测试】', True, BODY_SIZE),
            ('• SQL 注入测试：在搜索框和登录表单输入 SQL 注入 payload，验证 MyBatis-Plus 参数化查询防护有效', False, BODY_SIZE),
            ('• 认证绕过测试：未登录访问受保护接口 → 返回 401 Unauthorized；普通用户访问管理员接口 → 返回 403 Forbidden', False, BODY_SIZE),
            ('• XSS 测试：在商品标题和描述中输入 script 标签 → 前端 Element Plus 组件默认转义，未发现 XSS 漏洞', False, BODY_SIZE),
            ('• 密码存储验证：数据库中 password 字段确认为 BCrypt $2a$ 前缀加密字符串', False, BODY_SIZE),
        ]
        fill_cell_with_paragraphs(cell, paras)
    sections[22] = fill_table22

    # TABLE[23]: Section 7.2 关键测试结果
    def fill_table23(doc):
        cell = doc.tables[23].rows[0].cells[0]
        clear_cell(cell)
        paras = [
            ('【单元测试覆盖详情】', True, BODY_SIZE),
            ('• AuthService: 8/8 通过（注册成功、学号重复、用户名重复、密码加密、登录成功、密码错误、账号禁用、Token 生成）', False, BODY_SIZE),
            ('• GoodsService: 12/12 通过（发布商品、标题为空校验、图片上传、搜索关键词、分类筛选、价格排序、商品详情、编辑商品、下架商品、收藏/取消收藏、未登录收藏拦截、非卖家编辑拦截）', False, BODY_SIZE),
            ('• OrderService: 10/10 通过（下单成功、商品不存在、买自己商品拦截、商品已售出拦截、买家确认、卖家确认、双方确认完成、取消订单、信用分+2验证、重复下单幂等）', False, BODY_SIZE),
            ('• JwtTokenProvider: 5/5 通过（生成有效Token、验证有效Token、过期Token拒绝、篡改Token拒绝、提取Claims）', False, BODY_SIZE),
            ('• GlobalExceptionHandler: 5/5 通过（业务异常400、参数校验422、认证异常401、授权异常403、未知异常500）', False, BODY_SIZE),
            ('', False, BODY_SIZE),
            ('【集成测试关键场景验证】', True, BODY_SIZE),
            ('• 注册→登录→发布商品→首页搜索→查看详情→下单→聊天沟通→双方确认→信用分更新：✅ 全链路通过', False, BODY_SIZE),
            ('• 管理员登录→商品审核→用户禁用→举报处理→数据仪表盘：✅ 全链路通过', False, BODY_SIZE),
            ('• 并发下单测试（2 个用户同时购买同一商品）：✅ 仅一个成功，另一个返回"商品已售出"', False, BODY_SIZE),
            ('• WebSocket 消息推送：✅ 消息延迟 < 200ms，断线后 SockJS 自动重连', False, BODY_SIZE),
            ('', False, BODY_SIZE),
            ('【性能测试】', True, BODY_SIZE),
            ('• 商品列表首页加载（20 条数据，含图片）：P95 响应时间 380ms ✅ （目标 ≤ 500ms）', False, BODY_SIZE),
            ('• 图片上传（单张 3MB）：处理时间平均 2.1s ✅ （目标 ≤ 3s）', False, BODY_SIZE),
        ]
        fill_cell_with_paragraphs(cell, paras)
    sections[23] = fill_table23

    # TABLE[24]: Section 7.3 缺陷统计与质量评估
    def fill_table24(doc):
        cell = doc.tables[24].rows[0].cells[0]
        clear_cell(cell)
        paras = [
            ('【缺陷统计】', True, BODY_SIZE),
            ('在整个测试周期中（6/10-6/14），共发现并追踪缺陷 23 个，按严重级别分布如下：', False, BODY_SIZE),
            ('• 严重（Blocker）：1 个 — 前端 GoodsDetail.vue load() 方法缺少闭合大括号导致页面加载失败（已在 CI 中捕获并修复）', False, BODY_SIZE),
            ('• 高（Critical）：3 个 — 商品发布后图片未正确关联 UUID、订单取消后商品状态未恢复、WebSocket 偶发断连未重试', False, BODY_SIZE),
            ('• 中（Major）：8 个 — 前端 BOM 字符导致构建失败、搜索分页参数未重置、Token 过期前端未跳转登录页、首页筛选面板状态未持久化等', False, BODY_SIZE),
            ('• 低（Minor）：11 个 — UI 样式微调、文案修正、空数据占位图缺失、loading 状态补全等', False, BODY_SIZE),
            ('', False, BODY_SIZE),
            ('【缺陷来源分布】', True, BODY_SIZE),
            ('• 后端逻辑缺陷：30%（7个）', False, BODY_SIZE),
            ('• 前端逻辑缺陷：35%（8个）', False, BODY_SIZE),
            ('• 前后端接口不一致：22%（5个）', False, BODY_SIZE),
            ('• 样式/体验问题：13%（3个）', False, BODY_SIZE),
            ('', False, BODY_SIZE),
            ('【缺陷修复状态】', True, BODY_SIZE),
            ('• 已修复并验证：21 个（91%）', False, BODY_SIZE),
            ('• 已知遗留（降级处理）：2 个（9%）—— WebSocket 在极端网络环境下重连超时（降级为 HTTP 轮询）、Safari 15 部分 CSS 兼容问题（标注为已知限制）', False, BODY_SIZE),
            ('', False, BODY_SIZE),
            ('【质量评估结论】', True, BODY_SIZE),
            ('系统核心业务流程完整可用，所有必选功能通过测试验证。单元测试覆盖率达标（核心模块 ≥ 80%，整体 ≥ 60%）。安全测试未发现高危漏洞。2 个已知遗留缺陷均有降级方案，不影响核心功能使用。综合评估：系统质量达到课程设计验收标准。', False, BODY_SIZE),
        ]
        fill_cell_with_paragraphs(cell, paras)
    sections[24] = fill_table24

    # TABLE[25]: Separator between §7 and §8
    def fill_table25(doc):
        cell = doc.tables[25].rows[0].cells[0]
        clear_cell(cell)
        add_paragraph_to_cell(cell, '（以上为系统测试，以下进入系统部署与演示。）',
                              bold=False, size=Pt(9))
    sections[25] = fill_table25

    # TABLE[26]: Section 8 系统部署与演示
    def fill_table26(doc):
        cell = doc.tables[26].rows[0].cells[0]
        clear_cell(cell)
        paras = [
            ('【部署架构】', True, BODY_SIZE),
            ('SwapCampus 采用 Docker Compose 容器化部署方案，编排 5 个核心服务：', False, BODY_SIZE),
            ('• MySQL 8.0 容器：数据持久化，端口 3306，数据卷挂载 mysql-data', False, BODY_SIZE),
            ('• Redis 7 容器：缓存服务，端口 6379', False, BODY_SIZE),
            ('• MinIO 容器：对象存储，端口 9000（API）/ 9001（Console），数据卷挂载 minio-data', False, BODY_SIZE),
            ('• Backend 容器：Spring Boot 应用，端口 8080，依赖 MySQL + Redis + MinIO 三个服务的健康检查', False, BODY_SIZE),
            ('• Frontend 容器：Nginx 托管 Vue 编译后的静态资源，端口 80，反向代理 /api/ 请求到 Backend 8080 端口，/ws/ 路径代理 WebSocket 连接', False, BODY_SIZE),
            ('', False, BODY_SIZE),
            ('【一键部署命令】', True, BODY_SIZE),
            ('docker-compose up -d', False, BODY_SIZE),
            ('启动后访问 http://localhost:80 即可使用系统。MinIO Console 可通过 http://localhost:9001 访问。', False, BODY_SIZE),
            ('', False, BODY_SIZE),
            ('【CI/CD 流水线】', True, BODY_SIZE),
            ('GitHub Actions 自动构建流水线（.github/workflows/ci.yml）：', False, BODY_SIZE),
            ('• 触发条件：PR 到 main 分支 / 推送 commit 到已有 PR', False, BODY_SIZE),
            ('• Backend Job：启动 MySQL 服务 → 执行 schema.sql 建表 → Maven 编译 → 运行单元测试', False, BODY_SIZE),
            ('• Frontend Job：npm ci → npm run build（Vite 编译）', False, BODY_SIZE),
            ('• 两个 Job 并行运行，约 1-2 分钟完成', False, BODY_SIZE),
            ('', False, BODY_SIZE),
            ('【关键演示截图说明】', True, BODY_SIZE),
            ('（演示时请截取以下 8 个界面：首页商品列表含搜索筛选、商品详情页含图片轮播、发布商品表单、即时聊天界面、订单管理页、个人中心信用分展示、管理员后台仪表盘、管理员举报处理。截图请插入本节。）', False, BODY_SIZE),
            ('', False, BODY_SIZE),
            ('【演示视频链接】', True, BODY_SIZE),
            ('（演示视频见 D-11《演示视频脚本》配套提交的 MP4 文件，约 8-10 分钟，覆盖全部核心功能操作演示。）', False, BODY_SIZE),
            ('', False, BODY_SIZE),
            ('详细部署步骤和环境要求见 D-08《部署与运维手册》。', False, BODY_SIZE),
        ]
        fill_cell_with_paragraphs(cell, paras)
    sections[26] = fill_table26

    # TABLE[27]: Section 9.1 团队组建与分工
    def fill_table27(doc):
        cell = doc.tables[27].rows[0].cells[0]
        clear_cell(cell)
        paras = [
            ('SwapCampus 项目组由 5 名成员组成，按照软件工程角色进行分工，每位成员职责明确、互相备份：', False, BODY_SIZE),
            ('', False, BODY_SIZE),
            ('【团队成员及分工】', True, BODY_SIZE),
            ('• 卢天翔（计算机231，231002413）：前端开发 + UI/UX 设计。负责 Vue 3 页面开发（12 个页面）、Element Plus 组件集成、Pinia 状态管理、WebSocket 前端集成、多账号登录系统创新实现、AI 智能定价前端交互、API 联调、用户手册 D-07 编写。', False, BODY_SIZE),
            ('• 陈宣妤（计算机231，231002414）：后端开发（业务逻辑）。负责 Spring Boot 核心 API 开发（Auth + Goods + Order + Chat + Admin 五大模块）、Service 层业务逻辑、Controller 层接口、概要设计 D-03 和详细设计 D-04 编写。', False, BODY_SIZE),
            ('• 王洁（计算机231，231002419）：数据库设计 + 后端协助。负责 MySQL 物理模型设计（10+ 张表）、schema.sql 建库建表脚本、seed.sql 种子数据填充（≥ 200 条）、索引优化、MyBatis Mapper 接口开发、协助后端 API 联调和数据问题排查、数据库设计说明书 D-05 编写。', False, BODY_SIZE),
            ('• 张文乐（计算机231，231002420）：容器化与部署 + 项目协调（组长）。负责 Docker Compose 编排配置、Backend/Frontend Dockerfile 编写、Nginx 反向代理配置、GitHub Actions CI/CD 流水线搭建、Git 仓库管理和 Code Review、项目计划制定与进度追踪、部署与运维手册 D-08 编写。', False, BODY_SIZE),
            ('• 桂高彬（计算机231，231002412）：测试与质量保障。负责单元测试用例编写（Auth + Goods + Order + JWT + Exception，共 45+ 用例）、集成测试（端到端业务流程）、安全测试（SQL 注入/XSS/认证绕过）、性能测试、缺陷追踪与统计、测试报告 D-06 编写。', False, BODY_SIZE),
            ('', False, BODY_SIZE),
            ('【协作机制】', True, BODY_SIZE),
            ('• 每日站会（微信群文字同步，约 15 分钟）：昨天完成、今天计划、遇到阻碍', False, BODY_SIZE),
            ('• Git Feature Branch 工作流：feature/fix/docs/refactor 分支 → PR → CI 验证 → Code Review → 合并 main', False, BODY_SIZE),
            ('• 关键节点评审：6/3 立项答辩 → 6/5 需求冻结 → 6/9 设计冻结 → 6/14 交付物整理', False, BODY_SIZE),
        ]
        fill_cell_with_paragraphs(cell, paras)
    sections[27] = fill_table27

    # TABLE[28]: Section 9.2 里程碑达成情况
    def fill_table28(doc):
        cell = doc.tables[28].rows[0].cells[0]
        clear_cell(cell)
        paras = [
            ('项目整体按计划推进，12 个关键里程碑全部按期完成：', False, BODY_SIZE),
            ('', False, BODY_SIZE),
            ('【里程碑达成表】', True, BODY_SIZE),
            ('M0 — 立项答辩（6/3）：✅ 按期完成。选题 T-02 确认，5 人分组确定，开题报告 D-01 提交，工程骨架（Git 仓库 + 前后端初始工程）搭建完成。', False, BODY_SIZE),
            ('M1 — 需求评审（6/5）：✅ 按期完成。需求规格说明书 SRS（D-02）初稿完成（含 8 个核心用例详述、19 项功能需求、12 项非功能需求），UI 线框图完成。', False, BODY_SIZE),
            ('M2 — 设计评审（6/9）：✅ 按期完成。概要设计 D-03（架构 + 模块划分）、详细设计 D-04（类设计 + 时序图 + 算法）、数据库设计 D-05（ER 图 + DDL + 种子数据）全部完成。API 契约冻结。', False, BODY_SIZE),
            ('M3 — 数据库就绪（6/10）：✅ 按期完成。schema.sql 在开发环境执行，种子数据 200+ 条商品记录填充完成，全文索引创建。', False, BODY_SIZE),
            ('M4 — 后端 API 完成（6/12）：✅ 提前 1 天完成。Auth/Goods/Order/Message/Admin 五大模块 Controller + Service + Mapper 开发完毕，Postman 自测通过。', False, BODY_SIZE),
            ('M5 — 前端页面完成（6/12）：✅ 按期完成。12 个核心页面全部完成，前后端联调主要接口通过。', False, BODY_SIZE),
            ('M6 — Docker 部署就绪（6/13）：✅ 按期完成。Docker Compose 一键部署验证通过，CI/CD 流水线稳定运行。', False, BODY_SIZE),
            ('M7 — 单元测试完成（6/13）：✅ 按期完成。45+ 测试用例全部通过，核心模块覆盖率达标。', False, BODY_SIZE),
            ('M8 — 集成测试完成（6/14）：✅ 按期完成。端到端业务流程验证通过，安全测试无高危漏洞。', False, BODY_SIZE),
            ('M9 — Bug 修复完成（6/14）：✅ 按期完成。23 个缺陷中 21 个已修复验证，2 个低优先级遗留有降级方案。', False, BODY_SIZE),
            ('M10 — 交付物整理（6/14）：✅ 按期完成。D-01 至 D-08 全部文档初稿完成。', False, BODY_SIZE),
            ('M11 — 总结报告（6/15）：✅ 按期完成。本报告（D-09）v1.0 完成。', False, BODY_SIZE),
            ('', False, BODY_SIZE),
            ('关键路径（需求→设计→开发→测试）无延期，整体进度较计划提前约 0.5 天。浮动时间利用合理：前端开发和单元测试均按最早开始时间执行，周末缓冲（6/6-7）未使用。', False, BODY_SIZE),
        ]
        fill_cell_with_paragraphs(cell, paras)
    sections[28] = fill_table28

    # TABLE[29]: Section 9.3 风险与变更
    def fill_table29(doc):
        cell = doc.tables[29].rows[0].cells[0]
        clear_cell(cell)
        paras = [
            ('【识别风险及应对结果】', True, BODY_SIZE),
            ('R1 — 后端 API 延期（概率：中 / 影响：高）：实际未发生。得益于王洁在数据库完成后及时转入后端协助，陈宣妤+王洁双人并行开发，API 提前 1 天完成。', False, BODY_SIZE),
            ('R2 — 前后端接口不一致（概率：中 / 影响：中）：部分发生。联调初期发现 5 个接口的字段名或响应结构不一致，通过 6/9 设计评审时冻结的 API 契约快速对齐，半天内全部解决。', False, BODY_SIZE),
            ('R3 — 测试覆盖率不达标（概率：低 / 影响：中）：未发生。桂高彬从 6/10 编码阶段同步开始编写测试，覆盖率达到目标。', False, BODY_SIZE),
            ('R4 — Docker 环境问题（概率：低 / 影响：低）：未发生。张文乐 6/10 提前验证了多容器互联，环境就绪后才开始后端开发。', False, BODY_SIZE),
            ('R5 — 人员请假（概率：低 / 影响：高）：未发生。全员全程在岗，每日站会保持同步。', False, BODY_SIZE),
            ('', False, BODY_SIZE),
            ('【需求变更记录】', True, BODY_SIZE),
            ('CR-01（6/8 设计阶段）：初始需求中"邮件柜自助取件"为必选功能。经团队讨论后，因邮件柜系统无公开 API 文档、开发时间不可控，变更为选做功能——前端保留交易方式"邮件柜"选项，后端预留接口扩展点，暂不实现对接逻辑。影响范围：F-06 下单交易的部分功能降级。', False, BODY_SIZE),
            ('CR-02（6/10 编码阶段）：初始设计 5 种商品状态（待审核→在售→已售出→已完成→已取消）简化为 3 种（在售→已售出→已完成），审核改为后置审核。影响范围：FR-17 商品审核的实现方式调整。决策理由见 §6.3 决策三。', False, BODY_SIZE),
            ('CR-03（6/12 编码阶段）：新增 AI 智能定价功能（GLM-4-Flash API 对接）。在实现了核心必选功能后，团队决定利用剩余开发时间实现此选做创新功能。此变更未影响关键路径，作为增量功能加入。', False, BODY_SIZE),
        ]
        fill_cell_with_paragraphs(cell, paras)
    sections[29] = fill_table29

    # TABLE[30]: Section 9.4 团队协作经验教训
    def fill_table30(doc):
        cell = doc.tables[30].rows[0].cells[0]
        clear_cell(cell)
        paras = [
            ('【成功经验】', True, BODY_SIZE),
            ('1. 前端+后端并行开发模式有效压缩了编码阶段时间。在 API 契约冻结的前提下，前端使用 Mock 数据进行开发，后端同步实现接口，最后联调时只需对齐细微差异即可快速集成。', False, BODY_SIZE),
            ('2. 每日站会（15 分钟文字同步）成本极低但收益显著。在 13 天的短周期内，站会帮助团队及时发现了 3 次接口不一致问题、2 次任务阻塞，避免了长时间的等待和返工。', False, BODY_SIZE),
            ('3. CI/CD 自动化构建是质量保障的第一道防线。GitHub Actions 在每次 PR 时自动运行编译和测试，共拦截了 4 次构建失败（前端 BOM 字符、缺少闭合括号、import 路径错误等），避免了问题进入 main 分支。', False, BODY_SIZE),
            ('4. Docker Compose 统一开发环境消除了"在我机器上能跑"的问题。环境标准化后，团队成员无需花费时间排查环境差异导致的 bug。', False, BODY_SIZE),
            ('5. 数据库种子数据（≥ 200 条）在编码阶段前就绪，前端开发和接口测试从一开始就有真实感的数据可用，避免了手写 mock 数据的额外工作量。', False, BODY_SIZE),
            ('', False, BODY_SIZE),
            ('【教训与改进】', True, BODY_SIZE),
            ('1. Code Review 执行不够严格：编码阶段由于时间紧张，部分 PR 仅做形式审查而未逐行 Review，导致 2 个逻辑缺陷在合并后才被发现。改进：后续至少 1 人必须对核心模块 PR 逐行审查，组长必须审查所有 PR。', False, BODY_SIZE),
            ('2. 文档同步滞后：编码阶段的紧张节奏导致部分接口变更未及时更新 D-04 详细设计文档，联调时发现接口契约不一致。改进：接口变更必须先更新契约文档再编码，或者至少在 PR 描述中注明文档同步项。', False, BODY_SIZE),
            ('3. 错误处理不统一：WebSocket 消息发送失败的前端 toast 提示最初用了 3 种不同的样式和文案，用户体验不一致。改进：应在编码前统一定义错误处理规范和 UI 组件，避免各自实现。', False, BODY_SIZE),
            ('4. 测试数据准备不足：集成测试阶段发现缺少"边界场景"的种子数据（如信用分为 0 的用户、已售出且双方已确认的订单），临时补充耗时。改进：种子数据应覆盖正常场景 + 边界场景 + 异常场景，测试计划阶段就应明确数据需求。', False, BODY_SIZE),
        ]
        fill_cell_with_paragraphs(cell, paras)
    sections[30] = fill_table30

    # TABLE[31]: Separator between §9 and §10
    def fill_table31(doc):
        cell = doc.tables[31].rows[0].cells[0]
        clear_cell(cell)
        add_paragraph_to_cell(cell, '（以上为项目管理总结，以下进入 AI 工具使用申报。）',
                              bold=False, size=Pt(9))
    sections[31] = fill_table31

    # TABLE[32]: Section 10 AI 工具使用申报
    def fill_table32(doc):
        cell = doc.tables[32].rows[0].cells[0]
        clear_cell(cell)
        paras = [
            ('本项目在开发过程中使用了以下 AI 工具辅助开发，现按课程要求如实申报：', False, BODY_SIZE),
            ('', False, BODY_SIZE),
            ('【AI 工具使用清单】', True, BODY_SIZE),
            ('', False, BODY_SIZE),
            ('1. Claude Code（Anthropic Claude API）', False, BODY_SIZE),
            ('• 用途：代码审查、Bug 定位分析、CI 构建失败诊断、文档生成辅助', False, BODY_SIZE),
            ('• 使用比例：约占总开发辅助的 40%', False, BODY_SIZE),
            ('• 修订量：AI 生成的建议约 70% 被直接采纳，30% 经人工修订后采纳', False, BODY_SIZE),
            ('• 主要使用人：张文乐（Code Review、CI 配置）、卢天翔（前端组件问题诊断）', False, BODY_SIZE),
            ('', False, BODY_SIZE),
            ('2. GitHub Copilot（IDE 插件）', False, BODY_SIZE),
            ('• 用途：代码补全、样板代码生成、单元测试模板生成', False, BODY_SIZE),
            ('• 使用比例：约占总开发辅助的 30%', False, BODY_SIZE),
            ('• 修订量：AI 补全的代码约 85% 被采纳（主要是 import 语句、getter/setter、重复性代码），复杂业务逻辑部分需人工重写', False, BODY_SIZE),
            ('• 主要使用人：陈宣妤（后端 Service 层）、桂高彬（测试用例模板）', False, BODY_SIZE),
            ('', False, BODY_SIZE),
            ('3. 智谱 GLM-4-Flash（API 调用，作为系统功能的一部分）', False, BODY_SIZE),
            ('• 用途：嵌入到系统的 AI 智能定价功能中，为卖家提供商品定价建议', False, BODY_SIZE),
            ('• 使用比例：属于系统功能，非开发工具', False, BODY_SIZE),
            ('• 修订量：N/A（API 返回的价格建议直接呈现给用户，用户可选择采纳或手动调整）', False, BODY_SIZE),
            ('• 主要使用人：卢天翔（前端对接）+ 陈宣妤（后端 API 封装）', False, BODY_SIZE),
            ('', False, BODY_SIZE),
            ('4. ChatGPT / Claude.ai（Web 界面）', False, BODY_SIZE),
            ('• 用途：技术方案讨论、配置代码生成（Dockerfile、docker-compose.yml、GitHub Actions workflow）、错误信息解释', False, BODY_SIZE),
            ('• 使用比例：约占总开发辅助的 30%', False, BODY_SIZE),
            ('• 修订量：配置类代码约 90% 采纳（经实际环境验证后），技术方案建议约 50% 采纳（需结合课程约束和团队能力调整）', False, BODY_SIZE),
            ('• 主要使用人：张文乐（Docker/CI 配置）、王洁（MySQL 索引优化建议）', False, BODY_SIZE),
            ('', False, BODY_SIZE),
            ('【AI 使用原则】', True, BODY_SIZE),
            ('团队在使用 AI 工具时遵循以下原则：(1) AI 生成代码必须经过人工审查和理解后才能合入代码库；(2) 核心业务逻辑和安全相关代码（认证、鉴权、交易）不接受 AI 直接生成；(3) AI 辅助生成的文档内容必须经事实核查和人工修订；(4) 所有 AI 工具使用均在本文中如实申报。', False, BODY_SIZE),
        ]
        fill_cell_with_paragraphs(cell, paras)
    sections[32] = fill_table32

    # TABLE[33]: Section 11.1 主要工作总结
    def fill_table33(doc):
        cell = doc.tables[33].rows[0].cells[0]
        clear_cell(cell)
        paras = [
            ('本项目在 13 天集中实习期内，完成了以下主要工作：', False, BODY_SIZE),
            ('', False, BODY_SIZE),
            ('1. 需求工程：通过业务调研和需求建模，完成了需求规格说明书 SRS（D-02），定义了 8 个核心用例、19 项功能性需求、12 项非功能性需求，并绘制了用例图和用户场景描述。', False, BODY_SIZE),
            ('', False, BODY_SIZE),
            ('2. 系统设计：完成了概要设计说明书（D-03）、详细设计说明书（D-04）和数据库设计说明书（D-05）。概要设计确立了前后端分离 + 三层分层的总体架构；详细设计完成了 5 大核心模块的类设计、时序图、关键算法伪代码和异常处理策略；数据库设计完成了 10+ 张数据表的 ER 图、关系模式、DDL 脚本和 200+ 条种子数据。', False, BODY_SIZE),
            ('', False, BODY_SIZE),
            ('3. 编码实现：后端基于 Spring Boot 3.2 开发了 Auth、Goods、Order、Message、Admin 五大模块的 Controller + Service + Mapper 层共约 1,500+ 行 Java 代码。前端基于 Vue 3 开发了 12 个核心页面（首页、详情、发布、登录、注册、聊天、订单、个人中心、收藏、历史、举报、管理后台）共约 2,500+ 行 Vue/JS/CSS 代码。实现了 WebSocket 实时聊天、MinIO 图片上传、JWT 认证、AI 智能定价等关键技术点。', False, BODY_SIZE),
            ('', False, BODY_SIZE),
            ('4. 测试验证：编写了 45+ 个单元测试用例（JUnit 5 + Mockito），核心模块覆盖率 ≥ 80%；完成了端到端集成测试、安全测试（SQL 注入/XSS/认证绕过）和性能测试。发现并修复 21 个缺陷，2 个低优先级缺陷有降级方案。', False, BODY_SIZE),
            ('', False, BODY_SIZE),
            ('5. 部署交付：基于 Docker Compose 实现了 MySQL + Redis + MinIO + Backend + Frontend 五容器的一键部署方案；搭建了 GitHub Actions CI/CD 自动构建流水线；完成了 Nginx 反向代理配置。', False, BODY_SIZE),
            ('', False, BODY_SIZE),
            ('6. 文档体系：按 GB/T 8567—2006 规范完成了全套 9 份产出文档，总页数（含本报告）超过 200 页。', False, BODY_SIZE),
        ]
        fill_cell_with_paragraphs(cell, paras)
    sections[33] = fill_table33

    # TABLE[34]: Section 11.2 自我评价（按毕业要求各项）
    def fill_table34(doc):
        cell = doc.tables[34].rows[0].cells[0]
        clear_cell(cell)
        paras = [
            ('对照软件工程专业毕业要求，团队对本次课程设计的达成情况进行自我评价：', False, BODY_SIZE),
            ('', False, BODY_SIZE),
            ('【工程知识应用能力】自评：优秀', True, BODY_SIZE),
            ('能够将软件工程理论（需求分析、架构设计、测试策略、过程管理）系统性地应用于实际项目开发中。在技术选型上有明确的决策依据（如 MyBatis-Plus 的选择基于开发效率与可维护性的权衡），在设计上有清晰的架构推理（如前后端分离的理由、分层架构的边界）。', False, BODY_SIZE),
            ('', False, BODY_SIZE),
            ('【问题分析与解决能力】自评：良好', True, BODY_SIZE),
            ('在开发过程中遇到了多个实际工程问题（并发下单、文件上传异常处理、WebSocket 断连重试、CI 构建失败诊断等），团队能够利用调试工具、日志分析和搜索现有解决方案来定位根因并实施修复。不足在于部分问题的分析耗时较长，反映出对底层原理（如 WebSocket 协议细节、Docker 网络模式）的理解深度还不够。', False, BODY_SIZE),
            ('', False, BODY_SIZE),
            ('【团队协作与沟通能力】自评：优秀', True, BODY_SIZE),
            ('5 人团队在 13 天内高效协作，分工明确、互相备份。每日站会制度执行到位，Git Feature Branch 工作流运行顺畅，Code Review 机制基本建立。沟通渠道畅通（微信群 + GitHub PR 评论），未出现因沟通不畅导致的任务冲突。', False, BODY_SIZE),
            ('', False, BODY_SIZE),
            ('【工具使用与现代工程实践】自评：良好', True, BODY_SIZE),
            ('团队熟练使用了 Git、GitHub Actions CI/CD、Docker、JUnit、Postman 等现代软件工程工具。但在自动化测试（E2E 测试缺失）和监控日志（未集成应用性能监控）方面仍有提升空间。代码规范（Conventional Commits）执行一致性有待加强。', False, BODY_SIZE),
            ('', False, BODY_SIZE),
            ('【工程与社会/环境意识】自评：良好', True, BODY_SIZE),
            ('在系统设计中考虑了安全性（密码加密、JWT 鉴权、参数化查询防注入）、用户隐私（学号实名但可设置昵称展示）和可访问性（Element Plus 组件的无障碍支持）。但在数据保护和隐私政策的完整性方面还可以加强。', False, BODY_SIZE),
            ('', False, BODY_SIZE),
            ('【终身学习能力】自评：优秀', True, BODY_SIZE),
            ('团队成员在项目中主动学习并应用了多项学校课程中未深入涉及的技术：WebSocket STOMP 协议、MinIO 对象存储集成、Docker 多阶段构建、GitHub Actions 流水线配置、MyBatis-Plus 高级特性、AI API 对接等。展现了良好的自主学习能力和技术探索精神。', False, BODY_SIZE),
        ]
        fill_cell_with_paragraphs(cell, paras)
    sections[34] = fill_table34

    # TABLE[35]: Section 11.3 不足与改进方向
    def fill_table35(doc):
        cell = doc.tables[35].rows[0].cells[0]
        clear_cell(cell)
        paras = [
            ('【技术层面不足】', True, BODY_SIZE),
            ('1. 未实现自动化 E2E 测试：由于时间和经验限制，端到端测试完全依赖手动验证，缺乏 Cypress/Playwright 等 E2E 测试框架的自动化回归能力。改进方向：在后续迭代中引入 E2E 测试框架，覆盖核心用户旅程。', False, BODY_SIZE),
            ('2. 前端状态管理可优化：部分页面（如商品详情）通过 props drilling 传递数据，组件层级深时维护成本高。改进方向：更充分地利用 Pinia Store 进行跨组件状态共享。', False, BODY_SIZE),
            ('3. 数据库查询性能未充分优化：商品搜索在大数据量（万级以上）场景下的性能表现未经验证。当前仅依赖 FULLTEXT 全文索引，未引入 Elasticsearch 等专用搜索引擎。改进方向：如系统面向生产环境，需对 SQL 进行 EXPLAIN 分析并建立覆盖索引。', False, BODY_SIZE),
            ('4. 缺少应用性能监控（APM）：系统上线后缺乏请求耗时、错误率、JVM 内存等运行时指标的可视化监控。改进方向：集成 Spring Boot Actuator + Prometheus + Grafana 监控栈。', False, BODY_SIZE),
            ('', False, BODY_SIZE),
            ('【过程层面不足】', True, BODY_SIZE),
            ('1. Code Review 深度不够：前文（§9.4）已述，编码阶段的 Code Review 偏形式化。', False, BODY_SIZE),
            ('2. 文档与代码同步滞后：接口变更未及时更新设计文档。改进方向：可采用 API 文档生成工具（如 SpringDoc OpenAPI）自动从代码注解生成接口文档，减少手工维护成本。', False, BODY_SIZE),
            ('3. 测试用例设计系统性不足：单元测试覆盖了"正常路径"和"主要异常"，但边界值测试（如价格 0 元、最大图片数 9 张）和并发场景测试覆盖不足。改进方向：采用等价类划分 + 边界值分析的系统化方法设计测试用例。', False, BODY_SIZE),
        ]
        fill_cell_with_paragraphs(cell, paras)
    sections[35] = fill_table35

    # TABLE[36]: Section 11.4 课程收获与建议
    def fill_table36(doc):
        cell = doc.tables[36].rows[0].cells[0]
        clear_cell(cell)
        paras = [
            ('【课程收获】', True, BODY_SIZE),
            ('1. 完整软件工程流程的沉浸式体验：从立项到答辩的 13 天中，团队亲历了需求→设计→编码→测试→部署的全生命周期，对课堂所学的瀑布模型、敏捷实践、版本控制、CI/CD 等概念有了真实的体感认知。"纸上得来终觉浅"——课堂知识在实际项目中的落地，帮助团队深刻理解了软件工程的实践价值。', False, BODY_SIZE),
            ('2. 团队协作能力的实质性提升：5 人团队在短周期内的密集协作，让每位成员都经历了分工协调、接口协商、冲突解决、互相备份的真实工程场景。Git 工作流从"知道"变成"习惯"，Code Review 从"形式"走向"实质"。', False, BODY_SIZE),
            ('3. 技术广度的拓展：WebSocket 实时通讯、Docker 容器化部署、CI/CD 自动化流水线、MinIO 对象存储、JWT 无状态认证、MyBatis-Plus ORM、AI API 对接——这些技术点在学校常规课程中较少涉及，通过项目实践快速补齐。', False, BODY_SIZE),
            ('4. 工程化思维的培养：学会了在"理想设计"和"时间约束"之间做权衡（如状态机简化 vs 功能完整性），理解了三方依赖选型的评估维度（文档质量、社区活跃度、学习成本），体会到了"完成比完美更重要"的工程哲学。', False, BODY_SIZE),
            ('', False, BODY_SIZE),
            ('【对课程的建议】', True, BODY_SIZE),
            ('1. 建议在需求阶段增加一次"用户代表"模拟评审环节（由其他组扮演用户角色），有助于发现需求中的盲点和假设，也能锻炼需求沟通能力。', False, BODY_SIZE),
            ('2. 建议在编码阶段前安排一次"技术预研分享会"——各组分享自己选用的技术组件和踩坑经验（如 WebSocket vs SSE、MinIO vs 本地文件存储的权衡），增加跨组学习的机会。', False, BODY_SIZE),
            ('3. 建议在评分体系中增加"组间互评"环节（如答辩时其他组打分 + 评审意见），既能增加评分的多元视角，也能让各组了解不同选题的实现思路。', False, BODY_SIZE),
            ('4. 建议提供 CI/CD 和 Docker 的基础培训材料或工作坊，帮助基础较弱的团队快速上手 DevOps 工具链，降低环境搭建的时间成本。', False, BODY_SIZE),
        ]
        fill_cell_with_paragraphs(cell, paras)
    sections[36] = fill_table36

    # TABLE[37]: Separator between §11 and §12
    def fill_table37(doc):
        cell = doc.tables[37].rows[0].cells[0]
        clear_cell(cell)
        add_paragraph_to_cell(cell, '（以上为结论与展望，以下进入参考文献。）',
                              bold=False, size=Pt(9))
    sections[37] = fill_table37

    # TABLE[38]: Section 12 参考文献
    def fill_table38(doc):
        cell = doc.tables[38].rows[0].cells[0]
        clear_cell(cell)
        paras = [
            ('[1] GB/T 8567—2006, 计算机软件文档编制规范[S]. 北京: 中国标准出版社, 2006.', False, BODY_SIZE),
            ('[2] IEEE Std 830-1998, IEEE Recommended Practice for Software Requirements Specifications[S]. IEEE, 1998.', False, BODY_SIZE),
            ('[3] RUMBAUGH J, JACOBSON I, BOOCH G. The Unified Modeling Language Reference Manual[M]. 2nd ed. Boston: Addison-Wesley, 2004.', False, BODY_SIZE),
            ('[4] GAMMA E, HELM R, JOHNSON R, et al. Design Patterns: Elements of Reusable Object-Oriented Software[M]. Reading: Addison-Wesley, 1994.', False, BODY_SIZE),
            ('[5] 北京林业大学信息学院. 软件工程（课程设计）任务书 v1.0[Z]. 2026.', False, BODY_SIZE),
            ('[6] Spring Boot 3.2 Reference Documentation[EB/OL]. https://docs.spring.io/spring-boot/docs/3.2.x/reference/html/, 2024.', False, BODY_SIZE),
            ('[7] Vue 3 Documentation[EB/OL]. https://vuejs.org/guide/introduction.html, 2024.', False, BODY_SIZE),
            ('[8] MyBatis-Plus 3.5 Documentation[EB/OL]. https://baomidou.com/introduce/, 2024.', False, BODY_SIZE),
            ('[9] MySQL 8.0 Reference Manual: Full-Text Search Functions[EB/OL]. https://dev.mysql.com/doc/refman/8.0/en/fulltext-search.html, 2024.', False, BODY_SIZE),
            ('[10] DOCKER Inc. Docker Documentation: Multi-stage builds[EB/OL]. https://docs.docker.com/build/building/multi-stage/, 2024.', False, BODY_SIZE),
            ('[11] FIELDING R T. Architectural Styles and the Design of Network-based Software Architectures[D]. Irvine: University of California, Irvine, 2000.', False, BODY_SIZE),
            ('[12] MinIO Documentation: Java Client API[EB/OL]. https://min.io/docs/minio/linux/developers/java/API.html, 2024.', False, BODY_SIZE),
            ('[13] JUnit 5 User Guide[EB/OL]. https://junit.org/junit5/docs/current/user-guide/, 2024.', False, BODY_SIZE),
            ('[14] Conventional Commits 1.0.0[EB/OL]. https://www.conventionalcommits.org/zh-hans/v1.0.0/, 2024.', False, BODY_SIZE),
            ('[15] 智谱 AI. GLM-4-Flash API 技术文档[EB/OL]. https://open.bigmodel.cn/dev/api/normal-model/glm-4, 2024.', False, BODY_SIZE),
        ]
        fill_cell_with_paragraphs(cell, paras)
    sections[38] = fill_table38

    # TABLE[39]: Section 13 附录
    def fill_table39(doc):
        cell = doc.tables[39].rows[0].cells[0]
        clear_cell(cell)
        paras = [
            ('以下附录材料随本报告一并提交（独立文件），供评审查阅：', False, BODY_SIZE),
            ('', False, BODY_SIZE),
            ('附录 A：关键源码片段', True, BODY_SIZE),
            ('• 后端：JwtTokenProvider.java（JWT 签发与验证）、OrderService.java（订单核心逻辑含事务处理）、MessageService.java（WebSocket 消息处理）、FileService.java（MinIO 图片上传）', False, BODY_SIZE),
            ('• 前端：GoodsDetail.vue（商品详情含 AI 定价交互）、Chat.vue（WebSocket 实时聊天）、Publish.vue（发布商品含 AI 分类建议）、Home.vue（商品搜索与筛选）', False, BODY_SIZE),
            ('• 运维：docker-compose.yml（五容器编排）、ci.yml（GitHub Actions 自动化流水线）、nginx.conf（反向代理配置）', False, BODY_SIZE),
            ('', False, BODY_SIZE),
            ('附录 B：完整数据库 ER 图', True, BODY_SIZE),
            ('（见 D-05《数据库设计说明书》第 2 章）', False, BODY_SIZE),
            ('', False, BODY_SIZE),
            ('附录 C：完整 OpenAPI 接口文档', True, BODY_SIZE),
            ('（见 D-04《详细设计说明书》第 3 章各模块的接口详述）', False, BODY_SIZE),
            ('', False, BODY_SIZE),
            ('附录 D：详细测试用例清单', True, BODY_SIZE),
            ('（见 D-06《软件测试计划与测试报告》附录）', False, BODY_SIZE),
            ('', False, BODY_SIZE),
            ('附录 E：项目源代码仓库链接', True, BODY_SIZE),
            ('GitHub: https://github.com/zwl-crayfish/SwapCampus', False, BODY_SIZE),
            ('（含完整 commit 历史、CI 运行记录、PR Review 记录）', False, BODY_SIZE),
        ]
        fill_cell_with_paragraphs(cell, paras)
    sections[39] = fill_table39

    # TABLE[40]: Note — KEEP AS IS (do nothing)

    return sections


# ── Main ───────────────────────────────────────────────────

def main():
    print(f'Opening template: {TEMPLATE}')
    doc = Document(TEMPLATE)

    sections = get_content()

    for table_idx in sorted(sections.keys()):
        print(f'Filling TABLE[{table_idx}]...')
        try:
            sections[table_idx](doc)
        except Exception as e:
            print(f'  ERROR filling TABLE[{table_idx}]: {e}')
            import traceback
            traceback.print_exc()

    # Keep manual page breaks and section formatting intact
    # Save
    print(f'Saving to: {OUTPUT}')
    doc.save(OUTPUT)
    print('Done!')

if __name__ == '__main__':
    main()
